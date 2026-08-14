"""审计 4.1-4.7 / 3.7 / 3.8 的 Word 缺陷回归测试。

这里的 docx 全部用 python-docx 真实构造出对应的 XML 结构（w:hyperlink、嵌套表格、
w:sdt、w:ins、w:numPr），不做结构上的 mock——这批缺陷的成因恰恰是"python-docx 的
视图与真实 XML 不一致"，用假结构断言等于把 bug 一起假掉。
"""

from __future__ import annotations

import tempfile
import unittest
from contextlib import ExitStack
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml

from core.api_config_check import ApiConfigCheckResult
from core.model_throughput import EffectiveModelThroughput
from core.task_runner import ErrorMsg
from core.translation_protocol import REPLACE_TRANSLATION_PREFIX
from core.word_batching import WordBatchRunStats, _split_long_word_text
from core.word_document import (
    WordFileItem,
    detect_hidden_word_content,
    normalize_docx_automatic_numbering,
    write_bilingual_docx,
)
from core.word_task_runner import (
    WordTaskRunner,
    _WordRecoveryOutcome,
    _WordRecoveryPool,
    _word_cell_line_mismatch_issue,
)
from settings import AppSettings, WordBatchSettings
from tests.app_data_isolation import IsolatedAppDataTestCase


def _replace(text: str) -> str:
    return f"{REPLACE_TRANSLATION_PREFIX}{text}"


def _add_hyperlink(paragraph, text: str, url: str):
    """在段落末尾追加一个真正的 w:hyperlink（带外部关系），而不是伪造的 run。"""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    text_element.set(qn("xml:space"), "preserve")
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _set_num_pr(paragraph, *, num_id: str = "9", ilvl: str = "0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_element = OxmlElement("w:ilvl")
    ilvl_element.set(qn("w:val"), ilvl)
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), num_id)
    num_pr.append(ilvl_element)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)


def _set_num_id_9_level(doc, *, number_format: str, level_text_value: str) -> None:
    numbering_root = doc.part.numbering_part.element
    target_abstract_id = None
    for num in numbering_root.findall(qn("w:num")):
        if num.get(qn("w:numId")) == "9":
            abstract_id = num.find(qn("w:abstractNumId"))
            target_abstract_id = abstract_id.get(qn("w:val")) if abstract_id is not None else None
            break
    for abstract_num in numbering_root.findall(qn("w:abstractNum")):
        if abstract_num.get(qn("w:abstractNumId")) != target_abstract_id:
            continue
        for level in abstract_num.findall(qn("w:lvl")):
            if int(level.get(qn("w:ilvl")) or 0) != 0:
                continue
            num_format = level.find(qn("w:numFmt"))
            if num_format is None:
                num_format = OxmlElement("w:numFmt")
                level.append(num_format)
            num_format.set(qn("w:val"), number_format)
            level_text = level.find(qn("w:lvlText"))
            if level_text is None:
                level_text = OxmlElement("w:lvlText")
                level.append(level_text)
            level_text.set(qn("w:val"), level_text_value)
            return


def _paragraph_num_id(paragraph) -> str:
    p_pr = getattr(paragraph._p, "pPr", None)
    if p_pr is None:
        return ""
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return ""
    num_id = num_pr.find(qn("w:numId"))
    return num_id.get(qn("w:val")) if num_id is not None else ""


class WordHyperlinkParagraphTests(unittest.TestCase):
    """4.1 / 4.2：paragraph.runs 不含超链接 run，paragraph.text 却含。"""

    def test_replace_only_paragraph_with_hyperlink_leaves_no_source_text(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "hyperlink.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("See the ")
            _add_hyperlink(paragraph, "official site", "https://example.com")
            paragraph.add_run(" for details.")
            doc.save(source_path)

            source_text = "See the official site for details."
            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=root / "out",
                translations={source_text: _replace("请见官方网站了解详情。")},
                target_lang="zh",
                source_lang="en",
            )

            out_doc = Document(str(out_path))
            self.assertEqual(out_doc.paragraphs[0].text, "请见官方网站了解详情。")
            # 修复前这里是 '请见官方网站了解详情。official site'：超链接文字既没被
            # 译文覆盖，也没留在原位，而是被甩到了译文末尾。
            self.assertNotIn("official site", out_doc.paragraphs[0].text)
            self.assertEqual(
                len(out_doc.paragraphs[0]._p.findall(qn("w:hyperlink"))),
                0,
                "整段文字已被译文取代，残留的空超链接是一块看不见却可点击的区域",
            )

    def test_replace_only_paragraph_made_only_of_a_hyperlink(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "only_link.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            _add_hyperlink(paragraph, "Design specification", "https://example.com/spec")
            doc.save(source_path)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=root / "out",
                translations={"Design specification": _replace("设计规范")},
                target_lang="zh",
                source_lang="en",
            )
            out_doc = Document(str(out_path))
            self.assertEqual(out_doc.paragraphs[0].text, "设计规范")

    def test_numbering_label_lands_before_a_leading_hyperlink(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "numbered_link.docx"
            doc = Document()
            _set_num_id_9_level(doc, number_format="decimal", level_text_value="%1.")
            paragraph = doc.add_paragraph()
            _add_hyperlink(paragraph, "Reference standard", "https://example.com/std")
            paragraph.add_run(" applies to this section.")
            _set_num_pr(paragraph)
            doc.save(source_path)

            normalized = normalize_docx_automatic_numbering(source_path)
            text = Document(str(normalized.path)).paragraphs[0].text
            # 修复前：标签写进 paragraph.runs[0]（超链接之后的那个 run），
            # 变成 'Reference standard 1. applies to this section.'
            self.assertTrue(
                text.startswith("1. Reference standard"),
                f"编号标签没有落在段首：{text!r}",
            )


class WordDocumentOrderTests(unittest.TestCase):
    """4.3：doc.paragraphs + 表格段落的拼接顺序不是文档真实顺序。"""

    def test_numbering_counts_follow_real_document_order(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "mixed_order.docx"
            doc = Document()
            _set_num_id_9_level(doc, number_format="decimal", level_text_value="%1.")
            first = doc.add_paragraph("Body item before the table")
            _set_num_pr(first)
            table = doc.add_table(rows=1, cols=1)
            cell_paragraph = table.rows[0].cells[0].paragraphs[0]
            cell_paragraph.add_run("Table item in the middle")
            _set_num_pr(cell_paragraph)
            last = doc.add_paragraph("Body item after the table")
            _set_num_pr(last)
            doc.save(source_path)

            normalized = normalize_docx_automatic_numbering(source_path)
            out_doc = Document(str(normalized.path))
            body_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
            table_text = out_doc.tables[0].rows[0].cells[0].paragraphs[0].text
            # 修复前表格段落被排到全部正文之后，拿到的是 3.，而表格后的正文拿到 2.
            self.assertEqual(body_texts[0], "1. Body item before the table")
            self.assertEqual(table_text, "2. Table item in the middle")
            self.assertEqual(body_texts[1], "3. Body item after the table")


class WordFrontMatterProtectionTests(unittest.TestCase):
    """4.4：保护边界算了，但扁平化对全文执行。"""

    def test_protected_front_matter_keeps_automatic_numbering_untouched(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "front_matter.docx"
            doc = Document()
            _set_num_id_9_level(doc, number_format="decimal", level_text_value="%1.")
            cover = doc.add_paragraph("Project cover list entry")
            _set_num_pr(cover)
            doc.add_paragraph("第一章 工程概况")
            body = doc.add_paragraph("Body list entry")
            _set_num_pr(body)
            doc.save(source_path)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=root / "out",
                translations={},
                target_lang="en",
                source_lang="zh",
                protect_front_matter=True,
            )
            out_doc = Document(str(out_path))
            protected, _heading, translated_body = out_doc.paragraphs[:3]

            self.assertEqual(protected.text, "Project cover list entry")
            self.assertEqual(
                _paragraph_num_id(protected),
                "9",
                "受保护段落的 numPr 被删掉了——那也是一次改写",
            )
            # 保护边界之外仍按原设计扁平化。
            self.assertEqual(translated_body.text, "2. Body list entry")
            self.assertEqual(_paragraph_num_id(translated_body), "0")

    def test_unprotected_run_still_flattens_the_whole_document(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "front_matter.docx"
            doc = Document()
            _set_num_id_9_level(doc, number_format="decimal", level_text_value="%1.")
            cover = doc.add_paragraph("Project cover list entry")
            _set_num_pr(cover)
            doc.add_paragraph("第一章 工程概况")
            doc.save(source_path)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=root / "out",
                translations={},
                target_lang="en",
                source_lang="zh",
                protect_front_matter=False,
            )
            out_doc = Document(str(out_path))
            self.assertEqual(out_doc.paragraphs[0].text, "1. Project cover list entry")


class WordTableCellReplaceTests(unittest.TestCase):
    """4.5：replace_only 用 cell.text 赋值会压平整格。"""

    def test_replace_only_cell_keeps_nested_table_and_paragraph_structure(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "nested_cell.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].add_run("First line")
            second = cell.add_paragraph()
            bold_run = second.add_run("Second line")
            bold_run.bold = True
            nested = cell.add_table(rows=1, cols=1)
            nested.rows[0].cells[0].paragraphs[0].add_run("Nested cell text")
            doc.save(source_path)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=root / "out",
                translations={
                    "First line\nSecond line": _replace("第一行\n第二行"),
                    "Nested cell text": _replace("嵌套单元格文本"),
                },
                target_lang="zh",
                source_lang="en",
            )

            source_cell = Document(str(source_path)).tables[0].rows[0].cells[0]
            source_layout = [
                child.tag
                for child in source_cell._tc.iterchildren()
                if child.tag in {qn("w:p"), qn("w:tbl")}
            ]

            out_doc = Document(str(out_path))
            out_cell = out_doc.tables[0].rows[0].cells[0]
            out_layout = [
                child.tag
                for child in out_cell._tc.iterchildren()
                if child.tag in {qn("w:p"), qn("w:tbl")}
            ]
            self.assertEqual(
                out_layout,
                source_layout,
                "修复前整格被压成一个纯文本段落，段落与嵌套表格的排布全部丢失",
            )
            self.assertEqual(
                len(out_cell.tables),
                1,
                "嵌套表格被 cell.text 赋值整块删掉了",
            )
            self.assertEqual(
                [paragraph.text for paragraph in out_cell.paragraphs][:2],
                ["第一行", "第二行"],
            )
            self.assertTrue(out_cell.paragraphs[1].runs[0].bold)
            self.assertEqual(
                out_cell.tables[0].rows[0].cells[0].paragraphs[0].text,
                "嵌套单元格文本",
            )

    def test_replace_only_cell_with_mismatched_line_count_keeps_source_and_appends(self) -> None:
        """行数不齐（模型丢了第三行）时保留全部原文、译文整体追加，并向报告留痕。

        修复前的回退是"译文全塞第一段、其余源文段清空"：译文没覆盖到的第三段
        原文被无声抹掉，报告里一个字都不提。
        """
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source_path = root / "mismatch_cell.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].add_run("First line")
            cell.add_paragraph().add_run("Second line")
            cell.add_paragraph().add_run("Third line")
            doc.save(source_path)

            logs: list[str] = []
            issues: list[dict] = []
            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=root / "out",
                translations={
                    "First line\nSecond line\nThird line": _replace("第一行\n第二行")
                },
                target_lang="zh",
                source_lang="en",
                log_callback=logs.append,
                issue_callback=issues.append,
            )
            out_cell = Document(str(out_path)).tables[0].rows[0].cells[0]
            texts = [paragraph.text for paragraph in out_cell.paragraphs]
            self.assertEqual(
                texts[:3],
                ["First line", "Second line", "Third line"],
                "行数不齐时源文段被清空——译文没覆盖到的内容就此丢失",
            )
            self.assertEqual(texts[-1], "第一行\n第二行", "译文应整体追加在单元格末尾")
            self.assertEqual(len(issues), 1, "行数不齐必须向报告层留痕")
            self.assertEqual(issues[0]["location"], "table[0].cell[0]")
            self.assertEqual(issues[0]["source"], "First line\nSecond line\nThird line")
            self.assertEqual(issues[0]["translation"], "第一行\n第二行")
            self.assertTrue(
                any("行数与原文段数不齐" in line for line in logs),
                f"日志汇总也要提到行数不齐：{logs}",
            )

    def test_cell_line_mismatch_issue_matches_report_shape(self) -> None:
        issue = _word_cell_line_mismatch_issue(
            file_name="a.docx",
            info={
                "location": "table[0].cell[2]",
                "source": "甲\n乙\n丙",
                "translation": "A\nB",
            },
        )
        self.assertEqual(issue["file"], "a.docx")
        self.assertEqual(issue["kind"], "table_cell")
        self.assertEqual(issue["location_label"], "表格 1 / 单元格 3")
        self.assertEqual(issue["severity"], "needs_review")
        self.assertIn("行数不一致", issue["problem"])
        self.assertIn("甲", issue["snippet"])


class WordSentencePackingTests(unittest.TestCase):
    """4.6：打包时 "".join 丢掉句间分隔符。"""

    def test_packing_preserves_original_separators(self) -> None:
        self.assertEqual(
            _split_long_word_text("A.\nB.\nC.", split_threshold=1, part_char_budget=100),
            ["A.\nB.\nC."],
        )

    def test_packing_preserves_spaces_between_sentences(self) -> None:
        text = "First sentence. Second sentence. Third sentence."
        self.assertEqual(
            _split_long_word_text(text, split_threshold=1, part_char_budget=200),
            [text],
        )

    def test_oversized_text_still_splits_within_budget(self) -> None:
        text = ("Sentence number one is here. " * 200).strip()
        parts = _split_long_word_text(text, split_threshold=3000, part_char_budget=1500)
        self.assertGreater(len(parts), 1)
        for part in parts:
            self.assertLessEqual(len(part), 1500)
        self.assertEqual("".join(part.strip() for part in parts).replace(" ", ""),
                         text.replace(" ", "").replace("\n", ""))


class WordHiddenContentDetectionTests(unittest.TestCase):
    """4.7：w:sdt / w:ins 包住的内容对 python-docx 完全不可见。"""

    @staticmethod
    def _build_hidden_content_docx(path: Path) -> None:
        doc = Document()
        doc.add_paragraph("Visible body paragraph.")
        section_properties = doc.element.body.find(qn("w:sectPr"))
        section_properties.addprevious(
            parse_xml(
                f'<w:sdt {nsdecls("w")}><w:sdtPr/><w:sdtContent>'
                "<w:p><w:r><w:t>Content control paragraph.</w:t></w:r></w:p>"
                "</w:sdtContent></w:sdt>"
            )
        )
        section_properties.addprevious(
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:r><w:t xml:space="preserve">Kept text. </w:t></w:r>'
                '<w:ins w:id="99" w:author="tester" w:date="2026-01-01T00:00:00Z">'
                "<w:r><w:t>Tracked inserted sentence.</w:t></w:r></w:ins></w:p>"
            )
        )
        doc.save(path)

    def test_content_controls_and_tracked_insertions_are_reported(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "hidden.docx"
            self._build_hidden_content_docx(path)

            report = detect_hidden_word_content(path)
            self.assertEqual(report.content_control_count, 1)
            self.assertEqual(report.tracked_insertion_count, 1)
            self.assertTrue(report.found)
            self.assertIn("内容控件", report.describe())
            self.assertIn("修订", report.describe())

    def test_plain_document_reports_nothing(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "plain.docx"
            doc = Document()
            doc.add_paragraph("Only ordinary content here.")
            doc.save(path)

            report = detect_hidden_word_content(path)
            self.assertFalse(report.found)
            self.assertEqual(report.total, 0)

    def test_paragraph_mark_insertions_are_not_reported(self) -> None:
        """w:pPr/w:rPr/w:ins 只标记段落标记本身，没有任何正文被藏起来。"""
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "mark_only.docx"
            doc = Document()
            doc.add_paragraph("Ordinary paragraph.")
            p_pr = doc.paragraphs[0]._p.get_or_add_pPr()
            p_pr.append(
                parse_xml(
                    f'<w:rPr {nsdecls("w")}>'
                    '<w:ins w:id="7" w:author="tester" w:date="2026-01-01T00:00:00Z"/>'
                    "</w:rPr>"
                )
            )
            doc.save(path)

            self.assertFalse(detect_hidden_word_content(path).found)


class WordRecoveryPoolShutdownTests(unittest.TestCase):
    """3.8：executor 只在 wait_for_completion 内部 shutdown，异常路径全漏。"""

    @staticmethod
    def _pool() -> _WordRecoveryPool:
        return _WordRecoveryPool(
            engine=SimpleNamespace(engine_name="test/pool"),
            target_lang="en",
            retry_prompt="retry",
            retry_batch_settings=WordBatchSettings(max_paragraphs_per_batch=1),
            retry_attempts=1,
            source_lang="zh",
            api_scheduler=None,
            concurrency=3,
            should_stop=lambda: False,
        )

    def test_executor_is_shut_down_when_the_wait_loop_raises(self) -> None:
        """等待循环里抛出时，HEAD 版本一次 shutdown 都跑不到——泄漏 concurrency 个线程。"""
        pool = self._pool()
        with patch.object(
            _WordRecoveryPool,
            "_all_complete_locked",
            side_effect=RuntimeError("boom"),
        ):
            with self.assertRaises(RuntimeError):
                pool.wait_for_completion()
        self.assertTrue(pool._executor._shutdown)

    def test_executor_is_shut_down_on_the_normal_path(self) -> None:
        pool = self._pool()
        pool.wait_for_completion()
        self.assertTrue(pool._executor._shutdown)

    def test_shutdown_is_idempotent_and_usable_as_a_context_manager(self) -> None:
        pool = self._pool()
        with pool:
            pass
        pool.shutdown()
        self.assertTrue(pool._executor._shutdown)


class WordRunFallbackTests(IsolatedAppDataTestCase):
    """3.7：_run 无兜底 except/finally，写回异常 → 线程静默死亡、任务永久挂起。"""

    @staticmethod
    def _settings() -> AppSettings:
        return AppSettings(source_lang="zh", target_lang="en")

    def _runner_patches(self, stack: ExitStack, *, root: Path) -> None:
        stack.enter_context(
            patch("core.word_task_runner.TaskLogger", return_value=MagicMock(task_id="defect-fixes"))
        )
        stack.enter_context(
            patch(
                "core.word_task_runner.check_translation_api_config",
                return_value=ApiConfigCheckResult(ok=True),
            )
        )
        stack.enter_context(
            patch(
                "core.word_task_runner.build_engine",
                return_value=SimpleNamespace(engine_name="defect/mock"),
            )
        )
        stack.enter_context(patch("core.word_task_runner.get_system_prompt", return_value="system"))
        stack.enter_context(
            patch("core.word_task_runner.resolve_effective_model_config", return_value=object())
        )
        stack.enter_context(
            patch(
                "core.word_task_runner.get_model_throughput",
                return_value=EffectiveModelThroughput(
                    profile_key="defect", batch_size=10, concurrency=1
                ),
            )
        )
        stack.enter_context(
            patch("core.word_task_runner.build_word_output_dir", return_value=root / "out")
        )

    def test_unexpected_failure_still_emits_a_terminal_message_and_cleans_up(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("施工范围")
            document.save(source)
            leaked_temp = root / "converted_temp.docx"
            leaked_temp.write_bytes(b"temporary conversion artifact")

            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(leaked_temp,),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                self._settings(),
                source_root=root,
            )
            with ExitStack() as stack:
                self._runner_patches(stack, root=root)
                stack.enter_context(
                    patch(
                        "core.word_task_runner._prepare_word_source_for_translation",
                        return_value=prepared,
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.lookup_batch",
                        side_effect=PermissionError("disk went away"),
                    )
                )
                # 修复前这里会把异常直接抛出工作线程。
                runner._run()

            messages = list(runner._queue.queue)
            errors = [item for item in messages if isinstance(item, ErrorMsg)]
            self.assertEqual(len(errors), 1, "队列必须收到且只收到一条终止消息")
            # 终止横幅上只出现中文说明，PermissionError 的原文留给 debug 日志。
            self.assertIn("Word 翻译任务异常中止", errors[0].message)
            self.assertNotIn("disk went away", errors[0].message)
            self.assertFalse(
                leaked_temp.exists(),
                "临时 docx 必须在 finally 里被清理，否则每次异常都留一份垃圾",
            )

    def test_recovery_pool_is_shut_down_when_the_main_translation_raises(self) -> None:
        """3.8 审计描述的那条路径：主翻译抛异常，wait_for_completion 根本没被调用。"""
        created_pools: list[_WordRecoveryPool] = []

        def _make_pool(**kwargs):
            pool = _WordRecoveryPool(**kwargs)
            created_pools.append(pool)
            return pool

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("施工范围")
            document.save(source)

            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                self._settings(),
                source_root=root,
            )
            with ExitStack() as stack:
                self._runner_patches(stack, root=root)
                stack.enter_context(
                    patch(
                        "core.word_task_runner._prepare_word_source_for_translation",
                        return_value=prepared,
                    )
                )
                # TM 全不命中 → 走云端翻译 → 恢复池被建出来。
                stack.enter_context(
                    patch("core.word_task_runner.tm_manager.lookup_batch", return_value={})
                )
                stack.enter_context(
                    patch("core.word_task_runner.tm_manager.insert_batch", return_value=0)
                )
                stack.enter_context(
                    patch("core.word_task_runner._WordRecoveryPool", side_effect=_make_pool)
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.translate_word_texts",
                        side_effect=RuntimeError("upstream exploded"),
                    )
                )
                runner._run()

            self.assertEqual(len(created_pools), 1)
            self.assertTrue(
                created_pools[0]._executor._shutdown,
                "主翻译抛异常后线程池仍活着，每次任务泄漏 concurrency 个非守护线程",
            )
            errors = [item for item in runner._queue.queue if isinstance(item, ErrorMsg)]
            self.assertEqual(len(errors), 1)
            self.assertIn("Word 翻译任务异常中止", errors[0].message)
            self.assertNotIn("upstream exploded", errors[0].message)

    def test_terminal_message_survives_a_failure_while_building_the_result(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("施工范围")
            document.save(source)

            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                self._settings(),
                source_root=root,
            )
            with ExitStack() as stack:
                self._runner_patches(stack, root=root)
                stack.enter_context(
                    patch(
                        "core.word_task_runner._prepare_word_source_for_translation",
                        return_value=prepared,
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.lookup_batch",
                        side_effect=lambda texts, _pair: {text: "TM translation" for text in texts},
                    )
                )
                stack.enter_context(
                    patch("core.word_task_runner.tm_manager.insert_batch", return_value=0)
                )
                stack.enter_context(
                    patch("core.word_task_runner.write_bilingual_docx", return_value=root / "out.docx")
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner._append_post_write_coverage_issues",
                        return_value=0,
                    )
                )
                stack.enter_context(
                    patch.object(
                        WordTaskRunner,
                        "_build_result_contract",
                        side_effect=RuntimeError("contract exploded"),
                    )
                )
                runner._run()

            errors = [item for item in runner._queue.queue if isinstance(item, ErrorMsg)]
            self.assertEqual(len(errors), 1)
            self.assertIn("Word 翻译任务收尾失败", errors[0].message)
            self.assertNotIn("contract exploded", errors[0].message)


class WordHiddenContentWarningTests(IsolatedAppDataTestCase):
    """4.7 的用户可见部分：漏译必须留下告警，不能静悄悄。"""

    def test_runner_warns_about_content_controls_and_tracked_insertions(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "hidden.docx"
            WordHiddenContentDetectionTests._build_hidden_content_docx(source)

            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                AppSettings(source_lang="en", target_lang="zh"),
                source_root=root,
            )
            with ExitStack() as stack:
                WordRunFallbackTests()._runner_patches(stack, root=root)
                stack.enter_context(
                    patch(
                        "core.word_task_runner._prepare_word_source_for_translation",
                        return_value=prepared,
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.lookup_batch",
                        side_effect=lambda texts, _pair: {text: "译文" for text in texts},
                    )
                )
                stack.enter_context(
                    patch("core.word_task_runner.tm_manager.insert_batch", return_value=0)
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.write_bilingual_docx",
                        return_value=root / "out.docx",
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner._append_post_write_coverage_issues",
                        return_value=0,
                    )
                )
                runner._run()

            warnings = [
                message.message
                for message in runner._queue.queue
                if getattr(message, "level", "") == "WARN"
            ]
            self.assertTrue(
                any("未翻译内容" in text and "内容控件" in text for text in warnings),
                f"没有发出漏译告警：{warnings}",
            )


class WordReviewPositionCountTests(IsolatedAppDataTestCase):
    """走查发现：同一段落的两条判定被数成两处需复核，用户拿着数字去文档里找不到第二处。"""

    def _contract(self, quality_issues: list[dict], source: Path, root: Path) -> dict:
        runner = WordTaskRunner(
            [WordFileItem(path=source, name=source.name, size_kb=1.0)],
            AppSettings(source_lang="zh", target_lang="en"),
            source_root=root,
        )
        return runner._build_result_contract(
            file_results=[{"source_path": str(source), "success": True, "output": str(root / "out.docx")}],
            output_dir=str(root / "out"),
            elapsed_sec=1.0,
            tm_hit_count=0,
            api_text_count=1,
            source_lang="zh",
            target_lang="en",
            preflights={},
            file_texts=[set()],
            quality_issues=quality_issues,
            recovery_outcome=_WordRecoveryOutcome(
                fixed_sources=[],
                unresolved_sources=[],
                accepted_translations={},
                recovery_review_results={},
                semantic_review_results={},
                unresolved_validation_results={},
            ),
            word_batch_stats=WordBatchRunStats(),
            model_source_results={},
            stopped=False,
        )

    def test_two_judgments_on_one_paragraph_count_as_one_position(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("施工范围")
            document.save(source)

            issues = [
                {
                    "file": "source.docx",
                    "section_path": "正文",
                    "location_label": "正文段落 5",
                    "snippet": "施工范围……安装",
                    "problem": "重试后仍未获得有效译文",
                    "status": "保留原文，待人工复核",
                    "severity": "needs_review",
                },
                {
                    # 写出之后的覆盖率复查对同一段再记一条，说的是同一件事。
                    "file": "source.docx",
                    "section_path": "正文",
                    "location_label": "正文段落 5",
                    "snippet": "施工范围……安装",
                    "problem": "输出文档仍存在未译源文",
                    "status": "保留原文，待人工复核",
                    "severity": "needs_review",
                },
                {
                    "file": "source.docx",
                    "section_path": "正文",
                    "location_label": "正文段落 9",
                    "snippet": "设备安装",
                    "problem": "重试后仍未获得有效译文",
                    "status": "保留原文，待人工复核",
                    "severity": "needs_review",
                },
            ]
            contract = self._contract(issues, source, root)
            review = contract["review"]

            self.assertEqual(review["total_count"], 2)
            self.assertEqual(review["counts"], {"needs_review": 2})
            self.assertEqual(contract["kpi"]["review_text_count"], 2)
            self.assertEqual(len(review["items"]), 2)

            merged = next(item for item in review["items"] if item["location"] == "正文段落 5")
            # 合并的是行数，不是内容：两句判定都要留在这一行里。
            self.assertIn("重试后仍未获得有效译文", str(merged["problem"]))
            self.assertIn("输出文档仍存在未译源文", str(merged["problem"]))
            self.assertEqual(merged["action"], "保留原文，待人工复核")

    def test_resolved_and_needs_review_at_one_position_stay_separate(self) -> None:
        # 同一段既有「已自动处理」又有「待复核」时不能并成一行：颜色和结论相反，
        # 并起来等于把一条待办涂成绿色。
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("施工范围")
            document.save(source)

            issues = [
                {
                    "file": "source.docx",
                    "location_label": "正文段落 5",
                    "snippet": "施工范围",
                    "problem": "首次未获得译文",
                    "status": "严格重试已恢复译文",
                    "severity": "resolved",
                },
                {
                    "file": "source.docx",
                    "location_label": "正文段落 5",
                    "snippet": "施工范围",
                    "problem": "译文长度异常",
                    "status": "保留原文，待人工复核",
                    "severity": "needs_review",
                },
            ]
            review = self._contract(issues, source, root)["review"]
            self.assertEqual(review["total_count"], 2)
            self.assertEqual(review["counts"], {"resolved": 1, "needs_review": 1})


if __name__ == "__main__":
    unittest.main()
