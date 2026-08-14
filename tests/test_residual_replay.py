# -*- coding: utf-8 -*-
"""双语文档回放体检的回归（设计文档 §六.3 的合成语料版）。

真实的 9 文档语料在用户机器上、不入库，这里用合成 docx/xlsx 复刻同样的
失败形态：序号残留（含同级惯例投票）、术语尾巴、节标题两套写法、干净段
0 误报。真实语料的全量回放在发布验收时人工跑 scripts/replay_residual_check.py。
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from core.residual_classifier import (
    CATEGORY_NUMBERING_PREFIX,
    CATEGORY_TERM_FRAGMENT,
    CONVENTION_PAREN_ROMAN,
)
from core.residual_replay import replay_file
from core.unit_ledger import STATE_CLEAN, STATE_NEEDS_REVIEW


def _build_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    # 干净配对段 ×2
    doc.add_paragraph("对地面裂缝进行注浆修复。")
    doc.add_paragraph("Réparation des fissures du sol par injection.")
    doc.add_paragraph("修复完成后进行养护。")
    doc.add_paragraph("Assurer la cure après la réparation.")
    # （X）族同级：三段译好（(I)(II)(III) → paren_roman 惯例），第四段序号残留
    doc.add_paragraph("（一）表层裂缝")
    doc.add_paragraph("(I) Fissures superficielles")
    doc.add_paragraph("（二）结构裂缝")
    doc.add_paragraph("(II) Fissures structurelles")
    doc.add_paragraph("（三）贯穿裂缝")
    doc.add_paragraph("(III) Fissures traversantes")
    doc.add_paragraph("（四）沉降裂缝")
    doc.add_paragraph("（四）Fissures de tassement")
    # 术语尾巴复读
    doc.add_paragraph("沿裂缝开V型槽并清理浮灰。")
    doc.add_paragraph("Ouvrir une rainure en V 型槽 le long de la fissure et nettoyer.")
    # 节标题：多数派 Section N，一处离群
    doc.add_paragraph("第一节 总则")
    doc.add_paragraph("Section 1 Dispositions générales")
    doc.add_paragraph("第二节 工程概况")
    doc.add_paragraph("Deuxième section — Aperçu des travaux")
    doc.add_paragraph("第三节 施工方案")
    doc.add_paragraph("Section 3 Plan d'exécution")
    doc.save(str(path))


def _build_xlsx(path: Path) -> None:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "工程量"
    # 同格双语（本项目 Excel 输出排版）：干净 ×1、序号残留 ×1
    sheet["A1"] = "混凝土浇筑\nCoulage du béton"
    sheet["B2"] = "（二）钢筋绑扎\n（二）Ligature des armatures"
    sheet["C3"] = "单价"  # 纯源文（未翻译单元格，不算配对）
    workbook.save(str(path))


class DocxReplayTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls._tmp = tempfile.TemporaryDirectory()
        cls.docx_path = Path(cls._tmp.name) / "sample_bilingual.docx"
        _build_docx(cls.docx_path)
        cls.result = replay_file(cls.docx_path, target_lang="fr")

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_pairing_and_convention(self):
        self.assertEqual(self.result.kind, "docx")
        self.assertEqual(self.result.pair_count, 10)
        self.assertEqual(self.result.convention, CONVENTION_PAREN_ROMAN)

    def test_findings_cover_known_failures_only(self):
        categories = sorted(f.category for f in self.result.findings)
        self.assertEqual(
            categories, [CATEGORY_NUMBERING_PREFIX, CATEGORY_TERM_FRAGMENT]
        )
        numbering = next(
            f for f in self.result.findings
            if f.category == CATEGORY_NUMBERING_PREFIX
        )
        # 确定性修复对齐同级惯例 (I)(II)(III)
        self.assertEqual(
            numbering.deterministic_fix, "(IV) Fissures de tassement"
        )

    def test_heading_outlier_detected_with_fix(self):
        self.assertEqual(len(self.result.heading_outliers), 1)
        outlier = self.result.heading_outliers[0]
        self.assertEqual(outlier.fix, "Section 2 Aperçu des travaux")

    def test_ledger_states(self):
        counts = self.result.ledger.counts()
        # 10 配对段：2 残留 + 1 标题离群 = 3 needs_review，其余 clean
        self.assertEqual(counts[STATE_NEEDS_REVIEW], 3)
        self.assertEqual(counts[STATE_CLEAN], 7)
        self.assertEqual(self.result.ledger.unresolved(), [])

    def test_clean_paragraphs_produce_no_findings(self):
        flagged = {f.source_anchor for f in self.result.findings}
        clean_units = [
            r for r in self.result.ledger.records() if r.state == STATE_CLEAN
        ]
        for record in clean_units:
            self.assertNotIn(record.source_anchor, flagged)


class XlsxReplayTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls._tmp = tempfile.TemporaryDirectory()
        cls.xlsx_path = Path(cls._tmp.name) / "sample_bilingual.xlsx"
        _build_xlsx(cls.xlsx_path)
        cls.result = replay_file(cls.xlsx_path, target_lang="fr")

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_pairing_uses_sheet_coordinates(self):
        self.assertEqual(self.result.kind, "xlsx")
        anchors = sorted(
            record.source_anchor for record in self.result.ledger.records()
        )
        self.assertEqual(anchors, ["工程量!A1", "工程量!B2"])

    def test_residual_cell_flagged_clean_cell_not(self):
        self.assertEqual(len(self.result.findings), 1)
        finding = self.result.findings[0]
        self.assertEqual(finding.source_anchor, "工程量!B2")
        self.assertEqual(finding.category, CATEGORY_NUMBERING_PREFIX)
        self.assertEqual(
            self.result.ledger.get("工程量!A1").state, STATE_CLEAN
        )


class UnsupportedTypeTest(unittest.TestCase):
    def test_unsupported_suffix_raises(self):
        with self.assertRaises(ValueError):
            replay_file("whatever.txt", target_lang="fr")


if __name__ == "__main__":
    unittest.main()
