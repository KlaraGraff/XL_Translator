"""补译模式「这一对真的是原文＋译文吗」复核的回归。

判错的代价不对称：判成"已有译文"而其实不是，那段中文就永远留在文档里，体检也发现
不了（体检用的是同一套启发式）；判反了顶多多插一条译文，看得见、删得掉。所以下面
每一条都在守同一件事——宁可多翻，不可漏翻，同时别为此白烧模型调用。
"""

from __future__ import annotations

import tempfile
import threading
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

from docx import Document

from core import word_task_runner

from core.coverage_arbitration import (
    RETRANSLATE_MODEL,
    RETRANSLATE_UNCERTAIN,
    TRUST_KNOWN_TRANSLATION,
    TRUST_LENGTH_RATIO,
    TRUST_MODEL,
    TRUST_SHORT_TOKEN,
    apply_arbitration,
    collect_arbitration_candidates,
    review_coverage_pairs,
)
from core.translation_coverage import (
    COVERAGE_COVERED,
    COVERAGE_IGNORED,
    COVERAGE_SOURCE_ONLY,
    CoverageUnit,
    looks_like_foreign_acronym,
    looks_like_target_text,
)
from core.word_coverage import build_word_coverage_plan, write_untranslated_docx

# 一段真实长度的中文，和一条明显只翻了个开头的"译文"——长度比 0.2，可疑。
LONG_SOURCE = (
    "本工程受甲供短柱供货严重滞后影响，土建结构施工无法按原计划穿插进行，"
    "经与监理及业主协商，工期顺延一百六十九天。"
)
PARTIAL_TARGET = "Retard de livraison."
FULL_TARGET = (
    "En raison du retard important de livraison des potelets courts fournis par le "
    "maître d’ouvrage, les travaux de structure n’ont pas pu être menés selon le "
    "planning initial ; après concertation avec le maître d’œuvre et le maître "
    "d’ouvrage, le délai est prolongé de cent soixante-neuf jours."
)


def _pair(source: str, target: str, *, index: int = 0) -> CoverageUnit:
    return CoverageUnit(
        source_text=source,
        target_text=target,
        status=COVERAGE_COVERED,
        location=f"body.paragraph[{index}]",
        kind="paragraph",
        reason="下一段为目标语言译文。",
        data={"paragraph_index": index},
    )


class ForeignAcronymTests(unittest.TestCase):
    def test_all_caps_abbreviations_count_as_translation(self) -> None:
        for token in ("CCTEB", "ONEE", "SARL", "PV", "N/A", "BTR-ANODE-CCTEB-032"):
            self.assertTrue(looks_like_foreign_acronym(token), token)
            self.assertTrue(
                looks_like_target_text(token, source_lang="zh", target_lang="fr"),
                token,
            )

    def test_ordinary_words_and_chinese_are_not_abbreviations(self) -> None:
        for token in ("Béton", "Travaux", "施工单位", "A", "工程"):
            self.assertFalse(looks_like_foreign_acronym(token), token)


class CandidateSelectionTests(unittest.TestCase):
    def test_only_covered_paragraph_pairs_are_reviewed(self) -> None:
        units = [
            _pair("工程概况", "Présentation du projet", index=0),
            CoverageUnit(
                source_text="未译内容",
                status=COVERAGE_SOURCE_ONLY,
                location="body.paragraph[2]",
                kind="paragraph",
                reason="",
            ),
            CoverageUnit(
                source_text="表格原文",
                target_text="Texte du tableau",
                status=COVERAGE_COVERED,
                location="table[0].cell[0]",
                kind="table_cell",
                reason="",
            ),
            CoverageUnit(
                source_text="目录",
                status=COVERAGE_IGNORED,
                location="body.paragraph[3]",
                kind="paragraph",
                reason="",
            ),
        ]

        candidates = collect_arbitration_candidates(units)

        self.assertEqual([unit.location for unit in candidates], ["body.paragraph[0]"])


class CheapRuleTests(unittest.TestCase):
    """免费规则命中的对，一次模型调用都不能花。"""

    def test_short_foreign_token_is_trusted_without_the_model(self) -> None:
        calls: list[tuple[str, str]] = []

        def arbitrate(source: str, candidate: str) -> str:
            calls.append((source, candidate))
            return "not_equivalent"

        outcome = review_coverage_pairs(
            [_pair("序号", "N°"), _pair("施工单位", "CCTEB")],
            arbitrate=arbitrate,
        )

        self.assertEqual(calls, [])
        self.assertEqual(outcome.model_check_count, 0)
        self.assertEqual(outcome.retranslated, [])
        self.assertEqual(
            {review.reason for review in outcome.reviews}, {TRUST_SHORT_TOKEN}
        )

    def test_known_translation_match_is_trusted_without_the_model(self) -> None:
        calls: list[tuple[str, str]] = []

        def arbitrate(source: str, candidate: str) -> str:
            calls.append((source, candidate))
            return "not_equivalent"

        outcome = review_coverage_pairs(
            [_pair(LONG_SOURCE, PARTIAL_TARGET)],
            # 记忆库里就是这条译名——译名是特定的，比对字符串就够了。
            known_translations={LONG_SOURCE: PARTIAL_TARGET},
            arbitrate=arbitrate,
        )

        self.assertEqual(calls, [])
        self.assertEqual(outcome.reviews[0].reason, TRUST_KNOWN_TRANSLATION)

    def test_known_translation_comparison_ignores_case_and_spacing(self) -> None:
        outcome = review_coverage_pairs(
            [_pair(LONG_SOURCE, "  retard   DE livraison.  ")],
            known_translations={LONG_SOURCE: PARTIAL_TARGET},
            arbitrate=lambda source, candidate: "not_equivalent",
        )

        self.assertEqual(outcome.reviews[0].reason, TRUST_KNOWN_TRANSLATION)

    def test_known_translation_mismatch_still_goes_to_the_model(self) -> None:
        outcome = review_coverage_pairs(
            [_pair(LONG_SOURCE, PARTIAL_TARGET)],
            known_translations={LONG_SOURCE: FULL_TARGET},
            arbitrate=lambda source, candidate: "not_equivalent",
        )

        self.assertEqual(outcome.model_check_count, 1)
        self.assertEqual(outcome.reviews[0].reason, RETRANSLATE_MODEL)

    def test_normal_length_ratio_is_trusted_without_the_model(self) -> None:
        calls: list[str] = []

        outcome = review_coverage_pairs(
            [_pair(LONG_SOURCE, FULL_TARGET)],
            arbitrate=lambda source, candidate: calls.append(source) or "not_equivalent",
        )

        self.assertEqual(calls, [])
        self.assertEqual(outcome.reviews[0].reason, TRUST_LENGTH_RATIO)

    def test_a_whole_paragraph_followed_by_a_stray_marker_still_goes_to_the_model(
        self,
    ) -> None:
        """整段中文后面跟着一个孤零零的 II / PV，是排版残留，不是它的译文。"""
        for stray in ("II", "PV", "N°", "kg"):
            with self.subTest(stray=stray):
                outcome = review_coverage_pairs(
                    [_pair(LONG_SOURCE, stray)],
                    arbitrate=lambda source, candidate: "not_equivalent",
                )

                self.assertEqual(outcome.model_check_count, 1)
                self.assertEqual(len(outcome.retranslated), 1)

    def test_short_source_never_uses_the_length_ratio_as_a_suspicion(self) -> None:
        """"工期" vs "Délai" 比值本来就飘，短原文不靠比值判断。"""
        outcome = review_coverage_pairs(
            [_pair("工期", "Délai")],
            arbitrate=lambda source, candidate: "not_equivalent",
        )

        self.assertEqual(outcome.model_check_count, 0)
        self.assertEqual(outcome.reviews[0].reason, TRUST_LENGTH_RATIO)


class ModelVerdictTests(unittest.TestCase):
    def test_equivalent_keeps_the_pair_covered(self) -> None:
        outcome = review_coverage_pairs(
            [_pair(LONG_SOURCE, PARTIAL_TARGET)],
            arbitrate=lambda source, candidate: "equivalent",
        )

        self.assertEqual(outcome.reviews[0].reason, TRUST_MODEL)
        self.assertEqual(apply_arbitration(outcome), [])
        self.assertEqual(outcome.reviews[0].unit.status, COVERAGE_COVERED)

    def test_not_equivalent_flips_the_pair_back_to_untranslated(self) -> None:
        unit = _pair(LONG_SOURCE, PARTIAL_TARGET)
        outcome = review_coverage_pairs(
            [unit], arbitrate=lambda source, candidate: "not_equivalent"
        )

        flipped = apply_arbitration(outcome)

        self.assertEqual(flipped, [unit])
        self.assertEqual(unit.status, COVERAGE_SOURCE_ONLY)
        self.assertEqual(unit.data["arbitration"], RETRANSLATE_MODEL)

    def test_uncertain_also_retranslates(self) -> None:
        """拿不准就翻——漏翻是看不见的，多翻一条是看得见的。"""
        unit = _pair(LONG_SOURCE, PARTIAL_TARGET)
        outcome = review_coverage_pairs(
            [unit], arbitrate=lambda source, candidate: "uncertain"
        )

        self.assertEqual(outcome.reviews[0].reason, RETRANSLATE_UNCERTAIN)
        self.assertEqual(apply_arbitration(outcome), [unit])
        self.assertEqual(unit.status, COVERAGE_SOURCE_ONLY)

    def test_unusable_engine_keeps_the_heuristic_verdict(self) -> None:
        """本地引擎没有 chat，不能因为"没法确认"就把整份翻好的文档重翻一遍。"""
        unit = _pair(LONG_SOURCE, PARTIAL_TARGET)

        outcome = review_coverage_pairs([unit], arbitrate=None)

        self.assertEqual(outcome.model_check_count, 0)
        self.assertEqual(apply_arbitration(outcome), [])
        self.assertEqual(unit.status, COVERAGE_COVERED)

    def test_model_check_cap_keeps_the_rest_covered_and_is_reported(self) -> None:
        units = [_pair(LONG_SOURCE, PARTIAL_TARGET, index=i) for i in range(5)]

        outcome = review_coverage_pairs(
            units,
            arbitrate=lambda source, candidate: "not_equivalent",
            max_model_checks=2,
        )

        self.assertEqual(outcome.model_check_count, 2)
        self.assertEqual(outcome.skipped_over_cap, 3)
        self.assertEqual(len(apply_arbitration(outcome)), 2)

    def test_model_batch_is_announced_before_it_starts(self) -> None:
        """预处理阶段唯一的网络等待，界面上必须有交代。"""
        announced: list[int] = []

        review_coverage_pairs(
            [_pair(LONG_SOURCE, PARTIAL_TARGET, index=i) for i in range(3)]
            + [_pair("序号", "N°", index=9)],
            arbitrate=lambda source, candidate: "equivalent",
            notify_model_checks=announced.append,
        )

        self.assertEqual(announced, [3])

    def test_nothing_is_announced_when_no_pair_reaches_the_model(self) -> None:
        announced: list[int] = []

        review_coverage_pairs(
            [_pair("序号", "N°")],
            arbitrate=lambda source, candidate: "equivalent",
            notify_model_checks=announced.append,
        )

        self.assertEqual(announced, [])


class WriterIntegrationTests(unittest.TestCase):
    """改判之后，那一段必须真的被翻译并写进输出文档——否则整套复核是空转。"""

    def test_flipped_pair_is_translated_into_the_output_document(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source_path = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph(LONG_SOURCE)
            doc.add_paragraph("PV")  # 排版残留，被启发式误判成译文
            doc.save(source_path)

            plan = build_word_coverage_plan(
                source_path, target_lang="fr", source_lang="zh"
            )
            self.assertEqual(plan.source_texts, [])  # 启发式：整段都算已翻好

            outcome = review_coverage_pairs(
                plan.units, arbitrate=lambda source, candidate: "not_equivalent"
            )
            apply_arbitration(outcome)

            # 改判必须体现在写入器读的那个属性上，而不只是留在 outcome 里。
            self.assertEqual(plan.source_texts, [LONG_SOURCE])

            out_path = write_untranslated_docx(
                source_path=source_path,
                output_dir=Path(tmp) / "out",
                plan=plan,
                translations={LONG_SOURCE: FULL_TARGET},
                target_lang="fr",
            )

            texts = [p.text for p in Document(out_path).paragraphs]
            self.assertEqual(texts[:3], [LONG_SOURCE, FULL_TARGET, "PV"])


class RunnerGlueTests(unittest.TestCase):
    def _runner(self, logs: list[tuple[str, str]]):
        runner = word_task_runner.WordTaskRunner.__new__(
            word_task_runner.WordTaskRunner
        )
        runner._log = lambda level, message: logs.append((level, message))
        runner._stop_event = threading.Event()
        return runner

    def test_arbitration_failure_does_not_take_the_whole_file_down(self) -> None:
        """复核挂了只能退回原判——外层会把异常当成"文件打不开"，那样连译文都不出。"""
        logs: list[tuple[str, str]] = []
        unit = _pair(LONG_SOURCE, PARTIAL_TARGET)
        plan = SimpleNamespace(units=[unit])
        issues: list[dict] = []

        def boom(*args, **kwargs):
            raise RuntimeError("接口限流")

        with mock.patch.object(word_task_runner, "review_coverage_pairs", boom):
            self._runner(logs)._arbitrate_coverage_pairs(
                plan,
                engine=object(),
                api_scheduler=None,
                target_lang="fr",
                source_lang="zh",
                lang_pair=None,
                concurrency=4,
                file_name="方案.docx",
                file_identity="方案.docx",
                quality_issues=issues,
            )

        self.assertEqual(unit.status, COVERAGE_COVERED)
        self.assertEqual(issues, [])
        self.assertTrue(
            any(level == "WARNING" and "接口限流" in message for level, message in logs),
            logs,
        )


if __name__ == "__main__":
    unittest.main(verbosity=2)
