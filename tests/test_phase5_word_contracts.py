"""Phase 5 Word workflow contracts, all executable without local Office.

The tests deliberately use isolated DOCX fixtures plus mocked conversion and
translation services.  They make the Word safety boundaries repeatable while
real Office, real keys and macOS 12 hardware verification remain release-gate
work rather than normal development dependencies.
"""

from __future__ import annotations

from collections import deque
from contextlib import ExitStack
import hashlib
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from pydantic import ValidationError

from api.app import create_app
from api.task_manager import TaskOptions, TranslationTaskManager
from core import word_converter
from core.api_config_check import ApiConfigCheckResult
from core.language_preflight import TranslationLanguageResult
from core.model_api_identity import TaskApiContext
from core.model_throughput import EffectiveModelThroughput
from core.task_runner import DoneMsg, WordRecoveryStatusMsg
from core.word_converter import WordConversionError, convert_doc_to_docx
from core.word_document import WordFileItem, WordSegment, extract_word_segments, write_bilingual_docx
from core.word_task_runner import WordTaskRunner
from settings import AppSettings, WordBatchSettings


class _DormantRunner:
    """Captures frozen start inputs without allowing background work to begin."""

    def start(self) -> None:
        return None

    def stop(self) -> None:
        return None

    def needs_poll(self) -> bool:
        return True

    def get_message(self, timeout: float = 0.05):  # noqa: ARG002
        return None


class _WordRecoveryEventRunner:
    """A deterministic Word runner used to test the public SSE bridge."""

    def __init__(self) -> None:
        self._messages = deque(
            [
                WordRecoveryStatusMsg(
                    retry_round=2,
                    retry_total=3,
                    retry_processing_count=1,
                    retry_recovered_count=4,
                    retry_unresolved_count=1,
                    semantic_processing_count=1,
                    semantic_checked_count=5,
                    semantic_accepted_count=3,
                    semantic_uncertain_count=2,
                ),
                DoneMsg(
                    output_dir="/tmp/word-out",
                    file_results=[],
                    elapsed_sec=0.1,
                    tm_hit_count=0,
                    api_call_count=0,
                ),
            ]
        )

    def start(self) -> None:
        return None

    def stop(self) -> None:
        self._messages.clear()

    def needs_poll(self) -> bool:
        return bool(self._messages)

    def get_message(self, timeout: float = 0.05):  # noqa: ARG002
        return self._messages.popleft() if self._messages else None


class WordDocumentSafetyContractTests(unittest.TestCase):
    def test_docx_keeps_source_structure_and_never_translates_toc_or_field_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            output_dir = root / "out"
            doc = Document()
            doc.add_heading("第一章 工程概况", level=1)
            doc.add_paragraph("本工程需要翻译。")
            field_paragraph = doc.add_paragraph()
            field_paragraph.add_run("目录中的中文标题")
            field_start = OxmlElement("w:fldChar")
            field_start.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " TOC \\o \"1-3\" "
            field_end = OxmlElement("w:fldChar")
            field_end.set(qn("w:fldCharType"), "end")
            field_paragraph._p.append(field_start)
            field_paragraph._p.append(instruction)
            field_paragraph._p.append(field_end)
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "表格中的施工要求"
            doc.save(source)
            source_sha256 = hashlib.sha256(source.read_bytes()).hexdigest()

            segments = extract_word_segments(source, target_lang="en", source_lang="zh")
            extracted = {item.source for item in segments}
            self.assertIn("本工程需要翻译。", extracted)
            self.assertIn("表格中的施工要求", extracted)
            self.assertNotIn("目录中的中文标题", extracted)

            output = write_bilingual_docx(
                source_path=source,
                output_dir=output_dir,
                translations={
                    "本工程需要翻译。": "This project requires translation.",
                    "表格中的施工要求": "Construction requirements in the table",
                    "目录中的中文标题": "This must never be inserted",
                },
                target_lang="en",
                source_lang="zh",
            )

            self.assertEqual(hashlib.sha256(source.read_bytes()).hexdigest(), source_sha256)
            written = Document(output)
            paragraphs = [paragraph.text for paragraph in written.paragraphs]
            self.assertIn("本工程需要翻译。", paragraphs)
            self.assertIn("This project requires translation.", paragraphs)
            self.assertIn("目录中的中文标题", paragraphs)
            self.assertNotIn("This must never be inserted", paragraphs)
            self.assertIn("表格中的施工要求", written.tables[0].cell(0, 0).text)
            self.assertIn(
                "Construction requirements in the table",
                written.tables[0].cell(0, 0).text,
            )


class WordBatchingContractTests(unittest.TestCase):
    def test_word_batch_threshold_is_automatically_raised_to_character_budget(self) -> None:
        settings = WordBatchSettings(
            max_paragraphs_per_batch=4,
            max_chars_per_batch=5000,
            split_paragraph_chars=1500,
            strict_retry_attempts=8,
        )
        self.assertEqual(settings.split_paragraph_chars, 5000)
        self.assertEqual(settings.strict_retry_attempts, 8)

    def test_word_strict_retry_accepts_only_one_to_eight_attempts(self) -> None:
        self.assertEqual(WordBatchSettings(strict_retry_attempts=1).strict_retry_attempts, 1)
        self.assertEqual(WordBatchSettings(strict_retry_attempts=8).strict_retry_attempts, 8)
        with self.assertRaises(ValidationError):
            WordBatchSettings(strict_retry_attempts=0)
        with self.assertRaises(ValidationError):
            WordBatchSettings(strict_retry_attempts=9)


class WordSettingsAndTaskSnapshotContractTests(unittest.TestCase):
    def test_word_output_settings_are_independent_and_frozen_in_the_task_snapshot(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            Document().save(source)
            settings = AppSettings()
            self.assertTrue(hasattr(settings, "word_output"))
            settings.word_output.use_custom_output_dir = True
            settings.word_output.custom_output_dir = str(root / "word-out")
            settings.excel_output.use_custom_output_dir = False
            captured: dict[str, object] = {}
            manager = TranslationTaskManager(settings_loader=lambda: settings)
            manager._scan = lambda *_args: [SimpleNamespace(path=source, format="docx")]
            manager._build_runner = lambda **kwargs: (
                captured.update(kwargs) or _DormantRunner()
            )
            context = TaskApiContext(frozenset(), {})

            with (
                patch("api.task_manager.tm_manager.init_db"),
                patch("api.task_manager.task_api_context_for_page", return_value=context),
                patch(
                    "api.task_manager.check_translation_api_config",
                    return_value=ApiConfigCheckResult(ok=True),
                ),
                patch("api.task_manager.threading.Thread") as thread_type,
            ):
                started = manager.start_task(
                    surface="word",
                    source_path=str(root),
                    options=TaskOptions(allow_doc_fallback=True),
                )
                thread_type.return_value.start.assert_called_once_with()

            frozen = captured["settings"]
            self.assertIsInstance(frozen, AppSettings)
            self.assertIsNot(frozen.word_output, frozen.excel_output)
            self.assertTrue(frozen.word_output.use_custom_output_dir)
            self.assertFalse(frozen.excel_output.use_custom_output_dir)
            snapshot = started["task_snapshot"]
            self.assertEqual(snapshot["surface"], "word")
            self.assertEqual(snapshot["selected_file_count"], 1)
            self.assertEqual(snapshot["doc_file_count"], 0)
            self.assertEqual(snapshot["doc_conversion_mode"], "compatibility")
            self.assertTrue(snapshot["word_output"]["use_custom_output_dir"])
            self.assertIn("word_batch", snapshot)
            self.assertIn("word_review", snapshot)
            self.assertIn("word_conversion", snapshot)

    def test_task_api_forwards_explicit_doc_compatibility_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "legacy.doc"
            source.write_bytes(b"legacy source")
            captured: dict[str, object] = {}
            manager = TranslationTaskManager(settings_loader=AppSettings)
            manager._scan = lambda *_args: [SimpleNamespace(path=source, format="doc")]
            manager._build_runner = lambda **kwargs: (
                captured.update(kwargs) or _WordRecoveryEventRunner()
            )
            context = TaskApiContext(frozenset(), {})
            client = TestClient(create_app(task_manager=manager))

            with (
                patch("api.task_manager.tm_manager.init_db"),
                patch("api.task_manager.task_api_context_for_page", return_value=context),
                patch(
                    "api.task_manager.check_translation_api_config",
                    return_value=ApiConfigCheckResult(ok=True),
                ),
            ):
                response = client.post(
                    "/api/tasks",
                    json={
                        "surface": "word",
                        "source_path": str(root),
                        "allow_doc_fallback": True,
                    },
                )

            self.assertEqual(response.status_code, 202, response.text)
            self.assertTrue(captured["options"].allow_doc_fallback)


class WordScanAndEventContractTests(unittest.TestCase):
    def test_scan_keeps_doc_unknown_statistics_visible_and_separate_from_docx_totals(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            known = root / "known.docx"
            legacy = root / "legacy.doc"
            document = Document()
            document.add_paragraph("可统计的正文")
            document.add_table(rows=1, cols=1).cell(0, 0).text = "可统计的表格"
            document.save(known)
            legacy.write_bytes(b"legacy binary Word")

            with TestClient(create_app()) as client:
                response = client.post(
                    "/api/sources/scan",
                    json={"surface": "word", "path": str(root)},
                )

            self.assertEqual(response.status_code, 200, response.text)
            payload = response.json()
            self.assertEqual(set(payload), {"items", "skipped", "summary", "risk", "result"})
            known_item = next(item for item in payload["items"] if item["format"] == "docx")
            legacy_item = next(item for item in payload["items"] if item["format"] == "doc")
            self.assertEqual(known_item["paragraph_count"], 1)
            self.assertEqual(known_item["table_count"], 1)
            self.assertTrue(legacy_item["needs_conversion"])
            self.assertEqual(
                legacy_item["statistics_status"],
                "conversion_required",
            )
            self.assertIsNone(legacy_item["paragraph_count"])
            self.assertIsNone(legacy_item["table_count"])
            self.assertEqual(payload["summary"]["scanned_count"], 2)
            self.assertEqual(payload["summary"]["doc_count"], 1)
            self.assertEqual(payload["summary"]["doc_unknown_count"], 1)
            self.assertEqual(payload["summary"]["paragraph_count"], 1)
            self.assertEqual(payload["summary"]["table_count"], 1)
            self.assertTrue(payload["risk"]["has_doc"])

    def test_word_recovery_status_is_replayed_as_a_dedicated_sse_event(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            Document().save(source)
            manager = TranslationTaskManager(settings_loader=AppSettings)
            manager._scan = lambda *_args: [SimpleNamespace(path=source, format="docx")]
            manager._build_runner = lambda **_kwargs: _WordRecoveryEventRunner()
            context = TaskApiContext(frozenset(), {})

            with (
                patch("api.task_manager.tm_manager.init_db"),
                patch("api.task_manager.task_api_context_for_page", return_value=context),
                patch(
                    "api.task_manager.check_translation_api_config",
                    return_value=ApiConfigCheckResult(ok=True),
                ),
            ):
                started = manager.start_task(surface="word", source_path=str(root))
                stream = "".join(manager.iter_sse(started["task_id"]))

            self.assertIn("event: word_recovery", stream)
            self.assertIn('"retry_round":2', stream)
            self.assertIn('"semantic_accepted_count":3', stream)
            self.assertIn("event: done", stream)


class WordTaskResultContractTests(unittest.TestCase):
    @staticmethod
    def _settings() -> AppSettings:
        return AppSettings(source_lang="zh", target_lang="en")

    @staticmethod
    def _terminal_message(runner: WordTaskRunner, message_type):
        matches = [
            message
            for message in list(runner._queue.queue)
            if isinstance(message, message_type)
        ]
        if len(matches) != 1:
            raise AssertionError(
                f"expected exactly one {message_type.__name__}, got {len(matches)}"
            )
        return matches[0]

    def _runner_patches(
        self,
        stack: ExitStack,
        *,
        root: Path,
        prepared_by_path: dict[Path, object] | None = None,
    ) -> MagicMock:
        writer = stack.enter_context(
            patch(
                "core.word_task_runner.write_bilingual_docx",
                side_effect=lambda **kwargs: root / "out" / kwargs["output_name"],
            )
        )
        stack.enter_context(
            patch("core.word_task_runner.TaskLogger", return_value=MagicMock(task_id="phase5-contract"))
        )
        stack.enter_context(
            patch(
                "core.word_task_runner.check_translation_api_config",
                return_value=ApiConfigCheckResult(ok=True),
            )
        )
        stack.enter_context(
            patch("core.word_task_runner.build_engine", return_value=SimpleNamespace(engine_name="phase5/mock"))
        )
        stack.enter_context(patch("core.word_task_runner.get_system_prompt", return_value="system"))
        stack.enter_context(patch("core.word_task_runner.resolve_effective_model_config", return_value=object()))
        stack.enter_context(
            patch(
                "core.word_task_runner.get_model_throughput",
                return_value=EffectiveModelThroughput(
                    profile_key="phase5", batch_size=10, concurrency=1
                ),
            )
        )
        stack.enter_context(
            patch("core.word_task_runner.build_word_output_dir", return_value=root / "out")
        )
        stack.enter_context(
            patch("core.word_task_runner._append_post_write_coverage_issues", return_value=0)
        )
        stack.enter_context(
            patch("core.word_task_runner._write_word_quality_report", return_value=root / "out" / "word_translation_report.md")
        )
        stack.enter_context(
            patch(
                "core.word_task_runner.tm_manager.lookup_batch",
                side_effect=lambda texts, _pair: {text: "TM translation" for text in texts},
            )
        )
        stack.enter_context(patch("core.word_task_runner.tm_manager.insert_batch", return_value=0))
        if prepared_by_path is not None:
            def prepare(path: Path, **_kwargs):
                prepared = prepared_by_path[path]
                if isinstance(prepared, BaseException):
                    raise prepared
                return prepared

            stack.enter_context(
                patch(
                    "core.word_task_runner._prepare_word_source_for_translation",
                    side_effect=prepare,
                )
            )
        return writer

    def test_terminal_result_keeps_success_and_failure_traceable_without_modifying_sources(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            legacy = root / "legacy.doc"
            healthy = root / "healthy.docx"
            legacy.write_bytes(b"legacy source bytes")
            document = Document()
            document.add_paragraph("施工范围")
            document.save(healthy)
            source_hashes = {
                path: hashlib.sha256(path.read_bytes()).hexdigest()
                for path in (legacy, healthy)
            }
            prepared_healthy = SimpleNamespace(
                path=healthy,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [
                    WordFileItem(
                        path=legacy,
                        name=legacy.name,
                        size_kb=1.0,
                        format="doc",
                        needs_conversion=True,
                        statistics_status="conversion_required",
                    ),
                    WordFileItem(path=healthy, name=healthy.name, size_kb=1.0),
                ],
                self._settings(),
                source_root=root,
                allow_doc_fallback=False,
            )
            with ExitStack() as stack:
                writer = self._runner_patches(
                    stack,
                    root=root,
                    prepared_by_path={
                        legacy: WordConversionError("Word automation permission denied"),
                        healthy: prepared_healthy,
                    },
                )
                runner._run()

            writer.assert_called_once()
            done = self._terminal_message(runner, DoneMsg)
            by_name = {item["name"]: item for item in done.files}
            self.assertEqual(set(by_name), {legacy.name, healthy.name})
            self.assertEqual(by_name[legacy.name]["status"], "failed")
            self.assertFalse(by_name[legacy.name]["success"])
            self.assertEqual(by_name[legacy.name]["source_relative_path"], legacy.name)
            self.assertEqual(by_name[legacy.name]["format"], "doc")
            self.assertIn("automation", by_name[legacy.name]["error"])
            self.assertEqual(by_name[healthy.name]["status"], "succeeded")
            self.assertTrue(by_name[healthy.name]["success"])
            self.assertEqual(by_name[healthy.name]["source_relative_path"], healthy.name)
            self.assertEqual(by_name[healthy.name]["format"], "docx")
            self.assertTrue(by_name[healthy.name]["output"].endswith("healthy.docx"))
            self.assertIn("conversion", by_name[healthy.name])
            self.assertIn("numbering", by_name[healthy.name])
            self.assertEqual(done.kpi["selected_file_count"], 2)
            self.assertEqual(done.kpi["succeeded_file_count"], 1)
            self.assertEqual(done.kpi["failed_file_count"], 1)
            self.assertEqual(done.kpi["unstarted_file_count"], 0)
            self.assertEqual(done.kpi["tm_hit_count"], 1)
            self.assertEqual(done.kpi["model_translation_text_count"], 0)
            self.assertIn("items", done.review)
            self.assertIn("files", done.language)
            self.assertEqual(
                {
                    path: hashlib.sha256(path.read_bytes()).hexdigest()
                    for path in (legacy, healthy)
                },
                source_hashes,
            )

    def test_stop_before_word_scanning_records_each_file_as_unstarted(self) -> None:
        from core.task_runner import StoppedMsg

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            first = root / "first.docx"
            second = root / "second.docx"
            Document().save(first)
            Document().save(second)
            source_hashes = {
                path: hashlib.sha256(path.read_bytes()).hexdigest()
                for path in (first, second)
            }
            runner = WordTaskRunner(
                [
                    WordFileItem(path=first, name=first.name, size_kb=1.0),
                    WordFileItem(path=second, name=second.name, size_kb=1.0),
                ],
                self._settings(),
                source_root=root,
            )
            with ExitStack() as stack:
                self._runner_patches(stack, root=root)
                runner.stop()
                runner._run()

            stopped = self._terminal_message(runner, StoppedMsg)
            self.assertEqual([item["status"] for item in stopped.files], ["unstarted", "unstarted"])
            self.assertEqual(
                [item["source_relative_path"] for item in stopped.files],
                [first.name, second.name],
            )
            self.assertEqual(stopped.kpi["unstarted_file_count"], 2)
            self.assertEqual(
                {
                    path: hashlib.sha256(path.read_bytes()).hexdigest()
                    for path in (first, second)
                },
                source_hashes,
            )

    def test_auto_tm_rejects_model_source_language_outside_the_file_preflight_candidates(self) -> None:
        """A model source tag must narrow auto-TM use, never expand it.

        The file preflight here allows French only, while the translation
        protocol reports English for the individual item.  The output can be
        written, but this disagreement must neither create an en-zh TM entry
        nor be silently relabelled as fr-zh.
        """
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            Document().save(source)
            french = "Bonjour, cette exigence doit être vérifiée."
            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            settings = self._settings()
            settings.source_lang = "auto"
            settings.target_lang = "zh"
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                settings,
                source_root=root,
                source_lang="auto",
            )
            inserted: list[tuple[str, list[tuple[str, str]]]] = []

            class _PreflightEngine:
                engine_name = "phase5/mock"

                def __init__(self) -> None:
                    self.calls: list[tuple[str, str]] = []

                def chat(self, system: str, user: str) -> str:
                    self.calls.append((system, user))
                    return '{"source_langs":["fr"]}'

            engine = _PreflightEngine()

            def translate_with_reported_source(texts, _engine, target_lang, *_args, **kwargs):
                self.assertTrue(kwargs["report_source_languages"])
                callback = kwargs["source_result_callback"]
                self.assertIsNotNone(callback)
                for text in texts:
                    callback(
                        text,
                        TranslationLanguageResult(
                            source_text=text,
                            translation="应核查此要求。",
                            source_lang="en",
                            target_lang=target_lang,
                            tm_eligible=True,
                        ),
                    )
                return {text: "应核查此要求。" for text in texts}

            def capture_insert(entries, pair, *_args, **_kwargs):
                inserted.append((pair, list(entries)))
                return len(entries)

            with ExitStack() as stack:
                self._runner_patches(
                    stack,
                    root=root,
                    prepared_by_path={source: prepared},
                )
                stack.enter_context(
                    patch("core.word_task_runner.build_engine", return_value=engine)
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.extract_word_segments",
                        return_value=[
                            WordSegment(
                                source=french,
                                kind="paragraph",
                                location="正文第 1 段",
                            )
                        ],
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.lookup_batch",
                        return_value={french: None},
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.translate_word_texts",
                        side_effect=translate_with_reported_source,
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.insert_batch",
                        side_effect=capture_insert,
                    )
                )
                runner._run()

            done = self._terminal_message(runner, DoneMsg)
            self.assertEqual(len(engine.calls), 1)
            self.assertEqual(inserted, [])
            self.assertEqual(done.language["mode"], "automatic")
            self.assertEqual(
                done.language["files"][0]["preflight"]["source_langs"], ["fr"]
            )
            self.assertEqual(done.language["files"][0]["actual_source_counts"], {"en": 1})

    def test_quality_report_write_failure_only_warns_and_keeps_successful_word_result(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("施工范围")
            document.save(source)
            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                self._settings(),
                source_root=root,
            )
            with ExitStack() as stack:
                self._runner_patches(
                    stack,
                    root=root,
                    prepared_by_path={source: prepared},
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner._write_word_quality_report",
                        side_effect=OSError("read-only output directory"),
                    )
                )
                runner._run()

            done = self._terminal_message(runner, DoneMsg)
            self.assertEqual(done.files[0]["status"], "succeeded")
            self.assertEqual(done.report_path, "")
            self.assertIn("read-only output directory", done.report_warning)
            warnings = [
                message.message
                for message in list(runner._queue.queue)
                if getattr(message, "level", "") == "WARN"
            ]
            self.assertTrue(any("报告" in message for message in warnings))

    def test_only_strict_retry_recoveries_are_eligible_for_word_tm_writes(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "source.docx"
            Document().save(source)
            strict = "严格重试恢复的段落必须允许写入翻译记忆库。"
            semantic = "语义仲裁接受的边界段落不得自动写入记忆库。"
            unresolved = "未恢复的段落必须保留原文并且不得写入记忆库。"
            translations = {
                strict: "A strict-retry recovery may enter translation memory.",
                semantic: "A semantically accepted borderline paragraph must not enter memory.",
                unresolved: "An unresolved paragraph must remain source text.",
            }
            prepared = SimpleNamespace(
                path=source,
                method="编号预处理：Python 兜底",
                temp_paths=(),
                fallback_messages=(),
                labels_seen=0,
                labels_prepended=0,
                conversion_method="not_required",
                conversion_fidelity="not_required",
                numbering_method="python_conservative",
                numbering_fallback_messages=(),
            )
            runner = WordTaskRunner(
                [WordFileItem(path=source, name=source.name, size_kb=1.0)],
                self._settings(),
                source_root=root,
            )
            inserted: list[tuple[str, list[tuple[str, str]]]] = []

            class _RecoveryPool:
                def add_candidate(self, *_args, **_kwargs) -> None:
                    return None

                def start(self) -> None:
                    return None

                def wait_for_completion(self):
                    return SimpleNamespace(
                        fixed_sources=[strict],
                        unresolved_sources=[unresolved],
                        accepted_translations={strict: translations[strict]},
                        recovery_review_results={},
                        semantic_review_results={
                            semantic: SimpleNamespace(review_fragments=())
                        },
                        unresolved_validation_results={
                            unresolved: SimpleNamespace(review_fragments=())
                        },
                        semantic_check_count=1,
                    )

            recovery_pool = _RecoveryPool()

            def translate_with_all_results(texts, *_args, **kwargs):
                kwargs["drained_callback"]()
                return {text: translations[text] for text in texts}

            def capture_insert(entries, pair, *_args, **_kwargs):
                inserted.append((pair, list(entries)))
                return len(entries)

            with ExitStack() as stack:
                self._runner_patches(
                    stack,
                    root=root,
                    prepared_by_path={source: prepared},
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.extract_word_segments",
                        return_value=[
                            WordSegment(strict, "paragraph", "正文第 1 段"),
                            WordSegment(semantic, "paragraph", "正文第 2 段"),
                            WordSegment(unresolved, "paragraph", "正文第 3 段"),
                        ],
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.lookup_batch",
                        return_value={text: None for text in translations},
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.translate_word_texts",
                        side_effect=translate_with_all_results,
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner._WordRecoveryPool",
                        return_value=recovery_pool,
                    )
                )
                stack.enter_context(
                    patch(
                        "core.word_task_runner.tm_manager.insert_batch",
                        side_effect=capture_insert,
                    )
                )
                runner._run()

            done = self._terminal_message(runner, DoneMsg)
            self.assertEqual(inserted, [("zh-en", [(strict, translations[strict])])])
            self.assertEqual(done.recovery["recovered_count"], 1)
            self.assertEqual(done.recovery["semantic_accepted_count"], 1)
            self.assertEqual(done.recovery["unresolved_count"], 1)


class WordLegacyConversionContractTests(unittest.TestCase):
    def test_macos_word_automation_help_uses_the_real_monterey_settings_path(self) -> None:
        with (
            patch.object(word_converter.platform, "system", return_value="Darwin"),
            patch.object(word_converter.platform, "mac_ver", return_value=("12.7.6", (), "")),
        ):
            self.assertEqual(
                word_converter.macos_word_automation_privacy_path(),
                "系统偏好设置 > 安全性与隐私 > 隐私 > 自动化",
            )
        with (
            patch.object(word_converter.platform, "system", return_value="Darwin"),
            patch.object(word_converter.platform, "mac_ver", return_value=("13.0", (), "")),
        ):
            self.assertEqual(
                word_converter.macos_word_automation_privacy_path(),
                "系统设置 > 隐私与安全性 > 自动化",
            )

    def test_doc_compatibility_conversion_requires_explicit_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "legacy.doc"
            converted = root / "converted.docx"
            source.write_bytes(b"legacy binary Word")
            Document().save(converted)
            compatibility_calls: list[Path] = []

            def unavailable_native(_path: Path) -> Path:
                raise WordConversionError("macOS automation permission denied")

            def compatibility_converter(path: Path) -> Path:
                compatibility_calls.append(path)
                return converted

            with patch.dict(
                convert_doc_to_docx.__globals__,
                {
                    "convert_with_native_word": unavailable_native,
                    "convert_with_libreoffice": compatibility_converter,
                    "convert_with_textutil": compatibility_converter,
                },
            ):
                with self.assertRaises(WordConversionError):
                    convert_doc_to_docx(
                        source,
                        prefer_native_word=True,
                        allow_compatibility_fallback=False,
                    )
                self.assertEqual(compatibility_calls, [])

                result = convert_doc_to_docx(
                    source,
                    prefer_native_word=True,
                    allow_compatibility_fallback=True,
                )

            self.assertEqual(result.path, converted)
            self.assertEqual(result.method, "LibreOffice")
            self.assertEqual(compatibility_calls, [source])


if __name__ == "__main__":
    unittest.main(verbosity=2)
