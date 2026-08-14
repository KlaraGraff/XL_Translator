from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.translation_coverage import (
    COVERAGE_COVERED,
    COVERAGE_IGNORED,
    COVERAGE_SOURCE_ONLY,
)
from core.mixed_language import MIXED_MARK_UNRESOLVED
from core.word_coverage import (
    apply_coverage_review_marks,
    build_word_coverage_plan,
    write_untranslated_docx,
)


def _append_auto_toc_field(paragraph) -> None:
    """把一个空跑（run）标记为 Word 自动生成的 TOC 域，模拟 Insert TOC 之后的段落。"""
    paragraph.add_run()
    field_start = OxmlElement("w:fldChar")
    field_start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" '
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(field_start)
    paragraph._p.append(instruction)
    paragraph._p.append(field_end)


class WordCoverageTests(unittest.TestCase):
    def test_adjacent_paragraph_covered_and_source_only_written(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph("项目名称")
            doc.add_paragraph("Project name")
            doc.add_paragraph("施工内容")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "设备安装"
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="en", source_lang="zh")
            by_location = {unit.location: unit for unit in plan.units}

            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_COVERED)
            self.assertEqual(by_location["body.paragraph[2]"].status, COVERAGE_SOURCE_ONLY)
            self.assertEqual(by_location["table[0].cell[0]"].status, COVERAGE_SOURCE_ONLY)
            self.assertEqual(plan.source_texts, ["施工内容", "设备安装"])

            out_path = write_untranslated_docx(
                source_path=source,
                output_dir=Path(tmp) / "out",
                plan=plan,
                translations={
                    "施工内容": "Construction scope",
                    "设备安装": "Equipment installation",
                },
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            self.assertEqual(
                [paragraph.text for paragraph in out_doc.paragraphs],
                ["项目名称", "Project name", "施工内容", "Construction scope"],
            )
            self.assertEqual(
                out_doc.tables[0].cell(0, 0).text,
                "设备安装\nEquipment installation",
            )

            second_plan = build_word_coverage_plan(
                out_path,
                target_lang="en",
                source_lang="zh",
            )
            self.assertEqual(second_plan.source_units, [])

    def test_duplicate_source_text_only_patches_untranslated_position(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph("项目名称")
            doc.add_paragraph("Project name")
            doc.add_paragraph("项目名称")
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="en", source_lang="zh")

            self.assertEqual(len(plan.source_units), 1)
            self.assertEqual(plan.source_units[0].location, "body.paragraph[2]")

            out_path = write_untranslated_docx(
                source_path=source,
                output_dir=Path(tmp) / "out",
                plan=plan,
                translations={"项目名称": "Project name"},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            self.assertEqual(
                [paragraph.text for paragraph in out_doc.paragraphs],
                ["项目名称", "Project name", "项目名称", "Project name"],
            )

    def test_table_unit_line_is_not_misclassified_as_target_translation(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "units.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "水泥\n(kg/m³)"
            table.cell(0, 1).text = "砂子\nsable\n(kg/m³)"
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")
            by_location = {unit.location: unit for unit in plan.units}

            self.assertEqual(by_location["table[0].cell[0]"].status, COVERAGE_SOURCE_ONLY)
            self.assertEqual(by_location["table[0].cell[1]"].status, COVERAGE_COVERED)

    def test_front_matter_protects_cover_and_auto_toc_before_first_chapter(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "封面加自动目录.docx"
            doc = Document()
            doc.add_paragraph("某某工程施工方案")
            doc.add_paragraph("编制单位：某某公司")
            _append_auto_toc_field(doc.add_paragraph())
            doc.add_paragraph("第一章 工程概况")
            doc.add_paragraph("施工内容")
            doc.save(source)

            plan = build_word_coverage_plan(
                source,
                target_lang="fr",
                source_lang="zh",
                protect_front_matter=True,
            )

            self.assertTrue(plan.front_matter.found)
            self.assertEqual(plan.front_matter.heading_text, "第一章 工程概况")
            self.assertEqual(plan.front_matter.protected_paragraph_count, 2)
            by_location = {unit.location: unit for unit in plan.units}
            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_IGNORED)
            self.assertEqual(by_location["body.paragraph[1]"].status, COVERAGE_IGNORED)
            # 域段落本身没有可见文字，_classify_body_paragraphs 从不为空段落生成 unit，
            # 这里不用断言 paragraph[2]；重点是它前后的边界判断没有被打乱。
            self.assertEqual(by_location["body.paragraph[3]"].status, COVERAGE_SOURCE_ONLY)
            self.assertEqual(by_location["body.paragraph[4]"].status, COVERAGE_SOURCE_ONLY)

    def test_front_matter_hand_typed_toc_leader_is_not_mistaken_for_body_heading(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "封面加手打目录.docx"
            doc = Document()
            doc.add_paragraph("某某工程施工方案")
            # 手打目录条目本身以"第一章"开头——这正是 guardrail A 要防的假阳性：
            # 不带这条正则的话，扫描会在这里就误判成正文标题，导致目录本身反而被当正文翻译。
            doc.add_paragraph("第一章 工程概况 ...... 1")
            doc.add_paragraph("第二章 施工部署 ...... 5")
            doc.add_paragraph("第一章 工程概况")
            doc.add_paragraph("施工内容")
            doc.save(source)

            plan = build_word_coverage_plan(
                source,
                target_lang="fr",
                source_lang="zh",
                protect_front_matter=True,
            )

            self.assertTrue(plan.front_matter.found)
            self.assertEqual(plan.front_matter.heading_text, "第一章 工程概况")
            self.assertEqual(plan.front_matter.protected_paragraph_count, 3)
            by_location = {unit.location: unit for unit in plan.units}
            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_IGNORED)
            self.assertEqual(by_location["body.paragraph[1]"].status, COVERAGE_IGNORED)
            self.assertEqual(by_location["body.paragraph[2]"].status, COVERAGE_IGNORED)
            self.assertEqual(by_location["body.paragraph[3]"].status, COVERAGE_SOURCE_ONLY)
            self.assertEqual(by_location["body.paragraph[4]"].status, COVERAGE_SOURCE_ONLY)

    def test_front_matter_bare_numeric_heading_without_decimal_point(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "无点号标题.docx"
            doc = Document()
            doc.add_paragraph("某某工程施工方案")
            doc.add_paragraph("1 概述")
            doc.add_paragraph("施工内容")
            doc.save(source)

            plan = build_word_coverage_plan(
                source,
                target_lang="fr",
                source_lang="zh",
                protect_front_matter=True,
            )

            self.assertTrue(plan.front_matter.found)
            self.assertEqual(plan.front_matter.heading_text, "1 概述")
            by_location = {unit.location: unit for unit in plan.units}
            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_IGNORED)
            self.assertEqual(by_location["body.paragraph[1]"].status, COVERAGE_SOURCE_ONLY)
            self.assertEqual(by_location["body.paragraph[2]"].status, COVERAGE_SOURCE_ONLY)

    def test_front_matter_word_heading_style_without_recognizable_text_pattern(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "样式标题.docx"
            doc = Document()
            doc.add_paragraph("某某工程施工方案")
            # Word 内置"标题 1"样式，但文字本身既不带章节前缀也不带数字——只能靠样式识别。
            doc.add_heading("工程概况说明", level=1)
            doc.add_paragraph("施工内容")
            doc.save(source)

            plan = build_word_coverage_plan(
                source,
                target_lang="fr",
                source_lang="zh",
                protect_front_matter=True,
            )

            self.assertTrue(plan.front_matter.found)
            self.assertEqual(plan.front_matter.heading_text, "工程概况说明")
            by_location = {unit.location: unit for unit in plan.units}
            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_IGNORED)
            self.assertEqual(by_location["body.paragraph[2]"].status, COVERAGE_SOURCE_ONLY)

    def test_front_matter_protects_nothing_when_no_heading_is_found(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "无标题文档.docx"
            doc = Document()
            doc.add_paragraph("项目背景")
            doc.add_paragraph("施工内容")
            doc.add_paragraph("验收标准")
            doc.save(source)

            plan = build_word_coverage_plan(
                source,
                target_lang="fr",
                source_lang="zh",
                protect_front_matter=True,
            )

            # 找不到正文标题时绝不能把整份文档当封面吞掉：断言"没保护任何内容"。
            self.assertFalse(plan.front_matter.found)
            self.assertEqual(plan.front_matter.protected_paragraph_count, 0)
            self.assertEqual(
                [unit.source_text for unit in plan.source_units],
                ["项目背景", "施工内容", "验收标准"],
            )

    def test_front_matter_protects_cover_table_before_first_chapter(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "表格封面.docx"
            doc = Document()
            cover_table = doc.add_table(rows=2, cols=2)
            cover_table.cell(0, 0).text = "文件编号"
            cover_table.cell(0, 1).text = "ONEBTR-MS-035"
            cover_table.cell(1, 0).text = "编制人"
            cover_table.cell(1, 1).text = "张三"
            doc.add_paragraph("第一章 工程概况")
            body_table = doc.add_table(rows=1, cols=1)
            body_table.cell(0, 0).text = "工程量清单"
            doc.save(source)

            plan = build_word_coverage_plan(
                source,
                target_lang="fr",
                source_lang="zh",
                protect_front_matter=True,
            )

            self.assertTrue(plan.front_matter.found)
            cover_units = [unit for unit in plan.units if unit.location.startswith("table[0].")]
            body_table_units = [unit for unit in plan.units if unit.location.startswith("table[1].")]
            self.assertTrue(cover_units)
            self.assertTrue(all(unit.status == COVERAGE_IGNORED for unit in cover_units))
            self.assertTrue(body_table_units)
            self.assertTrue(all(unit.status == COVERAGE_SOURCE_ONLY for unit in body_table_units))

    def test_translation_carrying_a_numbering_prefix_still_counts_as_covered(self) -> None:
        """译文开头留着中文编号（一.1）不能让整对段落被判成未译源文。"""
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "编号前缀.docx"
            doc = Document()
            doc.add_paragraph("一.1 抢工背景分析")
            doc.add_paragraph(
                "一.1 Analyse du contexte de l’accélération des travaux"
            )
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")
            by_location = {unit.location: unit for unit in plan.units}

            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_COVERED)
            self.assertEqual(plan.source_units, [])
            self.assertEqual(len(plan.residual_units), 1)
            self.assertEqual(plan.residual_units[0].data["residual_cjk"], ["一"])
            self.assertEqual(
                plan.residual_units[0].data["residual_location"], "body.paragraph[1]"
            )

    def test_untranslated_chinese_paragraph_is_still_reported(self) -> None:
        """放宽判定不能把真正没翻的中文段落也放过去。"""
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "漏译.docx"
            doc = Document()
            doc.add_paragraph("Analyse du contexte de l’accélération des travaux")
            doc.add_paragraph("夜间施工必须配备足够的照明设备，并落实专人值守。")
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")
            by_location = {unit.location: unit for unit in plan.units}

            self.assertEqual(
                by_location["body.paragraph[1]"].status, COVERAGE_SOURCE_ONLY
            )
            self.assertEqual(plan.residual_units, [])

    def test_front_matter_disabled_by_default_leaves_cover_translatable(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "默认不保护.docx"
            doc = Document()
            doc.add_paragraph("某某工程施工方案")
            doc.add_paragraph("第一章 工程概况")
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")

            self.assertFalse(plan.front_matter.found)
            by_location = {unit.location: unit for unit in plan.units}
            self.assertEqual(by_location["body.paragraph[0]"].status, COVERAGE_SOURCE_ONLY)


class ReviewMarkTests(unittest.TestCase):
    """报告里写了几条「需人工复核」，文件里就得看得见几处标记。

    真实缺陷：用户勾了「标记需复核内容」，报告报了 38 条，打开文档一个标记都没有。
    两条原因——补译模式的写入器压根没有标记参数；而且所有需复核项都是写完文件之后
    才体检出来的，写的时候还不存在。
    """

    @staticmethod
    def _highlighted_paragraph_texts(path: Path) -> list[str]:
        doc = Document(str(path))
        marked = [
            paragraph.text
            for paragraph in doc.paragraphs
            if any(run.font.highlight_color for run in paragraph.runs)
        ]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if any(run.font.highlight_color for run in paragraph.runs):
                            marked.append(paragraph.text)
        return marked

    def _build_sample(self, tmp: Path) -> Path:
        source = tmp / "抢工方案.docx"
        doc = Document()
        doc.add_paragraph("一、工程概况")
        doc.add_paragraph("二、施工部署")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "砂浆名称"
        doc.save(source)
        return source

    def test_untranslated_only_writer_paints_review_marks(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = self._build_sample(tmp_path)
            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")

            out_path = write_untranslated_docx(
                source_path=source,
                output_dir=tmp_path / "out",
                plan=plan,
                translations={
                    "一、工程概况": "1. Présentation du projet",
                    "二、施工部署": "2. Organisation des travaux",
                    "砂浆名称": "Nom du mortier",
                },
                target_lang="fr",
                source_lang="zh",
                review_marks={
                    "一、工程概况": MIXED_MARK_UNRESOLVED,
                    "砂浆名称": MIXED_MARK_UNRESOLVED,
                },
            )

            marked = self._highlighted_paragraph_texts(out_path)
            # 原文段和补进去的译文段都要涂，用户翻到哪一边都看得见。
            self.assertIn("一、工程概况", marked)
            self.assertIn("1. Présentation du projet", marked)
            self.assertIn("砂浆名称", marked)
            # 没上标记的那一段不能被牵连。
            self.assertNotIn("二、施工部署", marked)

    def test_untranslated_only_writer_without_marks_paints_nothing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = self._build_sample(tmp_path)
            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")

            out_path = write_untranslated_docx(
                source_path=source,
                output_dir=tmp_path / "out",
                plan=plan,
                translations={"一、工程概况": "1. Présentation du projet"},
                target_lang="fr",
                source_lang="zh",
            )

            self.assertEqual(self._highlighted_paragraph_texts(out_path), [])

    def test_post_write_audit_marks_untranslated_and_residual_positions(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "双语(法文)_成品.docx"
            doc = Document()
            doc.add_paragraph("三、质量目标")  # 未译源文
            doc.add_paragraph("本工程竣工日期为 2026 年 8 月 9 日。")
            doc.add_paragraph(
                "La date d’achèvement du présent projet est fixée au 2026年8月9日."
            )  # 已翻好，只残留中文日期
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "专用砌筑砂浆"  # 未译单元格
            doc.save(output_path)

            plan = build_word_coverage_plan(
                output_path, target_lang="fr", source_lang="zh"
            )
            marked_count = apply_coverage_review_marks(output_path, plan=plan)

            self.assertEqual(marked_count, 3)
            marked = self._highlighted_paragraph_texts(output_path)
            self.assertIn("三、质量目标", marked)
            self.assertIn("专用砌筑砂浆", marked)
            self.assertIn(
                "La date d’achèvement du présent projet est fixée au 2026年8月9日.",
                marked,
            )
            # 中文原文段本身没问题，不涂。
            self.assertNotIn("本工程竣工日期为 2026 年 8 月 9 日。", marked)

    def test_post_write_audit_does_not_double_mark_what_the_writer_marked(self) -> None:
        """写入时已经涂过的位置，体检那一趟不能按「已有底色」再叠一层红色下划线。"""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "重复标记.docx"
            doc = Document()
            doc.add_paragraph("四、应急预案")
            doc.save(source)

            plan = build_word_coverage_plan(source, target_lang="fr", source_lang="zh")
            out_path = write_untranslated_docx(
                source_path=source,
                output_dir=tmp_path / "out",
                plan=plan,
                translations={"四、应急预案": "4. 应急预案 2026年"},
                target_lang="fr",
                source_lang="zh",
                review_marks={"四、应急预案": MIXED_MARK_UNRESOLVED},
                existing_highlight_policy="red_underline",
            )

            post_plan = build_word_coverage_plan(
                out_path, target_lang="fr", source_lang="zh"
            )
            apply_coverage_review_marks(
                out_path, plan=post_plan, existing_highlight_policy="red_underline"
            )

            written = Document(str(out_path))
            for paragraph in written.paragraphs:
                for run in paragraph.runs:
                    self.assertFalse(
                        run.underline,
                        f"重复涂了红色下划线：{paragraph.text}",
                    )

    def test_post_write_audit_on_clean_document_touches_nothing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "干净.docx"
            doc = Document()
            doc.add_paragraph("五、竣工验收")
            doc.add_paragraph("5. Réception des travaux")
            doc.save(output_path)
            before = output_path.stat().st_mtime_ns

            plan = build_word_coverage_plan(
                output_path, target_lang="fr", source_lang="zh"
            )

            self.assertEqual(apply_coverage_review_marks(output_path, plan=plan), 0)
            self.assertEqual(output_path.stat().st_mtime_ns, before)
            self.assertEqual(self._highlighted_paragraph_texts(output_path), [])


if __name__ == "__main__":
    unittest.main(verbosity=2)
