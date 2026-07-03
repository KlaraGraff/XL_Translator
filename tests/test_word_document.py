from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from core.mixed_language import (
    MIXED_MARK_FOREIGN_NOISE,
    MIXED_MARK_SEMANTIC,
    MIXED_MARK_UNRESOLVED,
)
from core.translation_filter import (
    VALIDATION_PROFILE_STRICT,
    VALIDATION_PROFILE_WORD_RECOVERY,
    validate_translation,
    is_translation_redundant,
)
from core.word_document import (
    extract_word_segments,
    scan_word_path,
    _has_visible_numbering_prefix,
    normalize_docx_automatic_numbering,
    write_bilingual_docx,
)
from core.word_task_runner import _needs_word_translation_retry
from core.word_task_runner import (
    _append_post_write_coverage_issues,
    _WordRecoveryPool,
    _build_source_excerpt,
    _run_word_strict_retries,
    _write_word_quality_report,
)
from engines.base_engine import TranslationEngine
from settings import WordBatchSettings


class RetryWordEngine(TranslationEngine):
    def __init__(self, *, success_on_call: int) -> None:
        self.success_on_call = success_on_call
        self.calls: list[list[str]] = []

    @property
    def engine_name(self) -> str:
        return "fake/retry-word"

    def translate_batch(
        self,
        texts: list[str],
        target_lang: str,
        system_prompt: str,
        source_lang: str = "zh",
    ) -> dict[str, str]:
        self.calls.append(list(texts))
        if len(self.calls) < self.success_on_call:
            return {text: text for text in texts}
        return {text: f"Traduction valide {len(text)}" for text in texts}


class MappingWordEngine(TranslationEngine):
    def __init__(self, translations: dict[str, str]) -> None:
        self.translations = translations
        self.calls: list[list[str]] = []

    @property
    def engine_name(self) -> str:
        return "fake/mapping-word"

    def translate_batch(
        self,
        texts: list[str],
        target_lang: str,
        system_prompt: str,
        source_lang: str = "zh",
    ) -> dict[str, str]:
        self.calls.append(list(texts))
        return {text: self.translations.get(text, text) for text in texts}


class SemanticWordEngine(MappingWordEngine):
    def __init__(self, translations: dict[str, str], verdict: str = "equivalent") -> None:
        super().__init__(translations)
        self.verdict = verdict
        self.chat_calls: list[tuple[str, str]] = []

    def chat(self, system: str, user: str) -> str:
        self.chat_calls.append((system, user))
        return f'{{"verdict":"{self.verdict}","reason":"same meaning"}}'


class WordDocumentTests(unittest.TestCase):
    def test_extract_and_write_bilingual_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "sample.docx"
            output_dir = temp_path / "out"
            self._build_sample_docx(source_path)

            segments = extract_word_segments(
                source_path,
                target_lang="en",
                source_lang="zh",
            )
            sources = {segment.source for segment in segments}

            self.assertIn("项目名称：测试工程", sources)
            self.assertIn("设备\n安装", sources)
            self.assertNotIn("12345", sources)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={
                    "项目名称：测试工程": "Project name: Test Project",
                    "设备\n安装": "Equipment installation",
                },
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertIn("项目名称：测试工程", paragraph_texts)
            self.assertIn("Project name: Test Project", paragraph_texts)

            cell_text = out_doc.tables[0].cell(0, 0).text
            self.assertEqual(cell_text, "设备\n安装\nEquipment installation")

    def test_word_review_highlight_marks_unresolved_paragraph_and_cell(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "sample.docx"
            output_dir = temp_path / "out"
            self._build_sample_docx(source_path)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={},
                target_lang="en",
                source_lang="zh",
                review_highlight_sources={"项目名称：测试工程", "设备\n安装"},
                review_highlight_color="DDEBFF",
            )

            out_doc = Document(str(out_path))
            self.assertEqual(self._paragraph_first_run_highlight(out_doc.paragraphs[0]), "cyan")
            self.assertEqual(self._paragraph_first_run_shading_fill(out_doc.paragraphs[0]), "")
            self.assertTrue(
                all(
                    self._paragraph_first_run_highlight(paragraph) == "cyan"
                    for paragraph in out_doc.tables[0].cell(0, 0).paragraphs
                    if paragraph.text.strip()
                )
            )

    def test_word_table_iteration_keeps_all_unmerged_cells(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "table.docx"
            output_dir = temp_path / "out"

            doc = Document()
            table = doc.add_table(rows=7, cols=3)
            expected_sources = []
            translations = {}
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    source = f"表格项目{row_index}列{col_index}"
                    expected_sources.append(source)
                    translations[source] = f"Table item {row_index}-{col_index}"
                    cell.text = source
            doc.save(str(source_path))

            segments = extract_word_segments(
                source_path,
                target_lang="en",
                source_lang="zh",
            )
            table_sources = [segment.source for segment in segments if segment.kind == "table_cell"]
            self.assertEqual(table_sources, expected_sources)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations=translations,
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            for row_index, row in enumerate(out_doc.tables[0].rows):
                for col_index, cell in enumerate(row.cells):
                    self.assertIn(
                        f"Table item {row_index}-{col_index}",
                        cell.text,
                    )

    def test_nested_table_cells_are_extracted_and_written_independently(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "nested-table.docx"
            output_dir = temp_path / "out"

            doc = Document()
            outer_table = doc.add_table(rows=1, cols=1)
            outer_cell = outer_table.cell(0, 0)
            outer_cell.text = "外层说明"
            nested_table = outer_cell.add_table(rows=1, cols=1)
            nested_table.cell(0, 0).text = "内部表格内容"
            doc.save(str(source_path))

            segments = extract_word_segments(
                source_path,
                target_lang="en",
                source_lang="zh",
            )
            table_sources = [segment.source for segment in segments if segment.kind == "table_cell"]

            self.assertIn("外层说明", table_sources)
            self.assertIn("内部表格内容", table_sources)
            self.assertNotIn("外层说明\n内部表格内容", table_sources)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={
                    "外层说明": "Outer note",
                    "内部表格内容": "Nested table content",
                },
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            out_outer_cell = out_doc.tables[0].cell(0, 0)
            self.assertIn("外层说明", out_outer_cell.text)
            self.assertIn("Outer note", out_outer_cell.text)
            self.assertEqual(
                out_outer_cell.tables[0].cell(0, 0).text,
                "内部表格内容\nNested table content",
            )

    def test_word_review_marks_use_configured_color_map(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "sample.docx"
            output_dir = temp_path / "out"

            doc = Document()
            doc.add_paragraph("语义接受")
            doc.add_paragraph("保留原文")
            doc.add_paragraph("疑似异常")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={},
                target_lang="en",
                source_lang="zh",
                review_marks={
                    "语义接受": MIXED_MARK_SEMANTIC,
                    "保留原文": MIXED_MARK_UNRESOLVED,
                    "疑似异常": MIXED_MARK_FOREIGN_NOISE,
                },
                review_mark_colors={
                    MIXED_MARK_SEMANTIC: "DDEBFF",
                    MIXED_MARK_UNRESOLVED: "D9EAD3",
                    MIXED_MARK_FOREIGN_NOISE: "FCE4D6",
                },
            )

            out_doc = Document(str(out_path))
            self.assertEqual(
                self._paragraph_first_run_highlight(out_doc.paragraphs[0]),
                "cyan",
            )
            self.assertEqual(
                self._paragraph_first_run_highlight(out_doc.paragraphs[1]),
                "green",
            )
            self.assertEqual(
                self._paragraph_first_run_highlight(out_doc.paragraphs[2]),
                "yellow",
            )

    def test_word_review_highlight_preserves_existing_shading(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "shaded.docx"
            output_dir = temp_path / "out"

            doc = Document()
            paragraph = doc.add_paragraph("项目名称：测试工程")
            self._set_shading_fill(paragraph._p.get_or_add_pPr(), "AABBCC")
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "设备安装"
            self._set_shading_fill(cell._tc.get_or_add_tcPr(), "CCDDEE")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={},
                target_lang="en",
                source_lang="zh",
                review_highlight_sources={"项目名称：测试工程", "设备安装"},
                review_highlight_color="FFF2CC",
            )

            out_doc = Document(str(out_path))
            self.assertEqual(self._paragraph_shading_fill(out_doc.paragraphs[0]), "AABBCC")
            self.assertEqual(self._cell_shading_fill(out_doc.tables[0].cell(0, 0)), "CCDDEE")
            self.assertEqual(self._paragraph_first_run_shading_fill(out_doc.paragraphs[0]), "")
            self.assertEqual(self._paragraph_first_run_highlight(out_doc.paragraphs[0]), "")
            self.assertEqual(
                self._paragraph_first_run_shading_fill(out_doc.tables[0].cell(0, 0).paragraphs[0]),
                "",
            )

    def test_word_review_mark_can_use_red_underline_when_existing_shading_is_preserved(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "shaded.docx"
            output_dir = temp_path / "out"

            doc = Document()
            paragraph = doc.add_paragraph("项目名称：测试工程")
            self._set_shading_fill(paragraph._p.get_or_add_pPr(), "AABBCC")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={},
                target_lang="en",
                source_lang="zh",
                review_highlight_sources={"项目名称：测试工程"},
                existing_highlight_policy="red_underline",
            )

            out_doc = Document(str(out_path))
            run = out_doc.paragraphs[0].runs[0]
            self.assertEqual(self._paragraph_shading_fill(out_doc.paragraphs[0]), "AABBCC")
            self.assertTrue(run.underline)
            self.assertEqual(str(run.font.color.rgb), "C00000")

    def test_word_automatic_numbering_is_flattened_in_bilingual_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "numbered.docx"
            output_dir = temp_path / "out"
            self._build_lower_letter_numbered_docx(source_path)

            segments = extract_word_segments(
                source_path,
                target_lang="fr",
                source_lang="zh",
            )
            sources = {segment.source for segment in segments}
            self.assertIn("包括但不限于根据合同约定或确定的款项；", sources)
            self.assertNotIn("a. 包括但不限于根据合同约定或确定的款项；", sources)

            normalized = normalize_docx_automatic_numbering(source_path)
            normalized_segments = extract_word_segments(
                normalized.path,
                target_lang="fr",
                source_lang="zh",
            )
            normalized_sources = {segment.source for segment in normalized_segments}
            self.assertIn("a. 包括但不限于根据合同约定或确定的款项；", normalized_sources)

            out_path = write_bilingual_docx(
                source_path=normalized.path,
                output_dir=output_dir,
                translations={
                    "a. 包括但不限于根据合同约定或确定的款项；": (
                        "a. y compris les montants prévus au contrat ;"
                    ),
                    "b. 乙方未履行整改通知规定的违约补救义务；": (
                        "b. La partie B n'a pas rempli les obligations de remédiation ;"
                    ),
                },
                target_lang="fr",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertIn("a. 包括但不限于根据合同约定或确定的款项；", paragraph_texts)
            self.assertIn("a. y compris les montants prévus au contrat ;", paragraph_texts)
            self.assertIn("b. 乙方未履行整改通知规定的违约补救义务；", paragraph_texts)
            self.assertIn(
                "b. La partie B n'a pas rempli les obligations de remédiation ;",
                paragraph_texts,
            )

            for paragraph in out_doc.paragraphs[:4]:
                self.assertEqual(self._paragraph_num_id(paragraph), "0")

    def test_chinese_numbering_fallback_keeps_chinese_labels_and_skips_blank_list_items(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "chinese-numbered.docx"
            self._build_chinese_section_numbered_docx(source_path)

            normalized = normalize_docx_automatic_numbering(source_path)
            out_doc = Document(str(normalized.path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]

            self.assertIn("", paragraph_texts)
            self.assertIn("第一章 裂缝开裂现状及成因分析", paragraph_texts)
            self.assertIn("第二章 修复总体原则与目标", paragraph_texts)
            self.assertNotIn("第1节 裂缝开裂现状及成因分析", paragraph_texts)
            self.assertNotIn("第2节 裂缝开裂现状及成因分析", paragraph_texts)
            self.assertTrue(all(self._paragraph_num_id(paragraph) == "0" for paragraph in out_doc.paragraphs))

    def test_decimal_chinese_section_fallback_uses_chinese_numerals(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "decimal-chinese-section.docx"
            self._build_decimal_chinese_section_numbered_docx(source_path)

            normalized = normalize_docx_automatic_numbering(source_path)
            out_doc = Document(str(normalized.path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]

            self.assertIn("第一章 裂缝开裂现状及成因分析", paragraph_texts)
            self.assertIn("第二章 修复总体原则与目标", paragraph_texts)
            self.assertNotIn("第1节 裂缝开裂现状及成因分析", paragraph_texts)
            self.assertNotIn("第2节 修复总体原则与目标", paragraph_texts)

    def test_nested_chinese_section_fallback_keeps_section_unit(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "nested-chinese-section.docx"
            self._build_nested_chinese_section_numbered_docx(source_path)

            normalized = normalize_docx_automatic_numbering(source_path)
            out_doc = Document(str(normalized.path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]

            self.assertIn("第一章 工程概况", paragraph_texts)
            self.assertIn("第一节 裂缝开裂现状及成因分析", paragraph_texts)

    def test_translation_inherits_manual_leading_spaces(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "spaces.docx"
            output_dir = temp_path / "out"

            doc = Document()
            doc.add_paragraph("   人工空格缩进段落")
            doc.add_paragraph("\u3000\u3000全角空格缩进段落")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={
                    "人工空格缩进段落": "Indented translation",
                    "全角空格缩进段落": "Full-width indented translation",
                },
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertIn("   Indented translation", paragraph_texts)
            self.assertIn("\u3000\u3000Full-width indented translation", paragraph_texts)

    def test_translation_trims_trailing_empty_body_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "trailing-empty.docx"
            output_dir = temp_path / "out"

            doc = Document()
            doc.add_paragraph("正文结束段落。")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"正文结束段落。": "Final body translation."},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertEqual(paragraph_texts[-1], "Final body translation.")
            self.assertNotIn("", paragraph_texts[-2:])

    def test_translation_preserves_style_character_first_line_indent(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "style-indent.docx"
            output_dir = temp_path / "out"

            doc = Document()
            style = doc.styles.add_style("施工正文", WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(12)
            p_pr = style.element.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLineChars"), "200")
            p_pr.append(ind)
            doc.add_paragraph("这是首行缩进正文段落。", style=style)
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"这是首行缩进正文段落。": "Indented body translation."},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text == "Indented body translation."
            )
            self.assertEqual(translation.style.name, "施工正文")
            self.assertEqual(self._paragraph_first_line_chars(translation), "")
            self.assertEqual(self._paragraph_first_line_indent(translation), "480")
            self.assertEqual(self._style_first_line_chars(translation), "200")

    def test_translation_materializes_default_style_character_indent_without_source_ppr(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "normal-style-indent.docx"
            output_dir = temp_path / "out"

            doc = Document()
            doc.styles["Normal"].font.size = Pt(12)
            p_pr = doc.styles["Normal"].element.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLineChars"), "200")
            p_pr.append(ind)
            paragraph = doc.add_paragraph("这是普通样式首行缩进正文。")
            self.assertIsNone(getattr(paragraph._p, "pPr", None))
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"这是普通样式首行缩进正文。": "Indented normal-style translation."},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text == "Indented normal-style translation."
            )
            self.assertEqual(self._paragraph_first_line_chars(translation), "")
            self.assertEqual(self._paragraph_first_line_indent(translation), "480")

    def test_translation_keeps_style_character_indent_when_direct_indent_has_zero_first_line(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "direct-zero-style-indent.docx"
            output_dir = temp_path / "out"

            doc = Document()
            style = doc.styles.add_style("方案正文", WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(12)
            style_p_pr = style.element.get_or_add_pPr()
            style_ind = OxmlElement("w:ind")
            style_ind.set(qn("w:firstLineChars"), "200")
            style_p_pr.append(style_ind)
            paragraph = doc.add_paragraph("地坪裂缝风险：受开挖影响。", style=style)
            paragraph_p_pr = paragraph._p.get_or_add_pPr()
            direct_ind = OxmlElement("w:ind")
            direct_ind.set(qn("w:left"), "420")
            direct_ind.set(qn("w:firstLine"), "0")
            paragraph_p_pr.append(direct_ind)
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"地坪裂缝风险：受开挖影响。": "Risk of slab cracking: affected by excavation."},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text == "Risk of slab cracking: affected by excavation."
            )
            self.assertEqual(translation.style.name, "方案正文")
            self.assertEqual(self._paragraph_left_indent(translation), "420")
            self.assertEqual(self._paragraph_first_line_chars(translation), "")
            self.assertEqual(self._paragraph_first_line_indent(translation), "480")

    def test_translation_child_style_zero_character_indent_clears_base_first_line(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "child-style-zero-indent.docx"
            output_dir = temp_path / "out"

            doc = Document()
            body_style = doc.styles.add_style("方案正文", WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(12)
            body_p_pr = body_style.element.get_or_add_pPr()
            body_ind = OxmlElement("w:ind")
            body_ind.set(qn("w:firstLine"), "200")
            body_ind.set(qn("w:firstLineChars"), "200")
            body_p_pr.append(body_ind)

            heading_style = doc.styles.add_style("方案三级", WD_STYLE_TYPE.PARAGRAPH)
            heading_style.base_style = body_style
            heading_p_pr = heading_style.element.get_or_add_pPr()
            heading_ind = OxmlElement("w:ind")
            heading_ind.set(qn("w:firstLineChars"), "0")
            heading_p_pr.append(heading_ind)

            doc.add_paragraph("现场施工条件现状分析", style=heading_style)
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"现场施工条件现状分析": "Analysis of site construction conditions"},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text == "Analysis of site construction conditions"
            )
            self.assertEqual(translation.style.name, "方案三级")
            self.assertEqual(self._paragraph_first_line_chars(translation), "")
            self.assertEqual(self._paragraph_first_line_indent(translation), "")

    def test_translation_does_not_expand_partial_bold_prefix_to_whole_paragraph(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "partial-bold.docx"
            output_dir = temp_path / "out"

            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("地坪裂缝风险：").bold = True
            paragraph.add_run("受开挖边坡变形及施工扰动影响。")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={
                    "地坪裂缝风险：受开挖边坡变形及施工扰动影响。": (
                        "Risques de fissures dans le dallage : sous l'effet des travaux."
                    )
                },
                target_lang="fr",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text.startswith("Risques de fissures")
            )
            self.assertIsNone(translation.runs[0].bold)

    def test_translation_does_not_repeat_direct_page_or_section_breaks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "direct-page-break.docx"
            output_dir = temp_path / "out"

            doc = Document()
            paragraph = doc.add_paragraph("分页标题")
            paragraph.paragraph_format.page_break_before = True
            paragraph._p.get_or_add_pPr().append(OxmlElement("w:sectPr"))
            doc.add_paragraph("后续正文")
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"分页标题": "Paged heading"},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            source = next(paragraph for paragraph in out_doc.paragraphs if paragraph.text == "分页标题")
            translation = next(paragraph for paragraph in out_doc.paragraphs if paragraph.text == "Paged heading")

            self.assertTrue(self._paragraph_page_break_before(source))
            self.assertTrue(self._paragraph_has_section_properties(source))
            self.assertIsNone(self._paragraph_page_break_before(translation))
            self.assertFalse(self._paragraph_has_section_properties(translation))

    def test_translation_overrides_style_page_break_before(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "style-page-break.docx"
            output_dir = temp_path / "out"

            doc = Document()
            style = doc.styles.add_style("分页样式", WD_STYLE_TYPE.PARAGRAPH)
            style.paragraph_format.page_break_before = True
            doc.add_paragraph("样式分页标题", style=style)
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"样式分页标题": "Styled paged heading"},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text == "Styled paged heading"
            )
            self.assertEqual(translation.style.name, "分页样式")
            self.assertFalse(self._paragraph_page_break_before(translation))

    def test_heading_translation_keeps_visual_emphasis_without_heading_style(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "heading.docx"
            output_dir = temp_path / "out"

            doc = Document()
            heading = doc.add_heading("第一节 工程概况分析", level=1)
            heading.style.font.bold = True
            heading.style.font.size = None
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"第一节 工程概况分析": "Section 1: Project overview"},
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            translation = next(
                paragraph
                for paragraph in out_doc.paragraphs
                if paragraph.text == "Section 1: Project overview"
            )
            self.assertFalse(translation.style.name.casefold().startswith("heading"))
            self.assertTrue(translation.runs[0].bold)

    def test_word_automatic_chapter_number_is_not_duplicated_when_text_already_has_it(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "chapter.docx"
            output_dir = temp_path / "out"
            self._build_chapter_numbered_docx(source_path)

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={
                    "第一章 编制依据": "Chapter 1 Compilation Basis",
                },
                target_lang="en",
                source_lang="zh",
            )

            out_doc = Document(str(out_path))
            paragraph_texts = [paragraph.text for paragraph in out_doc.paragraphs]
            self.assertIn("第一章 编制依据", paragraph_texts)
            self.assertIn("Chapter 1 Compilation Basis", paragraph_texts)
            self.assertNotIn("第1章 第一章 编制依据", paragraph_texts)
            self.assertNotIn("第1章 Chapter 1 Compilation Basis", paragraph_texts)

    def test_word_visible_numbering_prefix_detection_is_rule_based(self) -> None:
        self.assertTrue(_has_visible_numbering_prefix("第一章 编制依据"))
        self.assertTrue(_has_visible_numbering_prefix("第一节 工程概况"))
        self.assertTrue(_has_visible_numbering_prefix("1.施工图纸"))
        self.assertTrue(_has_visible_numbering_prefix("1.1主要材料概况"))
        self.assertTrue(_has_visible_numbering_prefix("一、工程概况"))
        self.assertTrue(_has_visible_numbering_prefix("（一）材料要求"))
        self.assertFalse(_has_visible_numbering_prefix("10mm控制线"))
        self.assertFalse(_has_visible_numbering_prefix("2#楼外墙"))
        self.assertFalse(_has_visible_numbering_prefix("2026年计划"))
        self.assertFalse(_has_visible_numbering_prefix("B20混凝土"))

    def test_scan_word_path_ignores_temp_and_generated_output_dirs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            legacy_path = root / "legacy.doc"
            source_path = root / "source.docx"
            temp_docx = root / "~$source.docx"
            output_dir = root / "source_翻译输出_20260513_120000"
            generated_path = output_dir / "generated.docx"
            output_dir.mkdir()

            self._build_sample_docx(source_path)
            self._build_sample_docx(generated_path)
            legacy_path.write_bytes(b"legacy word payload")
            temp_docx.write_text("not a real docx", encoding="utf-8")

            items = scan_word_path(root)

            self.assertEqual([item.path for item in items], [legacy_path, source_path])
            self.assertEqual(items[0].paragraph_count, 0)

    def test_write_bilingual_docx_uses_docx_output_name_for_legacy_doc(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            converted_path = temp_path / "legacy_temp.docx"
            output_dir = temp_path / "out"
            self._build_sample_docx(converted_path)

            out_path = write_bilingual_docx(
                source_path=converted_path,
                output_dir=output_dir,
                translations={"项目名称：测试工程": "Project name: Test Project"},
                target_lang="en",
                source_lang="zh",
                output_name="legacy.doc",
            )

            self.assertEqual(out_path.name, "双语(英文)_legacy.docx")
            self.assertTrue(out_path.exists())

    def test_word_retry_only_targets_unresolved_chinese_sources(self) -> None:
        self.assertTrue(
            _needs_word_translation_retry(
                "2、增设墙体厚度300mm",
                "2、增设墙体厚度300mm",
                source_lang="zh",
            )
        )
        self.assertTrue(
            _needs_word_translation_retry(
                "2、增设墙体厚度300mm",
                "",
                source_lang="zh",
            )
        )
        self.assertFalse(
            _needs_word_translation_retry(
                "Plan de construction pour la saison des pluies",
                "",
                source_lang="zh",
            )
        )
        self.assertFalse(
            _needs_word_translation_retry(
                "2、增设墙体厚度300mm",
                "2. Le mur ajouté aura une épaisseur de 300 mm.",
                source_lang="zh",
                target_lang="fr",
            )
        )
        self.assertTrue(
            _needs_word_translation_retry(
                "2、增设墙体厚度300mm",
                "2. Le mur ajouté 增设墙体 épaisseur 300 mm.",
                source_lang="zh",
                target_lang="fr",
            )
        )

    def test_word_strict_retry_repeats_until_translation_recovers(self) -> None:
        source = "短段落也应该通过多次重试恢复。"
        api_translations = {source: source}
        retry_settings = WordBatchSettings()
        retry_settings.max_paragraphs_per_batch = 1
        engine = RetryWordEngine(success_on_call=3)

        fixed_sources, unresolved_sources, recovery_reviews = _run_word_strict_retries(
            retry_sources=[source],
            api_translations=api_translations,
            engine=engine,
            target_lang="fr",
            retry_prompt="retry",
            retry_batch_settings=retry_settings,
            retry_attempts=3,
            source_lang="zh",
            should_stop=lambda: False,
        )

        self.assertEqual(fixed_sources, [source])
        self.assertEqual(unresolved_sources, [])
        self.assertEqual(recovery_reviews, {})
        self.assertEqual(len(engine.calls), 3)
        self.assertEqual(api_translations[source], f"Traduction valide {len(source)}")

    def test_word_strict_retry_reports_unresolved_after_attempts_exhausted(self) -> None:
        source = "短段落如果仍然返回原文则需要复核。"
        api_translations = {source: source}
        retry_settings = WordBatchSettings()
        retry_settings.max_paragraphs_per_batch = 1
        engine = RetryWordEngine(success_on_call=4)

        fixed_sources, unresolved_sources, recovery_reviews = _run_word_strict_retries(
            retry_sources=[source],
            api_translations=api_translations,
            engine=engine,
            target_lang="fr",
            retry_prompt="retry",
            retry_batch_settings=retry_settings,
            retry_attempts=3,
            source_lang="zh",
            should_stop=lambda: False,
        )

        self.assertEqual(fixed_sources, [])
        self.assertEqual(unresolved_sources, [source])
        self.assertEqual(recovery_reviews, {})
        self.assertEqual(len(engine.calls), 3)
        self.assertEqual(api_translations[source], source)

    def test_word_recovery_accepts_vnd_ten_thousand_amount_equivalence(self) -> None:
        source = (
            "6.7.4 乙方须严格遵守安全规定，违约行为将被处以"
            "每起违规行为 VND 2000万至 VND 3000万的罚款。"
        )
        translated = (
            "6.7.4 Party B shall strictly comply with safety regulations; "
            "each violation shall be subject to a fine from VND 20,000,000 "
            "to VND 30,000,000."
        )

        self.assertTrue(
            validate_translation(
                source,
                translated,
                target_lang="en",
                source_lang="zh",
                profile=VALIDATION_PROFILE_STRICT,
            ).is_fail
        )

        recovery = validate_translation(
            source,
            translated,
            target_lang="en",
            source_lang="zh",
            profile=VALIDATION_PROFILE_WORD_RECOVERY,
        )
        self.assertTrue(recovery.is_pass)

    def test_word_retry_soft_accepts_embedded_ocr_digit_with_review(self) -> None:
        source = "承包0商应被视为已将所有直接和间接成本计入合同金额。"
        translated = (
            "The contractor shall be deemed to have included all direct "
            "and indirect costs in the contract amount."
        )
        api_translations = {source: source}
        retry_settings = WordBatchSettings()
        retry_settings.max_paragraphs_per_batch = 1
        engine = MappingWordEngine({source: translated})

        fixed_sources, unresolved_sources, recovery_reviews = _run_word_strict_retries(
            retry_sources=[source],
            api_translations=api_translations,
            engine=engine,
            target_lang="en",
            retry_prompt="retry",
            retry_batch_settings=retry_settings,
            retry_attempts=3,
            source_lang="zh",
            should_stop=lambda: False,
        )

        self.assertEqual(fixed_sources, [source])
        self.assertEqual(unresolved_sources, [])
        self.assertIn(source, recovery_reviews)
        self.assertIn("承包0商", recovery_reviews[source].review_fragments)
        self.assertEqual(len(engine.calls), 1)
        self.assertEqual(api_translations[source], translated)

    def test_word_recovery_pool_semantic_accepts_valid_rejected_candidate(self) -> None:
        source = "签订时间 / Ngày：2026年 02月10日Ngày 10 tháng 02 năm 2026"
        candidate = "Signing Date / Date: February 10, 2026"
        engine = SemanticWordEngine({source: source}, verdict="equivalent")
        retry_settings = WordBatchSettings()
        retry_settings.max_paragraphs_per_batch = 1
        pool = _WordRecoveryPool(
            engine=engine,
            target_lang="en",
            retry_prompt="retry",
            retry_batch_settings=retry_settings,
            retry_attempts=1,
            source_lang="zh",
            api_scheduler=None,
            concurrency=2,
            should_stop=lambda: False,
            enable_semantic=True,
        )

        pool.add_candidate(source, candidate)
        outcome = pool.wait_for_completion()

        self.assertEqual(outcome.unresolved_sources, [])
        self.assertEqual(outcome.accepted_translations[source], candidate)
        self.assertIn(source, outcome.semantic_review_results)
        self.assertGreaterEqual(outcome.semantic_check_count, 1)
        self.assertTrue(engine.chat_calls)

    def test_word_recovery_pool_logs_locations_and_emits_summary(self) -> None:
        source = "签订时间 / Ngày：2026年 02月10日Ngày 10 tháng 02 năm 2026"
        candidate = "Signing Date / Date: February 10, 2026"
        engine = SemanticWordEngine({source: source}, verdict="equivalent")
        retry_settings = WordBatchSettings()
        retry_settings.max_paragraphs_per_batch = 1
        logs: list[tuple[str, str]] = []
        summaries = []
        pool = _WordRecoveryPool(
            engine=engine,
            target_lang="en",
            retry_prompt="retry",
            retry_batch_settings=retry_settings,
            retry_attempts=1,
            source_lang="zh",
            api_scheduler=None,
            concurrency=2,
            should_stop=lambda: False,
            log_callback=lambda level, message: logs.append((level, message)),
            status_callback=summaries.append,
            source_locations={
                source: [
                    {
                        "file": "方案.docx",
                        "section_path": "一、工程概况",
                        "location_label": "正文段落 8",
                    },
                    {
                        "file": "方案.docx",
                        "section_path": "二、表格",
                        "location_label": "表格 1 / 单元格 2",
                    },
                ]
            },
            enable_semantic=True,
        )

        pool.add_candidate(source, candidate)
        pool.wait_for_completion()

        log_text = "\n".join(message for _, message in logs)
        self.assertIn("方案.docx · 一、工程概况 · 正文段落 8 正在语义仲裁", log_text)
        self.assertIn("方案.docx · 二、表格 · 表格 1 / 单元格 2 正在语义仲裁", log_text)
        self.assertIn("方案.docx · 一、工程概况 · 正文段落 8 语义仲裁接受", log_text)
        self.assertTrue(any(summary.semantic_accepted_count == 2 for summary in summaries))

    def test_word_recovery_pool_keeps_review_when_semantic_is_uncertain(self) -> None:
        source = "签订时间 / Ngày：2026年 02月10日Ngày 10 tháng 02 năm 2026"
        candidate = "Signing Date / Date: February 10, 2026"
        engine = SemanticWordEngine({source: source}, verdict="uncertain")
        retry_settings = WordBatchSettings()
        retry_settings.max_paragraphs_per_batch = 1
        pool = _WordRecoveryPool(
            engine=engine,
            target_lang="en",
            retry_prompt="retry",
            retry_batch_settings=retry_settings,
            retry_attempts=1,
            source_lang="zh",
            api_scheduler=None,
            concurrency=2,
            should_stop=lambda: False,
            enable_semantic=True,
        )

        pool.add_candidate(source, candidate)
        outcome = pool.wait_for_completion()

        self.assertEqual(outcome.fixed_sources, [])
        self.assertEqual(outcome.unresolved_sources, [source])
        self.assertIn(source, outcome.unresolved_validation_results)

    def test_french_decimal_commas_pass_number_integrity_check(self) -> None:
        source = (
            "2、增设墙体厚度300mm，宽度8.925米，高度1.7米，"
            "底部标高-0.800米，顶部标高0.900米，预留200mm厚B30混凝土。"
        )
        translated = (
            "2. Ajouter un mur d'une épaisseur de 300 mm, largeur 8,925 m, "
            "hauteur 1,7 m, cote inférieure -0,800 m, cote supérieure 0,900 m, "
            "avec une réservation de 200 mm en béton B30."
        )

        self.assertFalse(
            is_translation_redundant(
                source,
                translated,
                target_lang="fr",
                source_lang="zh",
            )
        )

    def test_word_segments_keep_section_path(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "sections.docx"
            doc = Document()
            doc.add_heading("一、工程概况", level=1)
            doc.add_paragraph("施工范围：增设墙体。")
            doc.add_paragraph("（一）材料要求")
            doc.add_paragraph("混凝土强度等级为B30。")
            doc.save(str(source_path))

            segments = extract_word_segments(
                source_path,
                target_lang="fr",
                source_lang="zh",
            )
            paths = {segment.source: segment.section_path for segment in segments}

            self.assertEqual(paths["施工范围：增设墙体。"], "一、工程概况")
            self.assertEqual(paths["混凝土强度等级为B30。"], "一、工程概况 / （一）材料要求")

    def test_word_quality_report_records_location_and_excerpt(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir)
            source = "这是一个较长的施工段落，用于验证报告会记录开头和结尾，方便用户定位。"
            issue = {
                "file": "方案",
                "section_path": "一、工程概况 / （一）材料要求",
                "location_label": "正文段落 8",
                "snippet": _build_source_excerpt(source, head=10, tail=8),
                "problem": "初次翻译未获得有效译文",
                "status": "已自动单段重试并恢复译文。",
                "severity": "resolved",
                "review_fragments": ["承包0商"],
            }

            report_path = _write_word_quality_report(
                output_dir=output_dir,
                file_results=[{"name": "方案", "success": True}],
                issues=[issue],
                elapsed_sec=1.25,
                tm_hit_count=2,
                api_call_count=3,
            )

            self.assertIsNotNone(report_path)
            content = report_path.read_text(encoding="utf-8")
            self.assertIn("已自动处理", content)
            self.assertIn("一、工程概况 / （一）材料要求", content)
            self.assertIn("正文段落 8", content)
            self.assertIn(issue["snippet"], content)
            self.assertIn("问题片段：承包0商", content)

    def test_post_write_coverage_reports_untranslated_table_cells(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "table.docx"
            output_dir = temp_path / "out"

            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "砂浆名称"
            table.cell(0, 1).text = "专用砌筑砂浆"
            doc.save(str(source_path))

            out_path = write_bilingual_docx(
                source_path=source_path,
                output_dir=output_dir,
                translations={"砂浆名称": "Nom du mortier"},
                target_lang="fr",
                source_lang="zh",
            )
            issues: list[dict] = []

            residual_count = _append_post_write_coverage_issues(
                issues=issues,
                file_name="砂浆",
                output_path=out_path,
                target_lang="fr",
                source_lang="zh",
            )

            self.assertEqual(residual_count, 1)
            self.assertEqual(len(issues), 1)
            self.assertEqual(issues[0]["problem"], "输出文档仍存在未译源文")
            self.assertEqual(issues[0]["location_label"], "表格 1 / 单元格 2")
            self.assertIn("专用砌筑砂浆", issues[0]["snippet"])

    @staticmethod
    def _build_sample_docx(path: Path) -> None:
        doc = Document()
        doc.add_paragraph("项目名称：测试工程")
        doc.add_paragraph("12345")
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "设备"
        cell.add_paragraph("安装")
        doc.save(str(path))

    @staticmethod
    def _build_lower_letter_numbered_docx(path: Path) -> None:
        doc = Document()
        WordDocumentTests._set_default_num_id_9_level(
            doc,
            number_format="lowerLetter",
            level_text_value="%1.",
        )
        first = doc.add_paragraph("包括但不限于根据合同约定或确定的款项；")
        second = doc.add_paragraph("乙方未履行整改通知规定的违约补救义务；")
        for paragraph in (first, second):
            WordDocumentTests._set_paragraph_num_pr(paragraph, num_id="9", ilvl="0")
        doc.save(str(path))

    @staticmethod
    def _build_chapter_numbered_docx(path: Path) -> None:
        doc = Document()
        WordDocumentTests._set_default_num_id_9_level(
            doc,
            number_format="decimal",
            level_text_value="第%1章",
        )
        chapter = doc.add_paragraph("第一章 编制依据")
        WordDocumentTests._set_paragraph_num_pr(chapter, num_id="9", ilvl="0")
        doc.save(str(path))

    @staticmethod
    def _build_chinese_section_numbered_docx(path: Path) -> None:
        doc = Document()
        WordDocumentTests._set_default_num_id_9_level(
            doc,
            number_format="chineseCountingThousand",
            level_text_value="第%1节",
        )
        blank = doc.add_paragraph("")
        first = doc.add_paragraph("裂缝开裂现状及成因分析")
        second = doc.add_paragraph("修复总体原则与目标")
        for paragraph in (blank, first, second):
            WordDocumentTests._set_paragraph_num_pr(paragraph, num_id="9", ilvl="0")
        doc.save(str(path))

    @staticmethod
    def _build_decimal_chinese_section_numbered_docx(path: Path) -> None:
        doc = Document()
        WordDocumentTests._set_default_num_id_9_level(
            doc,
            number_format="decimal",
            level_text_value="第 %1 节",
        )
        first = doc.add_paragraph("裂缝开裂现状及成因分析")
        second = doc.add_paragraph("修复总体原则与目标")
        for paragraph in (first, second):
            WordDocumentTests._set_paragraph_num_pr(paragraph, num_id="9", ilvl="0")
        doc.save(str(path))

    @staticmethod
    def _build_nested_chinese_section_numbered_docx(path: Path) -> None:
        doc = Document()
        WordDocumentTests._set_num_id_9_levels(
            doc,
            {
                0: ("decimal", "第%1节"),
                1: ("decimal", "第%2节"),
            },
        )
        chapter = doc.add_paragraph("工程概况")
        section = doc.add_paragraph("裂缝开裂现状及成因分析")
        WordDocumentTests._set_paragraph_num_pr(chapter, num_id="9", ilvl="0")
        WordDocumentTests._set_paragraph_num_pr(section, num_id="9", ilvl="1")
        doc.save(str(path))

    @staticmethod
    def _set_default_num_id_9_level(
        doc: Document,
        *,
        number_format: str,
        level_text_value: str,
    ) -> None:
        WordDocumentTests._set_num_id_9_levels(doc, {0: (number_format, level_text_value)})

    @staticmethod
    def _set_num_id_9_levels(
        doc: Document,
        levels: dict[int, tuple[str, str]],
    ) -> None:
        numbering_root = doc.part.numbering_part.element
        target_abstract_id = None
        for num in numbering_root.findall(qn("w:num")):
            if num.get(qn("w:numId")) == "9":
                abstract_id = num.find(qn("w:abstractNumId"))
                target_abstract_id = abstract_id.get(qn("w:val")) if abstract_id is not None else None
                break
        if target_abstract_id is None:
            return
        for abstract_num in numbering_root.findall(qn("w:abstractNum")):
            if abstract_num.get(qn("w:abstractNumId")) != target_abstract_id:
                continue
            existing_levels = {
                int(level.get(qn("w:ilvl")) or 0): level
                for level in abstract_num.findall(qn("w:lvl"))
            }
            for ilvl, (number_format, level_text_value) in levels.items():
                level = existing_levels.get(ilvl)
                if level is None:
                    level = OxmlElement("w:lvl")
                    level.set(qn("w:ilvl"), str(ilvl))
                    abstract_num.append(level)
                num_format = level.find(qn("w:numFmt"))
                if num_format is None:
                    num_format = OxmlElement("w:numFmt")
                    level.append(num_format)
                num_format.set(qn("w:val"), number_format)
                level_text = level.find(qn("w:lvlText"))
                if level_text is None:
                    level_text = OxmlElement("w:lvlText")
                    level.append(level_text)
                level_text.set(qn("w:val"), level_text_value)
            return

    @staticmethod
    def _set_paragraph_num_pr(paragraph, *, num_id: str, ilvl: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl_element = OxmlElement("w:ilvl")
        ilvl_element.set(qn("w:val"), ilvl)
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), num_id)
        num_pr.append(ilvl_element)
        num_pr.append(num_id_element)
        p_pr.append(num_pr)

    @staticmethod
    def _paragraph_num_id(paragraph) -> str:
        p_pr = getattr(paragraph._p, "pPr", None)
        if p_pr is None:
            return ""
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            return ""
        num_id = num_pr.find(qn("w:numId"))
        return num_id.get(qn("w:val")) if num_id is not None else ""

    @staticmethod
    def _paragraph_first_line_indent(paragraph) -> str:
        p_pr = getattr(paragraph._p, "pPr", None)
        if p_pr is None:
            return ""
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            return ""
        return str(ind.get(qn("w:firstLine")) or "").strip()

    @staticmethod
    def _paragraph_first_line_chars(paragraph) -> str:
        p_pr = getattr(paragraph._p, "pPr", None)
        if p_pr is None:
            return ""
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            return ""
        return str(ind.get(qn("w:firstLineChars")) or "").strip()

    @staticmethod
    def _paragraph_left_indent(paragraph) -> str:
        p_pr = getattr(paragraph._p, "pPr", None)
        if p_pr is None:
            return ""
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            return ""
        return str(ind.get(qn("w:left")) or "").strip()

    @staticmethod
    def _paragraph_page_break_before(paragraph) -> bool | None:
        p_pr = getattr(paragraph._p, "pPr", None)
        if p_pr is None:
            return None
        page_break_before = p_pr.find(qn("w:pageBreakBefore"))
        if page_break_before is None:
            return None
        value = str(page_break_before.get(qn("w:val")) or "1").lower()
        return value not in {"0", "false", "off", "none"}

    @staticmethod
    def _paragraph_has_section_properties(paragraph) -> bool:
        p_pr = getattr(paragraph._p, "pPr", None)
        return p_pr is not None and p_pr.find(qn("w:sectPr")) is not None

    @staticmethod
    def _style_first_line_chars(paragraph) -> str:
        p_pr = paragraph.style.element.pPr
        if p_pr is None:
            return ""
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            return ""
        return str(ind.get(qn("w:firstLineChars")) or "").strip()

    @staticmethod
    def _paragraph_shading_fill(paragraph) -> str:
        p_pr = getattr(paragraph._p, "pPr", None)
        return WordDocumentTests._shading_fill(p_pr)

    @staticmethod
    def _cell_shading_fill(cell) -> str:
        tc_pr = getattr(cell._tc, "tcPr", None)
        return WordDocumentTests._shading_fill(tc_pr)

    @staticmethod
    def _paragraph_first_run_shading_fill(paragraph) -> str:
        run = next((item for item in paragraph.runs if item.text.strip()), None)
        if run is None:
            return ""
        return WordDocumentTests._shading_fill(getattr(run._element, "rPr", None))

    @staticmethod
    def _paragraph_first_run_highlight(paragraph) -> str:
        run = next((item for item in paragraph.runs if item.text.strip()), None)
        if run is None:
            return ""
        r_pr = getattr(run._element, "rPr", None)
        if r_pr is None:
            return ""
        highlight = r_pr.find(qn("w:highlight"))
        if highlight is None:
            return ""
        return str(highlight.get(qn("w:val")) or "").strip()

    @staticmethod
    def _shading_fill(parent_element) -> str:
        if parent_element is None:
            return ""
        shd = parent_element.find(qn("w:shd"))
        if shd is None:
            return ""
        return str(shd.get(qn("w:fill")) or "").strip().upper()

    @staticmethod
    def _set_shading_fill(parent_element, fill: str) -> None:
        shd = parent_element.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            parent_element.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)


if __name__ == "__main__":
    unittest.main(verbosity=2)
