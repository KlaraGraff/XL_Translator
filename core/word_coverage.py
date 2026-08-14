"""Word coverage detection and position-based untranslated-only writing."""

from __future__ import annotations

import shutil
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from core.language_registry import get_target_lang_display
from core.mixed_language import (
    MIXED_MARK_FOREIGN_NOISE,
    MIXED_MARK_UNRESOLVED,
)
from core.translation_coverage import (
    COVERAGE_AMBIGUOUS,
    COVERAGE_COVERED,
    COVERAGE_IGNORED,
    COVERAGE_SOURCE_ONLY,
    CoverageUnit,
    clean_coverage_text,
    coverage_summary,
    has_incidental_cjk,
    join_lines,
    looks_like_source_text,
    looks_like_target_text,
    residual_cjk_fragments,
    split_existing_bilingual_text,
)
from core.word_document import (
    EXISTING_HIGHLIGHT_POLICY_SKIP,
    WordFrontMatterBoundary,
    _append_translation_to_cell,
    _apply_cell_review_mark,
    _apply_paragraph_review_mark,
    _cell_has_review_highlight,
    _cell_source_text,
    _ensure_owner_writable,
    _insert_translation_paragraph_after,
    _is_toc_or_field_paragraph,
    _iter_unique_table_cells,
    _normalize_existing_highlight_policy,
    _normalize_review_mark_colors,
    _normalize_review_marks,
    _normalize_word_output_name,
    _paragraph_has_review_highlight,
    _paragraph_source_text,
    _review_mark_highlight_values,
    _sanitize_filename_fragment,
    find_word_front_matter_boundary,
)

# "未请求前置内容保护"时的占位边界：found 恒为 False，因此下面的分类逻辑天然
# 不会保护任何段落或表格——不需要另外一个 if protect_front_matter 分支。
_NO_FRONT_MATTER_PROTECTION = WordFrontMatterBoundary(body_start_index=None)


@dataclass
class WordCoveragePlan:
    path: Path
    units: list[CoverageUnit]
    front_matter: WordFrontMatterBoundary = field(default_factory=lambda: _NO_FRONT_MATTER_PROTECTION)

    @property
    def source_units(self) -> list[CoverageUnit]:
        return [unit for unit in self.units if unit.status == COVERAGE_SOURCE_ONLY]

    @property
    def residual_units(self) -> list[CoverageUnit]:
        """Covered positions whose translation still carries a few CJK characters."""
        return [unit for unit in self.units if unit.data.get("residual_cjk")]

    @property
    def source_texts(self) -> list[str]:
        seen: set[str] = set()
        texts: list[str] = []
        for unit in self.source_units:
            source = unit.source_text.strip()
            if source and source not in seen:
                seen.add(source)
                texts.append(source)
        return texts

    @property
    def summary(self) -> dict[str, int]:
        return coverage_summary(self.units)


def build_word_coverage_plan(
    path: str | Path,
    *,
    target_lang: str,
    source_lang: str = "zh",
    protect_front_matter: bool = False,
) -> WordCoveragePlan:
    """Classify app-style bilingual Word content by coverage status."""
    source_path = Path(path)
    doc = Document(str(source_path))
    front_matter = (
        find_word_front_matter_boundary(doc)
        if protect_front_matter
        else _NO_FRONT_MATTER_PROTECTION
    )
    units: list[CoverageUnit] = []
    units.extend(
        _classify_body_paragraphs(
            doc,
            target_lang=target_lang,
            source_lang=source_lang,
            front_matter=front_matter,
        )
    )
    units.extend(
        _classify_table_cells(
            doc,
            target_lang=target_lang,
            source_lang=source_lang,
            front_matter=front_matter,
        )
    )
    return WordCoveragePlan(path=source_path, units=units, front_matter=front_matter)


def write_untranslated_docx(
    *,
    source_path: str | Path,
    output_dir: str | Path,
    plan: WordCoveragePlan,
    translations: dict[str, str],
    target_lang: str,
    source_lang: str = "zh",
    output_name: str | None = None,
    review_marks: dict[str, str] | None = None,
    review_mark_colors: dict[str, str] | None = None,
    existing_highlight_policy: str = EXISTING_HIGHLIGHT_POLICY_SKIP,
    log_callback=None,
) -> Path:
    """Copy a Word document and insert translations only at source-only positions."""
    source_path = Path(source_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    review_mark_map = _normalize_review_marks(
        review_marks=review_marks,
        review_highlight_sources=None,
        review_highlight_color="",
    )
    review_color_map = _normalize_review_mark_colors(review_mark_colors)
    highlight_policy = _normalize_existing_highlight_policy(existing_highlight_policy)

    lang_display = _sanitize_filename_fragment(
        get_target_lang_display(target_lang, include_optional=True)
    )
    source_output_name = _normalize_word_output_name(output_name or source_path.name)
    out_path = output_dir / f"双语({lang_display})_{source_output_name}"
    shutil.copy2(source_path, out_path)
    _ensure_owner_writable(out_path)

    doc = Document(str(out_path))
    body_paragraphs = list(doc.paragraphs)
    table_cells = [
        cell
        for table in doc.tables
        for cell in _iter_unique_table_cells(table)
    ]

    paragraph_insertions = 0
    table_insertions = 0
    highlight_count = 0
    highlight_skip_count = 0
    for unit in reversed(plan.source_units):
        source_text = unit.source_text.strip()
        translation = str(translations.get(source_text) or "").strip()
        review_mark = review_mark_map.get(source_text)
        if not translation or translation.casefold() == source_text.casefold():
            continue
        kind = str(unit.kind or "")
        if kind == "paragraph":
            index = int(unit.data.get("paragraph_index", -1))
            if index < 0 or index >= len(body_paragraphs):
                continue
            paragraph = body_paragraphs[index]
            if _paragraph_source_text(paragraph) != source_text:
                continue
            translation_paragraph = _insert_translation_paragraph_after(
                paragraph,
                translation,
                target_lang=target_lang,
            )
            paragraph_insertions += 1
            if review_mark:
                # 原文段和刚插进去的译文段都涂上：用户在文档里翻到哪一边都看得见。
                marked = _apply_paragraph_review_mark(
                    paragraph,
                    review_mark,
                    highlight_policy,
                    review_color_map,
                )
                _apply_paragraph_review_mark(
                    translation_paragraph,
                    review_mark,
                    highlight_policy,
                    review_color_map,
                )
                if marked:
                    highlight_count += 1
                else:
                    highlight_skip_count += 1
        elif kind == "table_cell":
            index = int(unit.data.get("cell_index", -1))
            if index < 0 or index >= len(table_cells):
                continue
            cell = table_cells[index]
            if _cell_source_text(cell) != source_text:
                continue
            _append_translation_to_cell(
                cell,
                translation,
                target_lang=target_lang,
            )
            table_insertions += 1
            if review_mark:
                if _apply_cell_review_mark(
                    cell,
                    review_mark,
                    highlight_policy,
                    review_color_map,
                ):
                    highlight_count += 1
                else:
                    highlight_skip_count += 1

    doc.save(str(out_path))
    if log_callback:
        highlight_summary = ""
        if review_mark_map:
            highlight_summary = f"，复核标记 {highlight_count}"
            if highlight_skip_count:
                highlight_summary += f"，跳过已有标记 {highlight_skip_count}"
        log_callback(
            f"[OK] 已输出：{out_path.name}（补译段落 {paragraph_insertions}，"
            f"表格单元格 {table_insertions}{highlight_summary}）"
        )
    return out_path


def apply_coverage_review_marks(
    output_path: str | Path,
    *,
    plan: WordCoveragePlan,
    review_mark_colors: dict[str, str] | None = None,
    existing_highlight_policy: str = EXISTING_HIGHLIGHT_POLICY_SKIP,
) -> int:
    """把「成品文档体检」查出来的需复核位置，直接涂到成品文档上。

    报告里每写一条「需人工复核」，文件里就该有一处看得见的标记——否则用户拿到
    一张几十条的清单，却要自己在几百页里翻。体检只能在成品文档上做（要看的正是
    写完之后的结果），所以这一趟是"写完再涂"：按下标（段落序号／单元格序号）定位，
    不靠文字匹配，段落对里残留中文的那一段也能精确涂到。

    返回涂上的位置数。
    """
    output_path = Path(output_path)
    source_units = plan.source_units
    residual_units = plan.residual_units
    if not source_units and not residual_units:
        return 0

    color_map = _normalize_review_mark_colors(review_mark_colors)
    highlight_policy = _normalize_existing_highlight_policy(existing_highlight_policy)
    own_highlights = _review_mark_highlight_values(color_map)

    doc = Document(str(output_path))
    body_paragraphs = list(doc.paragraphs)
    table_cells = [
        cell
        for table in doc.tables
        for cell in _iter_unique_table_cells(table)
    ]

    marked = 0
    for unit, mark, index_key in (
        # 未译源文 → unresolved 色；译文里只残留零星中文（章节序号、单个汉字、日期片段
        # 都见过）是更轻的问题，换 foreign_noise 色，用户一眼分得清"整段没翻"和
        # "翻好了但留了几个字"。
        *((unit, MIXED_MARK_UNRESOLVED, None) for unit in source_units),
        *((unit, MIXED_MARK_FOREIGN_NOISE, "residual_index") for unit in residual_units),
    ):
        kind = str(unit.kind or "")
        if kind == "paragraph":
            raw_index = unit.data.get(index_key) if index_key else None
            if raw_index is None:
                raw_index = unit.data.get("paragraph_index", -1)
            index = int(raw_index)
            if index < 0 or index >= len(body_paragraphs):
                continue
            paragraph = body_paragraphs[index]
            if _paragraph_has_review_highlight(paragraph, own_highlights):
                continue
            if _apply_paragraph_review_mark(
                paragraph,
                mark,
                highlight_policy,
                color_map,
            ):
                marked += 1
        elif kind == "table_cell":
            index = int(unit.data.get("cell_index", -1))
            if index < 0 or index >= len(table_cells):
                continue
            cell = table_cells[index]
            if _cell_has_review_highlight(cell, own_highlights):
                continue
            if _apply_cell_review_mark(
                cell,
                mark,
                highlight_policy,
                color_map,
            ):
                marked += 1

    if marked:
        doc.save(str(output_path))
    return marked


def _mark_residual_cjk(
    data: dict,
    target_text: str,
    *,
    target_lang: str,
    location: str,
    index: int | None = None,
) -> None:
    """Record the CJK left inside an accepted translation, for a softer report item."""
    if not has_incidental_cjk(target_text, target_lang=target_lang):
        return
    fragments = residual_cjk_fragments(target_text)
    if not fragments:
        return
    data["residual_cjk"] = fragments
    data["residual_location"] = location
    data["residual_text"] = clean_coverage_text(target_text)
    # 残留位置往往不是这个单元本身：段落对里，残中文的是"下一段"那条译文。
    # 记下它的下标，写完体检那一趟才涂得准——否则只能拿文字去猜。
    if index is not None:
        data["residual_index"] = index


def _classify_body_paragraphs(
    doc: Document,
    *,
    target_lang: str,
    source_lang: str,
    front_matter: WordFrontMatterBoundary = _NO_FRONT_MATTER_PROTECTION,
) -> list[CoverageUnit]:
    units: list[CoverageUnit] = []
    paragraphs = list(doc.paragraphs)
    consumed_targets: set[int] = set()
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        text = _paragraph_source_text(paragraph)
        location = f"body.paragraph[{index}]"
        data = {"paragraph_index": index}
        if not text:
            index += 1
            continue
        if _is_toc_or_field_paragraph(paragraph):
            units.append(
                CoverageUnit(
                    source_text=text,
                    status=COVERAGE_IGNORED,
                    location=location,
                    kind="paragraph",
                    reason="目录或域段落默认跳过。",
                    data=data,
                )
            )
            index += 1
            continue
        if index in consumed_targets:
            index += 1
            continue
        if index in front_matter.protected_paragraph_indices:
            units.append(
                CoverageUnit(
                    source_text=text,
                    status=COVERAGE_IGNORED,
                    location=location,
                    kind="paragraph",
                    reason="前置内容保护：正文标题之前的封面/目录/前言默认跳过。",
                    data=data,
                )
            )
            index += 1
            continue
        if looks_like_source_text(text, source_lang=source_lang, target_lang=target_lang):
            next_index = index + 1
            next_text = (
                _paragraph_source_text(paragraphs[next_index])
                if next_index < len(paragraphs)
                else ""
            )
            if next_text and looks_like_target_text(
                next_text,
                source_lang=source_lang,
                target_lang=target_lang,
            ):
                _mark_residual_cjk(
                    data,
                    next_text,
                    target_lang=target_lang,
                    location=f"body.paragraph[{next_index}]",
                    index=next_index,
                )
                units.append(
                    CoverageUnit(
                        source_text=text,
                        target_text=next_text,
                        status=COVERAGE_COVERED,
                        location=location,
                        kind="paragraph",
                        reason="下一段为目标语言译文。",
                        data=data,
                    )
                )
                consumed_targets.add(next_index)
                index += 2
                continue
            units.append(
                CoverageUnit(
                    source_text=text,
                    status=COVERAGE_SOURCE_ONLY,
                    location=location,
                    kind="paragraph",
                    reason="源语言段落后未发现紧邻目标语言译文。",
                    data=data,
                )
            )
            index += 1
            continue
        if looks_like_target_text(text, source_lang=source_lang, target_lang=target_lang):
            _mark_residual_cjk(
                data, text, target_lang=target_lang, location=location, index=index
            )
            units.append(
                CoverageUnit(
                    source_text="",
                    target_text=text,
                    status=COVERAGE_IGNORED,
                    location=location,
                    kind="paragraph",
                    reason="段落看起来是目标语言译文，默认跳过。",
                    data=data,
                )
            )
            index += 1
            continue
        units.append(
            CoverageUnit(
                source_text=text,
                status=COVERAGE_IGNORED,
                location=location,
                kind="paragraph",
                reason="段落不符合补译候选规则。",
                data=data,
            )
        )
        index += 1
    return units


def _classify_table_cells(
    doc: Document,
    *,
    target_lang: str,
    source_lang: str,
    front_matter: WordFrontMatterBoundary = _NO_FRONT_MATTER_PROTECTION,
) -> list[CoverageUnit]:
    units: list[CoverageUnit] = []
    cell_index = 0
    for table_index, table in enumerate(doc.tables):
        protect_table = table_index in front_matter.protected_table_indices
        for cell in _iter_unique_table_cells(table):
            text = _cell_source_text(cell)
            location = f"table[{table_index}].cell[{cell_index}]"
            data = {"cell_index": cell_index, "table_index": table_index}
            if protect_table and clean_coverage_text(text):
                unit = CoverageUnit(
                    source_text=text,
                    status=COVERAGE_IGNORED,
                    location=location,
                    kind="table_cell",
                    section_path="表格",
                    reason="前置内容保护：正文标题之前的表格（常见于封面）默认跳过。",
                    data=data,
                )
            else:
                unit = _classify_cell_text(
                    text,
                    location=location,
                    data=data,
                    source_lang=source_lang,
                    target_lang=target_lang,
                )
            if unit is not None:
                units.append(unit)
            cell_index += 1
    return units


def _classify_cell_text(
    text: str,
    *,
    location: str,
    data: dict,
    source_lang: str,
    target_lang: str,
) -> CoverageUnit | None:
    cleaned = clean_coverage_text(text)
    if not cleaned:
        return None
    paragraph_lines = [line for line in cleaned.splitlines() if line.strip()]
    split = split_existing_bilingual_text(
        join_lines(paragraph_lines),
        source_lang=source_lang,
        target_lang=target_lang,
    )
    if split is not None:
        source, target = split
        _mark_residual_cjk(
            data,
            target,
            target_lang=target_lang,
            location=location,
            index=data.get("cell_index"),
        )
        return CoverageUnit(
            source_text=source,
            target_text=target,
            status=COVERAGE_COVERED,
            location=location,
            kind="table_cell",
            section_path="表格",
            reason="表格单元格内已包含源文和目标语言译文。",
            data=data,
        )
    if looks_like_source_text(cleaned, source_lang=source_lang, target_lang=target_lang):
        return CoverageUnit(
            source_text=cleaned,
            status=COVERAGE_SOURCE_ONLY,
            location=location,
            kind="table_cell",
            section_path="表格",
            reason="表格单元格包含源语言文本，未识别到目标语言译文。",
            data=data,
        )
    if looks_like_target_text(cleaned, source_lang=source_lang, target_lang=target_lang):
        _mark_residual_cjk(
            data,
            cleaned,
            target_lang=target_lang,
            location=location,
            index=data.get("cell_index"),
        )
        return CoverageUnit(
            source_text="",
            target_text=cleaned,
            status=COVERAGE_IGNORED,
            location=location,
            kind="table_cell",
            section_path="表格",
            reason="表格单元格看起来是目标语言译文，默认跳过。",
            data=data,
        )
    if len(paragraph_lines) > 1:
        return CoverageUnit(
            source_text=cleaned,
            status=COVERAGE_AMBIGUOUS,
            location=location,
            kind="table_cell",
            section_path="表格",
            reason="表格单元格无法可靠拆分源文和译文。",
            data=data,
        )
    return CoverageUnit(
        source_text=cleaned,
        status=COVERAGE_IGNORED,
        location=location,
        kind="table_cell",
        section_path="表格",
        reason="表格单元格不符合补译候选规则。",
        data=data,
    )
