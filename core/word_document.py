"""Word document scanning, extraction, and bilingual DOCX writing."""

from __future__ import annotations

import re
import shutil
import stat
import tempfile
import uuid
import zipfile
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from loguru import logger

from config import REVIEW_MARK_COLOR_DEFAULTS
from core.bilingual_writer import build_output_dir
from core.language_registry import get_target_lang_display
from core.mixed_language import (
    MIXED_MARK_FOREIGN_NOISE,
    MIXED_MARK_SEMANTIC,
    MIXED_MARK_UNRESOLVED,
)
from core.translation_filter import should_translate
from core.translation_protocol import extract_replace_translation, is_replace_translation
from core.word_converter import is_legacy_word_doc

SUPPORTED_WORD_SUFFIXES = {".docx", ".doc"}
GENERATED_OUTPUT_DIR_MARKER = "_翻译输出_"
EXISTING_HIGHLIGHT_POLICY_SKIP = "skip"
EXISTING_HIGHLIGHT_POLICY_OVERWRITE = "overwrite"
EXISTING_HIGHLIGHT_POLICY_RED_UNDERLINE = "red_underline"

_INVALID_FILENAME_FRAGMENT_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
_HEADING_STYLE_RE = re.compile(r"heading\s*(\d+)|标题\s*(\d+)", re.IGNORECASE)
_CHINESE_CHAPTER_RE = re.compile(r"^(?:第[一二三四五六七八九十百千万]+[章节篇]|[一二三四五六七八九十]+[、.．])")
_CHINESE_SECTION_RE = re.compile(r"^（[一二三四五六七八九十]+）")
_NUMBERED_SECTION_RE = re.compile(r"^\d+(?:\.\d+)+\s+")
_CHINESE_NUMBER_CHARS = "零〇一二两三四五六七八九十百千万"
_CHINESE_NUMBERING_PREFIX_RE = re.compile(
    rf"^第[{_CHINESE_NUMBER_CHARS}0-9０-９]+[章节篇卷部]"
)
_CHINESE_LIST_PREFIX_RE = re.compile(
    rf"^(?:[{_CHINESE_NUMBER_CHARS}]+[、.．]|[（(][{_CHINESE_NUMBER_CHARS}]+[）)])"
)
_ARABIC_LIST_PREFIX_RE = re.compile(
    r"^(?:\d{1,3}[、．.)）]|[（(]\d{1,3}[）)])"
)
_ALPHA_LIST_PREFIX_RE = re.compile(r"^[A-Za-z][.)）]")
_BULLET_PREFIX_RE = re.compile(r"^[•·▪▫●○◆◇■□]\s+")
_ARABIC_DECIMAL_PREFIX_RE = re.compile(r"^(\d{1,3}(?:\.\d{1,3})+)(.*)$")
_LEADING_INLINE_WHITESPACE_RE = re.compile(r"^[ \t\u3000]+")
# 无小数点的顶层数字标题，例如「1 概述」「1. 概述」「1、概述」。孤立的前导数字本身有
# 歧义——普通编号列表项（如「1. 检查产品合格证」）也是这个形状，所以额外要求分隔符后
# 的标题正文不含句末标点且长度受限，把典型的长句列表项挡在外面。这仍然是启发式规则，
# 挡不住「1 检查工作面清洁情况」这种又短又不带标点的列表项，代价见
# find_word_front_matter_boundary 的注释。
_BARE_NUMERIC_HEADING_RE = re.compile(r"^(\d{1,2})[ \u3000、.．](?!\d)([^。！？；;]{1,20})$")
# 手打目录条目：标题文字 + 点划线（或中圆点）+ 页码，例如「第一章 工程概况 ...... 1」。
# 自动生成的目录段落靠 _is_toc_or_field_paragraph 就能认出来，但工程文档的目录很多是
# 手工敲的，没有任何样式或域标记——扫描正文标题时若不排除它，目录里抄的那行「第一章
# ……」会被当成正文标题本身，结果只保护了封面、整份目录反倒被当正文翻译，与这个选项的
# 初衷正好相反。省略号「…」「⋯」单个字符就已经代表三个点，出现一个即可；两点引导符
# 「‥」信息量弱一些，要求 2 个以上才采信。
_HAND_TYPED_TOC_LEADER_RE = re.compile(
    r"(?:(?:[.．]\s*){3,}|(?:[·]\s*){3,}|(?:[…⋯]\s*){1,}|(?:[‥]\s*){2,})\s*\d+\s*$"
)

# 整段内容都由 Word 自己生成的域：目录、索引、题录、页码。这类段落翻译了也会被域
# 刷新覆盖，整段跳过。SEQ（题注编号）、STYLEREF（章号）、REF 不在此列——含这些域的
# 表题注、图题注属于正文，早期版本「见到任何域就跳过」会让它们一个字不翻还不进报告。
_GENERATED_BLOCK_FIELD_RE = re.compile(
    r"\b(?:TOC|TOA|INDEX|XE|TC|BIBLIOGRAPHY|PAGE|PAGEREF|NUMPAGES|SECTIONPAGES)\b"
)
# 去掉域结果后，段落自带的文字要有这么多才算「有正文要翻译」——挡住「详见 {REF}」
# 这种去掉域就只剩两个字的引用句。
_FIELD_PARAGRAPH_MIN_CJK_CHARS = 4
_FIELD_PARAGRAPH_MIN_LATIN_LETTERS = 8

_HEADER_FOOTER_PART_RE = re.compile(r"word/(?:header|footer)\d*\.xml$")

_WORD_HIGHLIGHT_RGB = {
    "black": (0x00, 0x00, 0x00),
    "blue": (0x00, 0x00, 0xFF),
    "cyan": (0x00, 0xFF, 0xFF),
    "green": (0x00, 0xFF, 0x00),
    "magenta": (0xFF, 0x00, 0xFF),
    "red": (0xFF, 0x00, 0x00),
    "yellow": (0xFF, 0xFF, 0x00),
    "white": (0xFF, 0xFF, 0xFF),
    "darkBlue": (0x00, 0x00, 0x80),
    "darkCyan": (0x00, 0x80, 0x80),
    "darkGreen": (0x00, 0x80, 0x00),
    "darkMagenta": (0x80, 0x00, 0x80),
    "darkRed": (0x80, 0x00, 0x00),
    "darkYellow": (0x80, 0x80, 0x00),
    "darkGray": (0x80, 0x80, 0x80),
    "lightGray": (0xC0, 0xC0, 0xC0),
}
_WORD_HIGHLIGHT_HEX_OVERRIDES = {
    "FFF2CC": "yellow",
    "FCE4D6": "yellow",
    "F4CCCC": "red",
    "DDEBFF": "cyan",
    "D9EAD3": "green",
}
_CHINESE_NUMBER_FORMATS = {
    "chineseCounting",
    "chineseCountingThousand",
    "chineseLegalSimplified",
    "chineseCountingThousand2",
}
_CHINESE_DIGITS = "零一二三四五六七八九"


@dataclass
class WordFileItem:
    path: Path
    name: str
    size_kb: float
    # Legacy .doc files cannot be inspected faithfully before conversion.
    # ``None`` is intentional: a zero here would misrepresent an unknown
    # document as an empty one.
    paragraph_count: int | None = None
    table_count: int | None = None
    translatable_count: int = 0
    original_path: Path | None = None
    relative_path: str = ""
    format: str = "docx"
    needs_conversion: bool = False
    statistics_status: str = "known"
    risk: dict[str, object] = field(default_factory=dict)


@dataclass(frozen=True)
class WordScanSkippedItem:
    """A selectable-looking Word source which failed scan-time inspection."""

    path: Path
    relative_path: str
    reason: str
    format: str = ""


@dataclass
class WordScanResult:
    """Typed Word source scan used by the API/UI task-start contract."""

    root: Path
    items: list[WordFileItem] = field(default_factory=list)
    skipped: list[WordScanSkippedItem] = field(default_factory=list)

    @property
    def summary(self) -> dict[str, int]:
        known_items = [item for item in self.items if item.statistics_status == "known"]
        doc_count = sum(1 for item in self.items if item.format == "doc")
        return {
            "scanned_count": len(self.items),
            "selected_count": len(self.items),
            "paragraph_count": sum(int(item.paragraph_count or 0) for item in known_items),
            "table_count": sum(int(item.table_count or 0) for item in known_items),
            "doc_count": doc_count,
            "doc_unknown_count": sum(
                1 for item in self.items if item.statistics_status != "known"
            ),
            "skipped_count": len(self.skipped),
        }

    @property
    def risk(self) -> dict[str, object]:
        doc_count = self.summary["doc_count"]
        return {
            "has_doc": bool(doc_count),
            "doc_count": doc_count,
            "requires_explicit_compatibility_confirmation": bool(doc_count),
            "message": (
                "检测到 .doc 文件：优先使用本机 Microsoft Word 高保真转换；"
                "只有明确确认兼容转换后才会改用 LibreOffice 或 macOS textutil，"
                "后者可能改变版式、域、图文或宏。"
                if doc_count
                else ""
            ),
        }


@dataclass(frozen=True)
class WordSegment:
    source: str
    kind: str
    location: str
    section_path: str = ""


@dataclass(frozen=True)
class ResolvedWordTranslation:
    text: str
    replace_only: bool = False


@dataclass(frozen=True)
class NumberingLevelDefinition:
    start: int = 1
    number_format: str = "decimal"
    level_text: str = "%1."
    legal_numbering: bool = False


@dataclass(frozen=True)
class WordNumberingNormalizationStats:
    labels_seen: int = 0
    labels_prepended: int = 0
    numbering_removed: int = 0


@dataclass(frozen=True)
class WordNumberingNormalizationResult:
    path: Path
    stats: WordNumberingNormalizationStats


@dataclass(frozen=True)
class WordHiddenContentReport:
    """python-docx 看不见、因而必然漏译的正文内容统计。

    python-docx 只把 ``w:p`` / ``w:tbl`` 的直接子节点当成文档内容，段落文本又只由
    ``w:r`` 与 ``w:hyperlink`` 拼成。被内容控件（``w:sdt``）或未接受的修订插入
    （``w:ins``）包住的文字因此完全不在扫描范围内——不报错、不计数，用户拿到的译文
    里就是缺了这几段，没有任何线索。这个报告的唯一职责是把"缺了多少"变成可见信息。
    """

    content_control_count: int = 0
    tracked_insertion_count: int = 0

    @property
    def total(self) -> int:
        return self.content_control_count + self.tracked_insertion_count

    @property
    def found(self) -> bool:
        return self.total > 0

    def describe(self) -> str:
        parts = []
        if self.content_control_count:
            parts.append(f"内容控件 {self.content_control_count} 处")
        if self.tracked_insertion_count:
            parts.append(f"未接受的修订插入 {self.tracked_insertion_count} 处")
        return "、".join(parts)

    def as_dict(self) -> dict[str, int]:
        return {
            "content_control_count": self.content_control_count,
            "tracked_insertion_count": self.tracked_insertion_count,
            "total": self.total,
        }


def is_supported_word_file(path: str | Path) -> bool:
    """Return whether a path points to a supported Word file."""
    path = Path(path)
    return (
        path.is_file()
        and path.suffix.lower() in SUPPORTED_WORD_SUFFIXES
        and not path.name.startswith("~")
    )


def scan_word_path(path: str | Path) -> list[WordFileItem]:
    """Compatibility entry point returning only selectable Word files."""
    return scan_word_sources(path).items


def scan_word_sources(path: str | Path) -> WordScanResult:
    """Scan Word input while preserving skipped-file and unknown-stat evidence."""
    source = Path(path).expanduser()
    root = source if source.is_dir() else source.parent
    result = WordScanResult(root=root)
    if not source.exists():
        reason = f"路径不存在：{source}"
        logger.warning(reason)
        result.skipped.append(WordScanSkippedItem(source, source.name, reason))
        return result

    if source.is_file():
        _scan_one_word_file(source, source.parent, result)
    elif source.is_dir():
        for candidate in sorted(source.rglob("*")):
            if not candidate.is_file():
                continue
            # Office lock files and Translator's own output folders are not
            # user-visible failed inputs: they are deliberately excluded from
            # the selectable source boundary.
            if candidate.name.startswith("~"):
                continue
            relative = _relative_word_path(candidate, source)
            if _is_generated_output(Path(relative)):
                continue
            if candidate.suffix.lower() not in SUPPORTED_WORD_SUFFIXES:
                continue
            _scan_one_word_file(candidate, source, result)
    else:
        result.skipped.append(
            WordScanSkippedItem(
                source,
                source.name,
                f"路径既不是文件也不是目录：{source}",
            )
        )

    result.items.sort(key=lambda item: str(item.path))
    result.skipped.sort(key=lambda item: str(item.path))
    logger.info(
        f"Word 扫描完成：{source}，可选 {len(result.items)}，跳过 {len(result.skipped)}"
    )
    return result


def scan_word_folder(root: str | Path) -> list[WordFileItem]:
    """Recursively scan a folder for supported Word files."""
    return scan_word_sources(root).items


def extract_word_segments(
    path: str | Path,
    *,
    target_lang: str,
    source_lang: str = "zh",
    protect_front_matter: bool = False,
) -> list[WordSegment]:
    """Extract unique body-paragraph and table-cell texts that need translation."""
    doc = Document(str(path))
    front_matter = find_word_front_matter_boundary(doc) if protect_front_matter else None
    protected_paragraphs = front_matter.protected_paragraph_indices if front_matter else frozenset()
    protected_tables = front_matter.protected_table_indices if front_matter else frozenset()
    seen: set[str] = set()
    segments: list[WordSegment] = []
    section_stack: dict[int, str] = {}

    for index, paragraph in enumerate(doc.paragraphs):
        if index in protected_paragraphs:
            continue
        source = _paragraph_source_text(paragraph)
        heading_level = _detect_heading_level(paragraph)
        if heading_level is not None and source:
            _update_section_stack(section_stack, heading_level, source)

        if not _is_translatable_source(
            source,
            target_lang=target_lang,
            source_lang=source_lang,
        ):
            continue
        if _is_toc_or_field_paragraph(paragraph):
            continue
        if source in seen:
            continue
        seen.add(source)
        segments.append(
            WordSegment(
                source=source,
                kind="paragraph",
                location=f"body.paragraph[{index}]",
                section_path=_format_section_path(section_stack),
            )
        )

    for table_index, table in enumerate(doc.tables):
        if table_index in protected_tables:
            continue
        for cell_index, cell in enumerate(_iter_unique_table_cells(table)):
            source = _cell_source_text(cell)
            if not _is_translatable_source(
                source,
                target_lang=target_lang,
                source_lang=source_lang,
            ):
                continue
            if source in seen:
                continue
            seen.add(source)
            segments.append(
                WordSegment(
                    source=source,
                    kind="table_cell",
                    location=f"table[{table_index}].cell[{cell_index}]",
                    section_path=f"表格 {table_index + 1}",
                )
            )

    return segments


def count_text_bearing_header_footer_parts(source: str | Path) -> int:
    """Count header/footer parts that carry real words (page numbers don't count).

    页眉页脚不参与翻译（见 extract_word_segments 的说明），报告里要如实写明这件事，
    但只在文档确实有页眉页脚文字时才提，避免每份报告都挂一句无关提示。直接读 zip
    里的 XML，不用 python-docx 再解析一遍整份文档。
    """
    try:
        with zipfile.ZipFile(str(source)) as archive:
            part_names = [
                name
                for name in archive.namelist()
                if _HEADER_FOOTER_PART_RE.match(name)
            ]
            count = 0
            for name in part_names:
                xml = archive.read(name).decode("utf-8", errors="ignore")
                texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, flags=re.DOTALL)
                if any(any(char.isalpha() for char in text) for text in texts):
                    count += 1
            return count
    except Exception:  # noqa: BLE001 - 报告里的说明性信息，读不到就不提。
        return 0


def detect_hidden_word_content(source: str | Path) -> WordHiddenContentReport:
    """统计文档里 python-docx 扫不到、因而会被静默漏译的内容。

    只看 body：页眉页脚本来就不在翻译范围内，把它们算进来只会制造无从处理的告警。
    """
    try:
        doc = Document(str(Path(source)))
    except Exception as exc:  # detection must never break a translatable file
        logger.warning(f"检测 Word 隐藏内容失败 {Path(source).name}：{exc}")
        return WordHiddenContentReport()
    return _detect_hidden_word_content(doc)


def _detect_hidden_word_content(doc: Document) -> WordHiddenContentReport:
    body = doc.element.body
    return WordHiddenContentReport(
        content_control_count=sum(
            1
            for element in _iter_outermost_elements(body, "w:sdt")
            if _element_has_visible_text(element)
        ),
        tracked_insertion_count=sum(
            1
            for element in _iter_outermost_elements(body, "w:ins")
            # w:ins 也用来标记"这一段的段落标记是新插入的"（w:pPr/w:rPr/w:ins），
            # 那种节点一个 w:t 都没有，不代表任何漏译的正文。
            if _element_has_visible_text(element)
            and not _has_ancestor_tag(element, "w:sdt", body)
        ),
    )


def _iter_outermost_elements(root, tag: str):
    """产出 root 下同名嵌套里最外层的那一个，避免嵌套结构被重复计数。"""
    qualified = qn(tag)
    for element in root.iter(qualified):
        if not _has_ancestor_tag(element, tag, root):
            yield element


def _has_ancestor_tag(element, tag: str, root) -> bool:
    qualified = qn(tag)
    parent = element.getparent()
    while parent is not None and parent is not root:
        if parent.tag == qualified:
            return True
        parent = parent.getparent()
    return False


def _element_has_visible_text(element) -> bool:
    return any((node.text or "").strip() for node in element.iter(qn("w:t")))


def write_bilingual_docx(
    *,
    source_path: str | Path,
    output_dir: str | Path,
    translations: dict[str, str],
    target_lang: str,
    source_lang: str = "zh",
    output_name: str | None = None,
    review_highlight_sources: set[str] | None = None,
    review_highlight_color: str = "FFF2CC",
    review_marks: dict[str, str] | None = None,
    review_mark_colors: dict[str, str] | None = None,
    existing_highlight_policy: str = EXISTING_HIGHLIGHT_POLICY_SKIP,
    log_callback=None,
    protect_front_matter: bool = False,
) -> Path:
    """Write a bilingual Word document to the output directory."""
    source_path = Path(source_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    review_mark_map = _normalize_review_marks(
        review_marks=review_marks,
        review_highlight_sources=review_highlight_sources,
        review_highlight_color=review_highlight_color,
    )
    review_color_map = _normalize_review_mark_colors(review_mark_colors)
    highlight_policy = _normalize_existing_highlight_policy(existing_highlight_policy)

    lang_display = _sanitize_filename_fragment(
        get_target_lang_display(target_lang, include_optional=True)
    )
    source_output_name = _normalize_word_output_name(output_name or source_path.name)
    out_path = output_dir / f"双语({lang_display})_{source_output_name}"
    shutil.copy2(source_path, out_path)
    _ensure_owner_writable(out_path)

    doc = Document(str(out_path))
    body_paragraphs = list(doc.paragraphs)
    front_matter = find_word_front_matter_boundary(doc) if protect_front_matter else None
    protected_paragraphs = front_matter.protected_paragraph_indices if front_matter else frozenset()
    # 表格没有"下标即位置"这种便利——_iter_unique_table_cells 会递归展开嵌套表格，
    # 所以按表格下标挑出受保护的表，再记下它产出的单元格。这里必须和下面遍历用的是
    # 同一批 _Cell 对象、并且由 table_cells 一直持有强引用：python-docx 每次访问都新建
    # 一个包装对象（doc.tables[0].cell(0, 0) is doc.tables[0].cell(0, 0) 为 False），
    # 先单独建一份 id 集合、再另建一份列表的话，集合里存的全是随即被回收的对象地址，
    # 命中与否只看内存分配器的运气——漏保护和误保护正文单元格都会发生。
    protected_tables = front_matter.protected_table_indices if front_matter else frozenset()
    table_cells: list = []
    protected_table_cell_ids: set[int] = set()
    # 段落身份同理：cell.paragraphs 每次访问都新建包装对象，必须把这一批留在列表里
    # 一直持有强引用，下面按 id() 判定"是否受保护"才有意义。
    cell_paragraphs: list[Paragraph] = []
    protected_paragraph_ids: set[int] = set()
    for table_index, table in enumerate(doc.tables):
        table_protected = table_index in protected_tables
        for cell in _iter_unique_table_cells(table):
            table_cells.append(cell)
            own_paragraphs = list(cell.paragraphs)
            cell_paragraphs.extend(own_paragraphs)
            if table_protected:
                protected_table_cell_ids.add(id(cell))
                protected_paragraph_ids.update(id(item) for item in own_paragraphs)
    for protected_index in protected_paragraphs:
        if 0 <= protected_index < len(body_paragraphs):
            protected_paragraph_ids.add(id(body_paragraphs[protected_index]))
    all_paragraphs = _order_paragraphs_by_document_position(
        doc,
        [*body_paragraphs, *cell_paragraphs],
    )
    numbering_labels = _collect_numbering_labels(doc, all_paragraphs)
    # 编号标签仍然按整篇文档统计（受保护段落也要参与计数，否则正文的序号会从头开始），
    # 但真正落笔改写只允许发生在保护边界之外——"保护封面和目录"承诺的是不改写，
    # 把封面的自动编号扁平化成正文文字并删掉 numPr 同样是改写。
    _flatten_automatic_numbering(
        all_paragraphs,
        numbering_labels,
        protected_paragraph_ids=protected_paragraph_ids,
    )
    original_paragraph_sources = {
        id(paragraph): _paragraph_source_text(paragraph)
        for paragraph in all_paragraphs
    }
    original_paragraph_prefixes = {
        id(paragraph): _paragraph_leading_inline_whitespace(paragraph)
        for paragraph in all_paragraphs
    }
    original_cell_sources = {
        id(cell): _cell_source_text(cell)
        for cell in table_cells
    }

    paragraph_insertions = 0
    table_insertions = 0
    highlight_count = 0
    highlight_skip_count = 0

    for paragraph_index, paragraph in enumerate(body_paragraphs):
        if paragraph_index in protected_paragraphs:
            continue
        paragraph_key = id(paragraph)
        source = original_paragraph_sources.get(
            paragraph_key,
            _paragraph_source_text(paragraph),
        )
        if not _is_translatable_source(
            source,
            target_lang=target_lang,
            source_lang=source_lang,
        ):
            continue
        if _is_toc_or_field_paragraph(paragraph):
            continue

        review_mark = review_mark_map.get(source.strip())
        if review_mark:
            if _apply_paragraph_review_mark(
                paragraph,
                review_mark,
                highlight_policy,
                review_color_map,
            ):
                highlight_count += 1
            else:
                highlight_skip_count += 1

        resolved = _resolve_translation(source, translations)
        if resolved is None:
            continue
        leading_prefix = original_paragraph_prefixes.get(paragraph_key, "")
        translated_text = _apply_leading_prefix(resolved.text, leading_prefix)
        # 原地改写会清空非锚点 run，段落里若有域（题注的 SEQ 编号）会被抹掉，
        # 这类段落一律改成在下一行插入译文。
        if resolved.replace_only and not _paragraph_has_field(paragraph):
            _replace_paragraph_text(
                paragraph,
                translated_text,
                target_lang=target_lang,
            )
            if review_mark:
                _apply_paragraph_review_mark(
                    paragraph,
                    review_mark,
                    highlight_policy,
                    review_color_map,
                )
        else:
            translation_paragraph = _insert_translation_paragraph_after(
                paragraph,
                translated_text,
                target_lang=target_lang,
            )
            if review_mark:
                _apply_paragraph_review_mark(
                    translation_paragraph,
                    review_mark,
                    highlight_policy,
                    review_color_map,
                )
        paragraph_insertions += 1

    for cell in table_cells:
        if id(cell) in protected_table_cell_ids:
            continue
        source = original_cell_sources.get(id(cell), _cell_source_text(cell))
        if not _is_translatable_source(
            source,
            target_lang=target_lang,
            source_lang=source_lang,
        ):
            continue
        review_mark = review_mark_map.get(source.strip())
        resolved = _resolve_translation(source, translations)
        if resolved is None:
            if review_mark:
                if _apply_cell_review_mark(
                    cell,
                    review_mark,
                    highlight_policy,
                    review_color_map,
                ):
                    highlight_count += 1
                else:
                    highlight_skip_count += 1
            continue
        if resolved.replace_only:
            _replace_cell_text(cell, resolved.text, target_lang=target_lang)
        else:
            _append_translation_to_cell(
                cell,
                resolved.text,
                target_lang=target_lang,
            )
        if review_mark:
            if _apply_cell_review_mark(
                cell,
                review_mark,
                highlight_policy,
                review_color_map,
            ):
                highlight_count += 1
            else:
                highlight_skip_count += 1
        table_insertions += 1

    _trim_trailing_empty_body_paragraphs(doc)
    doc.save(str(out_path))
    if log_callback:
        highlight_summary = ""
        if review_mark_map:
            highlight_summary = f"，复核标记 {highlight_count}"
            if highlight_skip_count:
                highlight_summary += f"，跳过已有标记 {highlight_skip_count}"
        log_callback(
            f"[OK] 已输出：{out_path.name}（段落 {paragraph_insertions}，表格单元格 {table_insertions}{highlight_summary}）"
        )
    return out_path


def build_word_output_dir(
    source_dir: str | Path,
    custom_output_dir: str | Path | None = None,
) -> Path:
    """Expose the shared output-directory convention for Word callers."""
    return build_output_dir(source_dir, custom_output_dir)


def normalize_docx_automatic_numbering(
    source_path: str | Path,
    output_path: str | Path | None = None,
) -> WordNumberingNormalizationResult:
    """Materialize automatic numbering once, then suppress residual list metadata."""
    source_path = Path(source_path)
    target_path = Path(output_path) if output_path is not None else _temp_docx_path(source_path)
    if source_path.resolve() != target_path.resolve():
        target_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_path, target_path)

    doc = Document(str(target_path))
    table_cells = [
        cell
        for table in doc.tables
        for cell in _iter_unique_table_cells(table)
    ]
    all_paragraphs = _order_paragraphs_by_document_position(
        doc,
        [
            *list(doc.paragraphs),
            *(paragraph for cell in table_cells for paragraph in cell.paragraphs),
        ],
    )
    numbering_labels = _collect_numbering_labels(doc, all_paragraphs)
    stats = _flatten_automatic_numbering(all_paragraphs, numbering_labels)
    doc.save(str(target_path))
    return WordNumberingNormalizationResult(path=target_path, stats=stats)


def _scan_one_word_file(path: Path, root: Path, result: WordScanResult) -> None:
    if not is_supported_word_file(path):
        result.skipped.append(
            WordScanSkippedItem(
                path=path,
                relative_path=_relative_word_path(path, root),
                reason="不支持的 Word 文件或 Office 临时文件",
                format=path.suffix.lower().lstrip("."),
            )
        )
        return
    try:
        result.items.append(_build_word_file_item(path, root=root))
    except Exception as exc:  # scan errors must remain visible and non-fatal
        reason = f"读取失败：{exc}"
        logger.warning(f"扫描 Word 文件失败 {path.name}：{exc}")
        result.skipped.append(
            WordScanSkippedItem(
                path=path,
                relative_path=_relative_word_path(path, root),
                reason=reason,
                format=path.suffix.lower().lstrip("."),
            )
        )


def _relative_word_path(path: Path, root: Path) -> str:
    try:
        return str(path.relative_to(root))
    except ValueError:
        return path.name


def _build_word_file_item(path: Path, *, root: Path | None = None) -> WordFileItem:
    if is_legacy_word_doc(path):
        return WordFileItem(
            path=path,
            name=path.stem,
            size_kb=round(path.stat().st_size / 1024, 1),
            original_path=path,
            relative_path=_relative_word_path(path, root or path.parent),
            format="doc",
            needs_conversion=True,
            statistics_status="conversion_required",
            risk={
                "compatibility_required": True,
                "message": (
                    ".doc 的段落和表格数需在转换为临时 .docx 后统计；"
                    "高保真不可用时必须先确认兼容转换风险。"
                ),
            },
        )

    doc = Document(str(path))
    segments = extract_word_segments(path, target_lang="en", source_lang="zh")
    return WordFileItem(
        path=path,
        name=path.stem,
        size_kb=round(path.stat().st_size / 1024, 1),
        paragraph_count=len(doc.paragraphs),
        table_count=len(doc.tables),
        translatable_count=len(segments),
        relative_path=_relative_word_path(path, root or path.parent),
        format="docx",
        needs_conversion=False,
        statistics_status="known",
    )


def _normalize_word_output_name(name: str) -> str:
    cleaned = _sanitize_filename_fragment(Path(str(name or "document")).name)
    if Path(cleaned).suffix.lower() == ".doc":
        return f"{Path(cleaned).stem}.docx"
    if Path(cleaned).suffix.lower() != ".docx":
        return f"{cleaned}.docx"
    return cleaned


def _temp_docx_path(original_path: Path) -> Path:
    temp_dir = Path(tempfile.gettempdir()) / "word_translator_temp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir / f"{original_path.stem}_normalized_{uuid.uuid4().hex[:8]}.docx"


def _is_generated_output(path: Path) -> bool:
    return any(GENERATED_OUTPUT_DIR_MARKER in part for part in path.parts)


def _iter_unique_table_cells(table: Table, seen: set | None = None):
    if seen is None:
        seen = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            yield cell
            for nested_table in cell.tables:
                yield from _iter_unique_table_cells(nested_table, seen)


def _iter_cell_direct_paragraphs(cell: _Cell):
    for child in cell._tc.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, cell)


def _order_paragraphs_by_document_position(
    doc: Document,
    paragraphs: list[Paragraph],
) -> list[Paragraph]:
    """把正文段落与表格内段落重排成 XML 里的真实先后顺序。

    ``doc.paragraphs`` 与 ``doc.tables`` 是两条互不相干的序列，把它们首尾相接等于
    断言"所有表格都排在全部正文之后"。自动编号的计数器正是按这个顺序推进的，于是
    "正文列表 → 表格里的列表 → 正文列表"（共用同一个 numId）算出来的序号会整体错位。
    body 的文档序才是 Word 自己的计数顺序。

    返回的是传入的同一批包装对象——上层用 ``id(paragraph)`` 关联原文、编号标签与保护
    边界，这里换新对象会让那些映射全部落空。
    """
    wrapper_by_element: dict = {}
    for paragraph in paragraphs:
        wrapper_by_element.setdefault(paragraph._p, paragraph)

    ordered: list[Paragraph] = []
    for element in doc.element.body.iter(qn("w:p")):
        wrapper = wrapper_by_element.pop(element, None)
        if wrapper is not None:
            ordered.append(wrapper)

    if wrapper_by_element:
        # 理论上不会发生（传进来的段落都取自这棵树）。真发生了也绝不能把这些段落丢掉
        # ——丢一个就是漏译一段，宁可让它们保持原有相对次序排在末尾。
        located = {id(item) for item in ordered}
        ordered.extend(item for item in paragraphs if id(item) not in located)
    return ordered


def _paragraph_content_runs(paragraph: Paragraph) -> list:
    """段落里所有承载可见文字的 run，按文档顺序，**包含 w:hyperlink 内部的 run**。

    python-docx 的 ``Paragraph.runs`` 只取 ``w:p`` 的直接 ``w:r`` 子元素，而
    ``Paragraph.text`` 是 ``w:r | w:hyperlink`` 拼出来的——读一套、写另一套。段落
    含超链接时改写译文就会漏掉超链接那截原文：译文写进第一个普通 run，超链接文字
    原封不动留在段落里，用户看到的是中英混排。段落文本的读与写必须共用这一个视图。
    """
    runs: list = []
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            runs.append(Run(child, paragraph))
        elif child.tag == qn("w:hyperlink"):
            runs.extend(
                Run(nested, paragraph)
                for nested in child.iterchildren(qn("w:r"))
            )
    return runs


def _paragraph_text_anchor_run(paragraph: Paragraph, runs: list):
    """挑出（必要时新建）承载整段文字的那个 run。

    段落以超链接开头时不能直接写进超链接内部的 run：那样整段译文都会变成可点击的
    链接。改为在该超链接之前插入一个段落级的普通 run。
    """
    first_run = runs[0]
    holder = first_run._r.getparent()
    if holder is paragraph._p:
        return first_run
    new_r = OxmlElement("w:r")
    holder.addprevious(new_r)
    return Run(new_r, paragraph)


def _drop_emptied_hyperlinks(paragraph: Paragraph) -> None:
    """删掉已经不含任何文字的超链接元素。

    改写整段文本后残留的空 ``w:hyperlink`` 是一个零长度的可点击区域：看不见、
    却仍然是链接。既然这一段的文字已经整体被译文取代，链接锚点也就无处附着了。
    """
    for hyperlink in list(paragraph._p.iterchildren(qn("w:hyperlink"))):
        if _element_has_visible_text(hyperlink):
            continue
        paragraph._p.remove(hyperlink)


def _paragraph_source_text(paragraph: Paragraph) -> str:
    return (paragraph.text or "").strip()


def _paragraph_raw_text(paragraph: Paragraph) -> str:
    return paragraph.text or ""


def _paragraph_leading_inline_whitespace(paragraph: Paragraph) -> str:
    match = _LEADING_INLINE_WHITESPACE_RE.match(_paragraph_raw_text(paragraph))
    return match.group(0) if match else ""


def _apply_leading_prefix(text: str, prefix: str) -> str:
    if not prefix:
        return text
    if str(text or "").startswith(prefix):
        return text
    return f"{prefix}{text}"


def _cell_source_text(cell: _Cell) -> str:
    paragraphs = [
        _paragraph_source_text(paragraph)
        for paragraph in _iter_cell_direct_paragraphs(cell)
    ]
    return "\n".join(text for text in paragraphs if text).strip()


def _is_translatable_source(
    source: str,
    *,
    target_lang: str,
    source_lang: str,
) -> bool:
    return bool(
        source
        and should_translate(
            source,
            target_lang=target_lang,
            source_lang=source_lang,
        )
    )


def _resolve_translation(
    source: str,
    translations: dict[str, str],
) -> ResolvedWordTranslation | None:
    raw = translations.get(source.strip())
    if raw is None:
        return None
    if is_replace_translation(raw):
        replacement = extract_replace_translation(raw).strip()
        return ResolvedWordTranslation(replacement, replace_only=True) if replacement else None

    translated = str(raw).strip()
    if not translated:
        return None
    if source.strip().casefold() == translated.casefold():
        return None
    return ResolvedWordTranslation(translated)


def _normalize_hex_fill(value: str, fallback: str = "FFF2CC") -> str:
    cleaned = str(value or "").strip().lstrip("#").upper()
    if len(cleaned) == 6 and all(char in "0123456789ABCDEF" for char in cleaned):
        return cleaned
    return fallback


def _normalize_review_mark_colors(colors: dict[str, str] | None) -> dict[str, str]:
    raw_colors = dict(colors or {})
    normalized: dict[str, str] = {}
    for mark, default_color in REVIEW_MARK_COLOR_DEFAULTS.items():
        normalized[mark] = _normalize_hex_fill(
            raw_colors.get(mark, ""),
            fallback=default_color,
        )
    return normalized


def _normalize_existing_highlight_policy(policy: str) -> str:
    value = str(policy or "").strip()
    if value in {
        EXISTING_HIGHLIGHT_POLICY_SKIP,
        EXISTING_HIGHLIGHT_POLICY_OVERWRITE,
        EXISTING_HIGHLIGHT_POLICY_RED_UNDERLINE,
    }:
        return value
    return EXISTING_HIGHLIGHT_POLICY_SKIP


def _normalize_review_marks(
    *,
    review_marks: dict[str, str] | None,
    review_highlight_sources: set[str] | None,
    review_highlight_color: str,
) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for source, mark in (review_marks or {}).items():
        cleaned_source = str(source or "").strip()
        if not cleaned_source:
            continue
        normalized[cleaned_source] = _normalize_review_mark(mark)

    if review_marks is None:
        legacy_fill = _normalize_hex_fill(review_highlight_color)
        for source in review_highlight_sources or set():
            cleaned_source = str(source or "").strip()
            if cleaned_source:
                normalized[cleaned_source] = legacy_fill
    return normalized


def _normalize_review_mark(mark: str) -> str:
    value = str(mark or "").strip()
    if value in {
        MIXED_MARK_UNRESOLVED,
        MIXED_MARK_FOREIGN_NOISE,
        MIXED_MARK_SEMANTIC,
    }:
        return value
    return _normalize_hex_fill(value)


def _review_mark_fill(mark: str, mark_colors: dict[str, str]) -> str:
    if mark in mark_colors:
        return mark_colors[mark]
    return _normalize_hex_fill(mark)


def _apply_paragraph_review_mark(
    paragraph: Paragraph,
    mark: str,
    existing_policy: str,
    mark_colors: dict[str, str],
) -> bool:
    if _paragraph_has_existing_highlight(paragraph):
        if existing_policy == EXISTING_HIGHLIGHT_POLICY_SKIP:
            return False
        if existing_policy == EXISTING_HIGHLIGHT_POLICY_RED_UNDERLINE:
            return _apply_paragraph_red_underline(paragraph)
    return _apply_paragraph_text_highlight(
        paragraph,
        _review_mark_highlight(mark, mark_colors),
    )


def _apply_cell_review_mark(
    cell: _Cell,
    mark: str,
    existing_policy: str,
    mark_colors: dict[str, str],
) -> bool:
    if _cell_has_existing_highlight(cell):
        if existing_policy == EXISTING_HIGHLIGHT_POLICY_SKIP:
            return False
        if existing_policy == EXISTING_HIGHLIGHT_POLICY_RED_UNDERLINE:
            return _apply_cell_red_underline(cell)
    applied = False
    highlight = _review_mark_highlight(mark, mark_colors)
    for paragraph in cell.paragraphs:
        applied = _apply_paragraph_text_highlight(paragraph, highlight) or applied
    return applied


def _review_mark_highlight(mark: str, mark_colors: dict[str, str]) -> str:
    fill = _review_mark_fill(mark, mark_colors)
    return _highlight_value_for_hex(fill)


def _highlight_value_for_hex(value: str) -> str:
    cleaned = _normalize_hex_fill(value)
    if cleaned in _WORD_HIGHLIGHT_HEX_OVERRIDES:
        return _WORD_HIGHLIGHT_HEX_OVERRIDES[cleaned]
    try:
        red, green, blue = tuple(int(cleaned[index : index + 2], 16) for index in (0, 2, 4))
    except Exception:
        return "yellow"
    high = max(red, green, blue)
    low = min(red, green, blue)
    if high <= 48:
        return "black"
    if high - low <= 18:
        return "lightGray" if high > 160 else "darkGray"

    if high == red:
        hue = ((green - blue) / (high - low)) % 6
    elif high == green:
        hue = ((blue - red) / (high - low)) + 2
    else:
        hue = ((red - green) / (high - low)) + 4
    hue *= 60

    if hue < 20 or hue >= 340:
        return "red"
    if hue < 70:
        return "yellow"
    if hue < 165:
        return "green"
    if hue < 205:
        return "cyan"
    if hue < 265:
        return "blue"
    if hue < 330:
        return "magenta"
    return "red"


def _apply_paragraph_text_highlight(paragraph: Paragraph, highlight: str) -> bool:
    applied = False
    for run in paragraph.runs:
        if not run.text:
            continue
        _set_highlight_value(run._element.get_or_add_rPr(), highlight)
        applied = True
    return applied


def _apply_paragraph_text_shading(paragraph: Paragraph, fill: str) -> bool:
    applied = False
    for run in paragraph.runs:
        if not run.text:
            continue
        _set_shading_fill(run._element.get_or_add_rPr(), fill)
        applied = True
    return applied


def _apply_paragraph_red_underline(paragraph: Paragraph) -> bool:
    applied = False
    for run in paragraph.runs:
        if not run.text:
            continue
        try:
            run.font.color.rgb = RGBColor(192, 0, 0)
            run.underline = True
            applied = True
        except Exception:
            continue
    return applied


def _apply_cell_red_underline(cell: _Cell) -> bool:
    applied = False
    for paragraph in cell.paragraphs:
        applied = _apply_paragraph_red_underline(paragraph) or applied
    return applied


def _paragraph_has_existing_highlight(paragraph: Paragraph) -> bool:
    p_pr = getattr(paragraph._p, "pPr", None)
    if _element_has_existing_shading(p_pr):
        return True
    try:
        if _element_has_existing_shading(getattr(paragraph.style._element, "pPr", None)):
            return True
    except Exception:
        pass

    for run in paragraph.runs:
        try:
            if run.font.highlight_color is not None:
                return True
        except Exception:
            pass

        r_pr = getattr(run._element, "rPr", None)
        if _element_has_existing_shading(r_pr):
            return True
        highlight = r_pr.find(qn("w:highlight")) if r_pr is not None else None
        if highlight is not None:
            value = str(highlight.get(qn("w:val")) or "").strip().upper()
            if value and value != "NONE":
                return True
    return False


def _cell_has_existing_highlight(cell: _Cell) -> bool:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if _element_has_existing_shading(tc_pr):
        return True
    return any(_paragraph_has_existing_highlight(paragraph) for paragraph in cell.paragraphs)


def _element_has_existing_shading(element) -> bool:
    if element is None:
        return False
    shd = element.find(qn("w:shd"))
    if shd is None:
        return False
    fill = str(shd.get(qn("w:fill")) or "").strip().upper()
    if fill and fill != "AUTO":
        return True
    value = str(shd.get(qn("w:val")) or "").strip().upper()
    return bool(value and value not in {"CLEAR", "NIL"})


def _set_shading_fill(parent_element, fill: str) -> None:
    shd = parent_element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        parent_element.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _normalize_hex_fill(fill))


def _set_highlight_value(parent_element, value: str) -> None:
    highlight = parent_element.find(qn("w:highlight"))
    if highlight is None:
        highlight = OxmlElement("w:highlight")
        parent_element.append(highlight)
    cleaned = str(value or "").strip()
    highlight.set(qn("w:val"), cleaned if cleaned in _WORD_HIGHLIGHT_RGB else "yellow")


def _is_toc_or_field_paragraph(paragraph: Paragraph) -> bool:
    """判断段落是否属于「不翻译的域段落」。

    只排除结果由 Word 自己生成或指向别处的域（目录、页码、交叉引用、超链接）——
    翻译它们没有意义，域一刷新就被覆盖。表题注、图题注用的是 SEQ 域，正文里也确实
    需要翻译，早期版本「见到任何域就跳过」会让这些题注一个字不翻、还不进报告。
    """
    style_name = _paragraph_style_name(paragraph).casefold()
    if "toc" in style_name or "目录" in style_name:
        return True

    instructions = _paragraph_field_instructions(paragraph)
    if not instructions:
        return False
    if _GENERATED_BLOCK_FIELD_RE.search(instructions):
        return True
    return not _has_translatable_literal_text(paragraph)


def _has_translatable_literal_text(paragraph: Paragraph) -> bool:
    """段落去掉域指令和域结果之后，是否还剩下值得翻译的文字。"""
    literal = _paragraph_literal_text(paragraph)
    cjk = 0
    latin = 0
    for char in literal:
        if not char.isalpha():
            continue
        if "一" <= char <= "鿿":
            cjk += 1
        else:
            latin += 1
    return cjk >= _FIELD_PARAGRAPH_MIN_CJK_CHARS or latin >= _FIELD_PARAGRAPH_MIN_LATIN_LETTERS


def _paragraph_literal_text(paragraph: Paragraph) -> str:
    """段落里作者自己敲的文字：域指令与域结果（含缓存的编号、页码）都不计入。"""
    element = paragraph._p
    field_simple_texts = {
        id(text_node)
        for field_simple in element.iter(qn("w:fldSimple"))
        for text_node in field_simple.iter(qn("w:t"))
    }
    parts: list[str] = []
    depth = 0
    for node in element.iter():
        tag = node.tag
        if tag == qn("w:fldChar"):
            char_type = node.get(qn("w:fldCharType"))
            if char_type == "begin":
                depth += 1
            elif char_type == "end":
                depth = max(0, depth - 1)
        elif tag == qn("w:t") and depth == 0 and id(node) not in field_simple_texts:
            parts.append(node.text or "")
    return "".join(parts)


def _paragraph_field_instructions(paragraph: Paragraph) -> str:
    """收集段落里所有域指令码，拼成一个大写串（w:instrText 常被拆到多个 run）。"""
    parts: list[str] = []
    for node in paragraph._p.iter():
        if node.tag == qn("w:fldSimple"):
            parts.append(str(node.get(qn("w:instr")) or ""))
        elif node.tag == qn("w:instrText"):
            parts.append(node.text or "")
    return " ".join(parts).upper()


def _paragraph_has_field(paragraph: Paragraph) -> bool:
    xml = paragraph._p.xml
    return "w:fldSimple" in xml or "w:fldChar" in xml


def _paragraph_style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def _is_heading_style(paragraph: Paragraph) -> bool:
    style_name = _paragraph_style_name(paragraph).casefold()
    return style_name.startswith("heading") or "标题" in style_name


def _detect_heading_level(paragraph: Paragraph) -> int | None:
    text = _paragraph_source_text(paragraph)
    if not text:
        return None

    style_name = _paragraph_style_name(paragraph)
    match = _HEADING_STYLE_RE.search(style_name)
    if match:
        raw_level = match.group(1) or match.group(2)
        try:
            return max(1, min(int(raw_level), 6))
        except (TypeError, ValueError):
            return 1

    if _CHINESE_CHAPTER_RE.match(text):
        return 1
    if _CHINESE_SECTION_RE.match(text):
        return 2
    if _NUMBERED_SECTION_RE.match(text):
        return min(text.split(maxsplit=1)[0].count(".") + 1, 6)
    if _BARE_NUMERIC_HEADING_RE.match(text):
        return 1

    return None


def _iter_body_block_items(doc: Document):
    """按 XML 文档流顺序交替产出顶层段落 / 表格及其各自的下标。

    python-docx 的 doc.paragraphs 和 doc.tables 各自独立编号，互相之间不保留
    "谁在文档里排在前面"的信息。要判断一张表格是否落在正文标题之前（从而应当被
    前置内容一并保护），必须按 body 元素的真实先后顺序重建这个交错序列。
    """
    body = doc.element.body
    paragraph_index = 0
    table_index = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", paragraph_index
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            yield "table", table_index
            table_index += 1


@dataclass(frozen=True)
class WordFrontMatterBoundary:
    """"保护封面和目录"选项划定出的正文起点。

    body_start_index 为 None 表示全文找不到任何可识别的正文标题——这种情况下调用方
    必须保持"不保护任何内容"，绝不能因为找不到边界就把整篇文档当成前置内容吞掉。
    """

    body_start_index: int | None
    heading_text: str = ""
    protected_paragraph_indices: frozenset[int] = frozenset()
    protected_table_indices: frozenset[int] = frozenset()

    @property
    def found(self) -> bool:
        return self.body_start_index is not None

    @property
    def protected_paragraph_count(self) -> int:
        return len(self.protected_paragraph_indices)


def find_word_front_matter_boundary(doc: Document) -> WordFrontMatterBoundary:
    """定位正文标题，划定其之前的封面 / 目录 / 前言范围。

    正文标题的判定直接复用 _detect_heading_level 的全部启发式（样式 + 三条正则 +
    无小数点的顶层数字标题），与 extract_word_segments 拆分章节路径时是同一套规则，
    避免出现两处"什么算标题"的定义分叉。扫描时必须先跳过：
    1. 自动生成的目录域段落（_is_toc_or_field_paragraph）；
    2. 手打目录条目（_HAND_TYPED_TOC_LEADER_RE）——否则目录里抄录的
       "第一章 工程概况 ...... 1" 会被误认成正文标题本身，导致只保护了封面、
       整份目录反而被当正文翻译。
    """
    paragraphs = list(doc.paragraphs)
    block_order = list(_iter_body_block_items(doc))

    body_start_block_index: int | None = None
    body_start_paragraph_index: int | None = None
    heading_text = ""
    for block_index, (kind, item_index) in enumerate(block_order):
        if kind != "paragraph":
            continue
        paragraph = paragraphs[item_index]
        text = _paragraph_source_text(paragraph)
        if not text:
            continue
        if _is_toc_or_field_paragraph(paragraph):
            continue
        if _HAND_TYPED_TOC_LEADER_RE.search(text):
            continue
        if _detect_heading_level(paragraph) is None:
            continue
        body_start_block_index = block_index
        body_start_paragraph_index = item_index
        heading_text = text
        break

    if body_start_block_index is None:
        return WordFrontMatterBoundary(body_start_index=None)

    preceding_blocks = block_order[:body_start_block_index]
    protected_paragraph_indices = frozenset(
        item_index
        for kind, item_index in preceding_blocks
        if kind == "paragraph" and _paragraph_source_text(paragraphs[item_index])
    )
    protected_table_indices = frozenset(
        item_index for kind, item_index in preceding_blocks if kind == "table"
    )
    return WordFrontMatterBoundary(
        body_start_index=body_start_paragraph_index,
        heading_text=heading_text,
        protected_paragraph_indices=protected_paragraph_indices,
        protected_table_indices=protected_table_indices,
    )


def find_word_front_matter_boundary_for_path(path: str | Path) -> WordFrontMatterBoundary:
    """便于只有源文件路径、尚未持有 Document 对象的调用方（如任务编排层）查询边界。"""
    return find_word_front_matter_boundary(Document(str(path)))


def _update_section_stack(section_stack: dict[int, str], level: int, text: str) -> None:
    cleaned = " ".join(str(text or "").split())
    if not cleaned:
        return
    for existing_level in list(section_stack):
        if existing_level >= level:
            section_stack.pop(existing_level, None)
    section_stack[level] = cleaned


def _format_section_path(section_stack: dict[int, str]) -> str:
    if not section_stack:
        return "正文"
    return " / ".join(section_stack[level] for level in sorted(section_stack))


def _collect_numbering_labels(
    doc: Document,
    paragraphs: list[Paragraph],
) -> dict[int, str]:
    level_definitions = _load_numbering_level_definitions(doc)
    counters: dict[tuple[str, int], int] = {}
    labels: dict[int, str] = {}

    for paragraph in paragraphs:
        numbering_info = _get_paragraph_numbering_info(paragraph)
        if numbering_info is None:
            continue
        if not _paragraph_source_text(paragraph):
            continue

        num_id, ilvl = numbering_info
        level_definition = level_definitions.get((num_id, ilvl))
        if level_definition is None:
            continue

        for key in list(counters):
            if key[0] == num_id and key[1] > ilvl:
                counters.pop(key, None)

        counter_key = (num_id, ilvl)
        if counter_key not in counters:
            counters[counter_key] = level_definition.start - 1
        counters[counter_key] += 1

        label = _format_numbering_label(
            num_id=num_id,
            ilvl=ilvl,
            counters=counters,
            level_definitions=level_definitions,
        )
        if label:
            labels[id(paragraph)] = label

    return labels


def _load_numbering_level_definitions(doc: Document) -> dict[tuple[str, int], NumberingLevelDefinition]:
    try:
        numbering_root = doc.part.numbering_part.element
    except Exception:
        return {}

    abstract_by_id = {
        abstract_num.get(qn("w:abstractNumId")): abstract_num
        for abstract_num in numbering_root.findall(qn("w:abstractNum"))
    }
    definitions: dict[tuple[str, int], NumberingLevelDefinition] = {}

    for num in numbering_root.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abstract_id = _child_val(num, "w:abstractNumId")
        abstract_num = abstract_by_id.get(abstract_id)
        if not num_id or abstract_num is None:
            continue

        levels = {
            int(level.get(qn("w:ilvl")) or 0): _read_numbering_level_definition(level)
            for level in abstract_num.findall(qn("w:lvl"))
        }

        for override in num.findall(qn("w:lvlOverride")):
            ilvl = _to_int(override.get(qn("w:ilvl")), fallback=0)
            if override.find(qn("w:lvl")) is not None:
                levels[ilvl] = _read_numbering_level_definition(override.find(qn("w:lvl")))
            elif override.find(qn("w:startOverride")) is not None:
                existing = levels.get(ilvl, NumberingLevelDefinition())
                levels[ilvl] = NumberingLevelDefinition(
                    start=_to_int(
                        _child_val(override, "w:startOverride"),
                        fallback=existing.start,
                    ),
                    number_format=existing.number_format,
                    level_text=existing.level_text,
                    legal_numbering=existing.legal_numbering,
                )

        for ilvl, definition in levels.items():
            definitions[(num_id, ilvl)] = definition

    return definitions


def _read_numbering_level_definition(level) -> NumberingLevelDefinition:
    return NumberingLevelDefinition(
        start=_to_int(_child_val(level, "w:start"), fallback=1),
        number_format=_child_val(level, "w:numFmt", default="decimal"),
        level_text=_child_val(level, "w:lvlText", default="%1."),
        legal_numbering=_on_off_value(level.find(qn("w:isLgl"))) is True,
    )


def _child_val(element, child_tag: str, default: str = "") -> str:
    if element is None:
        return default
    child = element.find(qn(child_tag))
    if child is None:
        return default
    return str(child.get(qn("w:val")) or default)


def _to_int(value, *, fallback: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return fallback


def _get_paragraph_numbering_info(paragraph: Paragraph) -> tuple[str, int] | None:
    direct_p_pr = getattr(paragraph._p, "pPr", None)
    direct = _read_numbering_info_from_ppr(direct_p_pr)
    if direct is not None:
        return direct
    if _ppr_suppresses_numbering(direct_p_pr):
        return None

    try:
        style_element = paragraph.style._element
    except Exception:
        return None
    return _read_numbering_info_from_ppr(getattr(style_element, "pPr", None))


def _ppr_suppresses_numbering(p_pr) -> bool:
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    return _child_val(num_pr, "w:numId") == "0"


def _read_numbering_info_from_ppr(p_pr) -> tuple[str, int] | None:
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None

    num_id = _child_val(num_pr, "w:numId")
    if not num_id or num_id == "0":
        return None
    return num_id, _to_int(_child_val(num_pr, "w:ilvl"), fallback=0)


def _format_numbering_label(
    *,
    num_id: str,
    ilvl: int,
    counters: dict[tuple[str, int], int],
    level_definitions: dict[tuple[str, int], NumberingLevelDefinition],
) -> str:
    definition = level_definitions.get((num_id, ilvl))
    if definition is None:
        return ""

    label = definition.level_text or "%1."
    if definition.number_format == "bullet":
        return _normalize_bullet_label(label)

    def replace_placeholder(match: re.Match[str]) -> str:
        level = max(0, _to_int(match.group(1), fallback=1) - 1)
        level_counter = counters.get((num_id, level), 0)
        if definition.legal_numbering:
            # w:isLgl：本层引用到的所有层级一律用阿拉伯数字，忽略被引用层自己的 numFmt。
            return _format_number(level_counter, "decimal")
        level_definition = level_definitions.get((num_id, level), definition)
        number_format = _number_format_for_level_text_placeholder(
            label,
            match,
            level_definition.number_format,
        )
        return _format_number(level_counter, number_format)

    formatted = re.sub(r"%(\d+)", replace_placeholder, label).strip()
    return _normalize_chinese_ordinal_numbering_label(formatted, top_level=ilvl == 0)


def _normalize_chinese_ordinal_numbering_label(label: str, *, top_level: bool = False) -> str:
    def replace_match(match: re.Match[str]) -> str:
        unit = match.group(2)
        if top_level and unit == "节":
            unit = "章"
        return f"第{match.group(1)}{unit}"

    return re.sub(
        rf"^第\s*([{_CHINESE_NUMBER_CHARS}]+)\s*([章节篇卷部])",
        replace_match,
        label,
    )


def _number_format_for_level_text_placeholder(
    level_text: str,
    match: re.Match[str],
    number_format: str,
) -> str:
    if number_format != "decimal":
        return number_format

    before = level_text[: match.start()]
    after = level_text[match.end() :]
    if before.rstrip().endswith("第") and after.lstrip().startswith(("章", "节", "篇", "卷", "部")):
        return "chineseCounting"
    return number_format


def _format_number(number: int, number_format: str) -> str:
    if number <= 0:
        number = 1
    if number_format in _CHINESE_NUMBER_FORMATS:
        return _format_chinese_number(number)
    if number_format == "decimalFullWidth":
        return _format_full_width_number(number)
    if number_format == "lowerLetter":
        return _format_alpha_number(number).lower()
    if number_format == "upperLetter":
        return _format_alpha_number(number).upper()
    if number_format == "lowerRoman":
        return _format_roman_number(number).lower()
    if number_format == "upperRoman":
        return _format_roman_number(number).upper()
    return str(number)


def _format_full_width_number(number: int) -> str:
    return str(number).translate(str.maketrans("0123456789", "０１２３４５６７８９"))


def _format_chinese_number(number: int) -> str:
    if number <= 0:
        return _CHINESE_DIGITS[0]
    if number < 10:
        return _CHINESE_DIGITS[number]
    if number < 100:
        tens, ones = divmod(number, 10)
        prefix = "" if tens == 1 else _CHINESE_DIGITS[tens]
        return f"{prefix}十{_CHINESE_DIGITS[ones] if ones else ''}"
    if number < 1000:
        hundreds, remainder = divmod(number, 100)
        if remainder == 0:
            return f"{_CHINESE_DIGITS[hundreds]}百"
        connector = "" if remainder >= 10 else "零"
        return f"{_CHINESE_DIGITS[hundreds]}百{connector}{_format_chinese_number(remainder)}"
    if number < 10000:
        thousands, remainder = divmod(number, 1000)
        if remainder == 0:
            return f"{_CHINESE_DIGITS[thousands]}千"
        connector = "" if remainder >= 100 else "零"
        return f"{_CHINESE_DIGITS[thousands]}千{connector}{_format_chinese_number(remainder)}"
    if number < 100000000:
        high, remainder = divmod(number, 10000)
        if remainder == 0:
            return f"{_format_chinese_number(high)}万"
        connector = "" if remainder >= 1000 else "零"
        return f"{_format_chinese_number(high)}万{connector}{_format_chinese_number(remainder)}"
    return str(number)


def _format_alpha_number(number: int) -> str:
    chars: list[str] = []
    while number > 0:
        number -= 1
        chars.append(chr(ord("A") + (number % 26)))
        number //= 26
    return "".join(reversed(chars)) or "A"


def _format_roman_number(number: int) -> str:
    values = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    result: list[str] = []
    for value, token in values:
        while number >= value:
            result.append(token)
            number -= value
    return "".join(result) or "I"


def _normalize_bullet_label(label: str) -> str:
    cleaned = (label or "").strip()
    if cleaned in {"\uf0b7", ""}:
        return "•"
    return cleaned or "•"


def _flatten_automatic_numbering(
    paragraphs: list[Paragraph],
    numbering_labels: dict[int, str],
    *,
    protected_paragraph_ids: set[int] | frozenset[int] | None = None,
) -> WordNumberingNormalizationStats:
    protected_ids = protected_paragraph_ids or frozenset()
    labels_seen = 0
    labels_prepended = 0
    numbering_removed = 0
    for paragraph in paragraphs:
        if id(paragraph) in protected_ids:
            continue
        label = numbering_labels.get(id(paragraph), "")
        if not label:
            if not _paragraph_source_text(paragraph) and _get_paragraph_numbering_info(paragraph):
                if _remove_paragraph_numbering(paragraph):
                    numbering_removed += 1
            continue
        labels_seen += 1
        if _prepend_paragraph_text(paragraph, label):
            labels_prepended += 1
        if _remove_paragraph_numbering(paragraph):
            numbering_removed += 1
    return WordNumberingNormalizationStats(
        labels_seen=labels_seen,
        labels_prepended=labels_prepended,
        numbering_removed=numbering_removed,
    )


def _prepend_paragraph_text(paragraph: Paragraph, label: str) -> bool:
    source = _paragraph_source_text(paragraph)
    if not source or _has_visible_numbering_prefix(source) or _text_starts_with_label(source, label):
        return False

    prefix = f"{label} "
    runs = _paragraph_content_runs(paragraph)
    if not runs:
        paragraph.add_run(prefix + source)
        return True
    first_run = runs[0]
    if first_run._r.getparent() is paragraph._p:
        first_run.text = prefix + first_run.text.lstrip()
        return True
    # 段落以超链接开头：编号标签写进 paragraph.runs[0] 会落到超链接之后，序号掉进
    # 段落中间；直接 add_run 又会追加到末尾。必须在超链接之前插入一个新的普通 run。
    label_run = _paragraph_text_anchor_run(paragraph, runs)
    label_run.text = prefix
    return True


def _remove_paragraph_numbering(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)
    suppress_num_pr = OxmlElement("w:numPr")
    suppress_num_id = OxmlElement("w:numId")
    suppress_num_id.set(qn("w:val"), "0")
    suppress_num_pr.append(suppress_num_id)
    p_pr.append(suppress_num_pr)
    return True


def _text_starts_with_label(text: str, label: str) -> bool:
    cleaned_text = str(text or "").strip()
    cleaned_label = str(label or "").strip()
    if not cleaned_text or not cleaned_label:
        return False
    return _text_starts_with_exact_label(cleaned_text, cleaned_label)


def _has_visible_numbering_prefix(text: str) -> bool:
    cleaned = str(text or "").strip()
    if not cleaned:
        return False
    if _BULLET_PREFIX_RE.match(cleaned):
        return True
    if _CHINESE_NUMBERING_PREFIX_RE.match(cleaned):
        return True
    if _CHINESE_LIST_PREFIX_RE.match(cleaned):
        return True
    if _ARABIC_LIST_PREFIX_RE.match(cleaned):
        return True
    if _ALPHA_LIST_PREFIX_RE.match(cleaned):
        return True

    decimal_match = _ARABIC_DECIMAL_PREFIX_RE.match(cleaned)
    if decimal_match is None:
        return False
    remainder = decimal_match.group(2)
    if not remainder:
        return True
    next_char = remainder[0]
    if next_char.isascii() and (next_char.isalpha() or next_char in {"%", "#"}):
        return False
    return True


def _text_starts_with_exact_label(text: str, label: str) -> bool:
    if not text.startswith(label):
        return False
    remainder = text[len(label) :]
    if not remainder:
        return True
    if label.startswith("第") and re.search(r"[章节篇]$", label):
        return True
    return _is_numbering_boundary(remainder[0])


def _is_numbering_boundary(char: str) -> bool:
    return char.isspace() or char in {
        ".",
        "．",
        "、",
        "，",
        ",",
        ":",
        "：",
        ")",
        "）",
        "-",
        "－",
        "—",
        "·",
    }


def _insert_translation_paragraph_after(
    paragraph: Paragraph,
    text: str,
    *,
    target_lang: str,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    _copy_translation_paragraph_shape(paragraph, new_para, target_lang=target_lang)
    _remove_paragraph_numbering(new_para)
    run = new_para.add_run(text)
    _copy_run_shape(paragraph, run, target_lang=target_lang)
    return new_para


def _append_translation_to_cell(
    cell: _Cell,
    text: str,
    *,
    target_lang: str,
) -> Paragraph:
    source_para = _last_non_empty_paragraph(cell.paragraphs)
    new_para = cell.add_paragraph()
    if source_para is not None:
        _copy_translation_paragraph_shape(source_para, new_para, target_lang=target_lang)
        _remove_paragraph_numbering(new_para)
    run = new_para.add_run(text)
    if source_para is not None:
        _copy_run_shape(source_para, run, target_lang=target_lang)
    else:
        _set_latin_run_font(run, target_lang=target_lang)
    return new_para


def _replace_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    target_lang: str,
) -> None:
    runs = _paragraph_content_runs(paragraph)
    if not runs:
        run = paragraph.add_run(text)
        _set_latin_run_font(run, target_lang=target_lang)
        return

    anchor = _paragraph_text_anchor_run(paragraph, runs)
    anchor.text = text
    _set_latin_run_font(anchor, target_lang=target_lang)
    for run in runs:
        if run._r is anchor._r:
            continue
        run.text = ""
    _drop_emptied_hyperlinks(paragraph)


def _clear_paragraph_text(paragraph: Paragraph) -> None:
    for run in _paragraph_content_runs(paragraph):
        run.text = ""
    _drop_emptied_hyperlinks(paragraph)


def _replace_cell_text(cell: _Cell, text: str, *, target_lang: str) -> None:
    """逐段改写单元格文本，保留嵌套表格、多段落与段落内的局部格式。

    ``cell.text = ...`` 会把整个 ``w:tc`` 清空、只留下一个纯文本段落：单元格里的
    嵌套表格、其余段落、局部加粗全部被压平，而这只是"把原文换成译文"而已。单元格
    原文本身就是各直接子段落的文本用换行拼起来的（见 ``_cell_source_text``），所以
    译文也按换行拆回去、一段对一段地写。
    """
    paragraphs = [
        paragraph
        for paragraph in _iter_cell_direct_paragraphs(cell)
        if _paragraph_source_text(paragraph)
    ]
    if not paragraphs:
        target = next(iter(_iter_cell_direct_paragraphs(cell)), None) or cell.add_paragraph()
        _replace_paragraph_text(target, text, target_lang=target_lang)
        return

    lines = str(text or "").split("\n")
    if len(lines) == len(paragraphs):
        for paragraph, line in zip(paragraphs, lines):
            _replace_paragraph_text(paragraph, line, target_lang=target_lang)
        return

    # 行数对不上（模型合并或拆分了换行）时不再猜哪一行对应哪一段：整段译文写进第一
    # 段，其余原文段清空。留下空段落好过让原文残留在译文旁边。
    _replace_paragraph_text(paragraphs[0], text, target_lang=target_lang)
    for paragraph in paragraphs[1:]:
        _clear_paragraph_text(paragraph)


def _trim_trailing_empty_body_paragraphs(doc) -> None:
    body = doc._body._element
    children = list(body)
    last_content_index = len(children) - 1
    if last_content_index >= 0 and children[last_content_index].tag == qn("w:sectPr"):
        last_content_index -= 1

    while last_content_index >= 0:
        element = children[last_content_index]
        if element.tag != qn("w:p") or not _is_removable_empty_paragraph_element(element):
            break
        body.remove(element)
        children.pop(last_content_index)
        last_content_index -= 1


def _is_removable_empty_paragraph_element(paragraph_element) -> bool:
    for tag in (
        "t",
        "tab",
        "br",
        "drawing",
        "pict",
        "object",
        "fldChar",
        "instrText",
        "sectPr",
        "bookmarkStart",
        "bookmarkEnd",
        "commentRangeStart",
        "commentRangeEnd",
    ):
        if paragraph_element.findall(f".//{qn(f'w:{tag}')}"):
            return False
    return True


def _last_non_empty_paragraph(paragraphs: list[Paragraph]) -> Paragraph | None:
    for paragraph in reversed(paragraphs):
        if paragraph.text.strip():
            return paragraph
    return paragraphs[-1] if paragraphs else None


def _copy_translation_paragraph_shape(
    source: Paragraph,
    target: Paragraph,
    *,
    target_lang: str,
) -> None:
    target_p_pr = getattr(target._p, "pPr", None)
    if target_p_pr is not None:
        target._p.remove(target_p_pr)

    copied_p_pr = _build_translation_paragraph_properties(source)
    if _is_heading_style(source):
        p_style = copied_p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            copied_p_pr.remove(p_style)
    _strip_translation_flow_controls(source, copied_p_pr)
    if target_lang != "zh":
        _materialize_character_first_line_indent(source, copied_p_pr)
    target._p.insert(0, copied_p_pr)


def _build_translation_paragraph_properties(source: Paragraph):
    p_pr = OxmlElement("w:pPr")
    for style in _paragraph_style_chain(source):
        style_p_pr = getattr(style.element, "pPr", None)
        if style_p_pr is not None:
            _merge_paragraph_properties(p_pr, style_p_pr)

    source_p_pr = getattr(source._p, "pPr", None)
    if source_p_pr is not None:
        _merge_paragraph_properties(p_pr, source_p_pr)
    return p_pr


def _paragraph_style_chain(paragraph: Paragraph) -> list:
    try:
        style = paragraph.style
    except Exception:
        return []
    chain = []
    seen: set[str] = set()
    while style is not None:
        style_id = str(getattr(style, "style_id", id(style)))
        if style_id in seen:
            break
        seen.add(style_id)
        chain.append(style)
        style = getattr(style, "base_style", None)
    chain.reverse()
    return chain


def _merge_paragraph_properties(target_p_pr, source_p_pr) -> None:
    for child in list(source_p_pr):
        if child.tag == qn("w:ind"):
            _merge_indentation(target_p_pr, child)
            continue
        existing = target_p_pr.find(child.tag)
        if existing is not None:
            target_p_pr.remove(existing)
        target_p_pr.append(deepcopy(child))


def _merge_indentation(target_p_pr, source_ind) -> None:
    target_ind = target_p_pr.find(qn("w:ind"))
    if target_ind is None:
        target_ind = OxmlElement("w:ind")
        target_p_pr.append(target_ind)

    source_attrs = dict(source_ind.attrib)
    if source_attrs.get(qn("w:firstLineChars")) == "0":
        _remove_indent_attrs(
            target_ind,
            ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"),
        )
    elif _indent_has_hanging(source_attrs):
        _remove_indent_attrs(target_ind, ("w:firstLine", "w:firstLineChars"))
    elif _indent_has_authoritative_first_line(source_attrs):
        _remove_indent_attrs(target_ind, ("w:hanging", "w:hangingChars"))
        if source_attrs.get(qn("w:firstLine")) not in {None, "0"}:
            _remove_indent_attrs(target_ind, ("w:firstLineChars",))

    for key, value in source_attrs.items():
        target_ind.set(key, value)


def _indent_has_hanging(attrs: dict) -> bool:
    return any(attrs.get(qn(attr)) is not None for attr in ("w:hanging", "w:hangingChars"))


def _indent_has_authoritative_first_line(attrs: dict) -> bool:
    return any(attrs.get(qn(attr)) is not None for attr in ("w:firstLine", "w:firstLineChars"))


def _remove_indent_attrs(ind, attrs: tuple[str, ...]) -> None:
    for attr in attrs:
        key = qn(attr)
        if key in ind.attrib:
            del ind.attrib[key]


def _strip_translation_flow_controls(source: Paragraph, p_pr) -> None:
    _remove_ppr_children(p_pr, "sectPr")
    _remove_ppr_children(p_pr, "pageBreakBefore")

    if _style_page_break_before(source) and p_pr.find(qn("w:pStyle")) is not None:
        page_break_before = OxmlElement("w:pageBreakBefore")
        page_break_before.set(qn("w:val"), "0")
        p_pr.append(page_break_before)


def _remove_ppr_children(p_pr, tag: str) -> None:
    for child in list(p_pr.findall(qn(f"w:{tag}"))):
        p_pr.remove(child)


def _copy_run_shape(source_paragraph: Paragraph, target_run, *, target_lang: str) -> None:
    source_runs = [run for run in source_paragraph.runs if run.text.strip()]
    if source_runs:
        try:
            target_run.bold = _first_defined(
                _uniform_run_value(source_runs, "bold"),
                _paragraph_default_bool(source_paragraph, "b"),
                _style_font_value(source_paragraph, "bold"),
            )
            target_run.italic = _first_defined(
                _uniform_run_value(source_runs, "italic"),
                _paragraph_default_bool(source_paragraph, "i"),
                _style_font_value(source_paragraph, "italic"),
            )
            target_run.underline = _first_defined(
                _uniform_run_value(source_runs, "underline"),
                _paragraph_default_underline(source_paragraph),
                _style_font_value(source_paragraph, "underline"),
            )
            target_run.font.size = _first_defined(
                _uniform_run_font_value(source_runs, "size"),
                _paragraph_default_size(source_paragraph),
                _style_font_value(source_paragraph, "size"),
            )
            target_run.font.color.rgb = _first_defined(
                _uniform_run_font_color(source_runs),
                _paragraph_default_color(source_paragraph),
                _style_font_color(source_paragraph),
            )
        except Exception:
            pass
    _set_latin_run_font(target_run, target_lang=target_lang)


def _materialize_character_first_line_indent(source: Paragraph, p_pr) -> None:
    raw_chars = _direct_first_line_chars(p_pr)
    if raw_chars is None and _direct_ind_keeps_style_first_line(p_pr):
        raw_chars = _style_first_line_chars(source)
    if raw_chars is None:
        return
    try:
        first_line_chars = int(raw_chars or "0")
    except ValueError:
        _clear_character_indent_attrs(p_pr)
        return
    if first_line_chars <= 0:
        _clear_character_indent_attrs(p_pr)
        return

    half_points = _paragraph_effective_half_points(source)
    twips = max(0, round((first_line_chars / 100) * (half_points / 2) * 20))
    if twips <= 0:
        _clear_character_indent_attrs(p_pr)
        return

    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    try:
        current = int(ind.get(qn("w:firstLine")) or "0")
    except ValueError:
        current = 0
    if current < twips:
        ind.set(qn("w:firstLine"), str(twips))
    _clear_character_indent_attrs(p_pr)


def _clear_character_indent_attrs(p_pr) -> None:
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        return
    for attr in ("w:firstLineChars", "w:hangingChars"):
        key = qn(attr)
        if key in ind.attrib:
            del ind.attrib[key]


def _direct_first_line_chars(p_pr) -> str | None:
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        return None
    return ind.get(qn("w:firstLineChars"))


def _direct_ind_keeps_style_first_line(p_pr) -> bool:
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        return True
    for attr in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
        if ind.get(qn(attr)) is not None:
            return False
    return True


def _style_first_line_chars(paragraph: Paragraph) -> str | None:
    try:
        style = paragraph.style
    except Exception:
        return None
    while style is not None:
        p_pr = style.element.pPr
        if p_pr is not None:
            value = _direct_first_line_chars(p_pr)
            if value is not None:
                return value
        style = getattr(style, "base_style", None)
    return None


def _style_page_break_before(paragraph: Paragraph) -> bool:
    try:
        style = paragraph.style
    except Exception:
        return False
    while style is not None:
        p_pr = style.element.pPr
        if p_pr is not None:
            value = _on_off_value(p_pr.find(qn("w:pageBreakBefore")))
            if value is not None:
                return value
        style = getattr(style, "base_style", None)
    return False


def _paragraph_effective_half_points(paragraph: Paragraph) -> int:
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        value = _run_half_points(run)
        if value is not None:
            return value
    value = _paragraph_default_half_points(paragraph)
    if value is not None:
        return value
    value = _style_half_points(paragraph)
    if value is not None:
        return value
    return 24


def _run_half_points(run) -> int | None:
    try:
        size = run.font.size
        if size is not None:
            return round(size.pt * 2)
    except Exception:
        pass
    return _rpr_half_points(getattr(run._element, "rPr", None))


def _paragraph_default_half_points(paragraph: Paragraph) -> int | None:
    return _rpr_half_points(_paragraph_rpr(paragraph))


def _style_half_points(paragraph: Paragraph) -> int | None:
    try:
        style = paragraph.style
    except Exception:
        return None
    while style is not None:
        value = _rpr_half_points(style.element.rPr)
        if value is not None:
            return value
        style = getattr(style, "base_style", None)
    return None


def _rpr_half_points(r_pr) -> int | None:
    if r_pr is None:
        return None
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        return None
    try:
        return int(sz.get(qn("w:val")) or "0")
    except ValueError:
        return None


def _uniform_run_value(runs, attr: str):
    values = [getattr(run, attr) for run in runs]
    if len(set(values)) == 1:
        return values[0]
    return None


def _uniform_run_font_value(runs, attr: str):
    values = [getattr(run.font, attr) for run in runs]
    if len(set(values)) == 1:
        return values[0]
    return None


def _uniform_run_font_color(runs):
    values = [run.font.color.rgb for run in runs]
    if len(set(values)) == 1:
        return values[0]
    return None


def _paragraph_rpr(paragraph: Paragraph):
    p_pr = getattr(paragraph._p, "pPr", None)
    if p_pr is None:
        return None
    return p_pr.find(qn("w:rPr"))


def _paragraph_default_bool(paragraph: Paragraph, tag: str) -> bool | None:
    r_pr = _paragraph_rpr(paragraph)
    if r_pr is None:
        return None
    return _on_off_value(r_pr.find(qn(f"w:{tag}")))


def _paragraph_default_underline(paragraph: Paragraph) -> bool | None:
    r_pr = _paragraph_rpr(paragraph)
    if r_pr is None:
        return None
    underline = r_pr.find(qn("w:u"))
    if underline is None:
        return None
    return str(underline.get(qn("w:val")) or "single").lower() != "none"


def _paragraph_default_size(paragraph: Paragraph):
    half_points = _paragraph_default_half_points(paragraph)
    if half_points is None:
        return None
    return half_points * 6350


def _paragraph_default_color(paragraph: Paragraph):
    r_pr = _paragraph_rpr(paragraph)
    if r_pr is None:
        return None
    color = r_pr.find(qn("w:color"))
    if color is None:
        return None
    value = str(color.get(qn("w:val")) or "").strip()
    if len(value) != 6 or value.lower() == "auto":
        return None
    try:
        return RGBColor.from_string(value)
    except ValueError:
        return None


def _on_off_value(element) -> bool | None:
    if element is None:
        return None
    value = str(element.get(qn("w:val")) or "1").lower()
    return value not in {"0", "false", "off", "none"}


def _first_defined(*values):
    for value in values:
        if value is not None:
            return value
    return None


def _style_font_value(paragraph: Paragraph, attr: str):
    try:
        return getattr(paragraph.style.font, attr)
    except Exception:
        return None


def _style_font_color(paragraph: Paragraph):
    try:
        return paragraph.style.font.color.rgb
    except Exception:
        return None


def _set_latin_run_font(run, *, target_lang: str) -> None:
    if target_lang == "zh":
        return
    try:
        run.font.name = "Times New Roman"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            r_fonts.set(qn(attr), "Times New Roman")
    except Exception:
        pass


def _ensure_owner_writable(path: Path) -> None:
    current_mode = path.stat().st_mode
    if current_mode & stat.S_IWUSR:
        return
    path.chmod(current_mode | stat.S_IWUSR)


def _sanitize_filename_fragment(value: str) -> str:
    cleaned = _INVALID_FILENAME_FRAGMENT_RE.sub("_", str(value or "")).strip().rstrip(". ")
    return cleaned or "目标语言"
