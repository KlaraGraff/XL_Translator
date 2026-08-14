# -*- coding: utf-8 -*-
"""
双语文档回放体检（只验不译，0 API）。

设计背景见 docs/redesign/2026-08-14_residual_repair_pipeline.md §六.3：
对已产出的双语 docx/xlsx 复跑残留分类器 + 文档级标题一致性巡检，输出
台账化的体检结果。用途有二：

1. 回放护栏——对历史交付语料回放，任何流水线改动使干净段落出现新动作
   即视为回归（tests/ 里有合成语料的机器化版本，真实语料在本机跑）；
2. 体检工具——对任意存量双语文档做残留巡检，不重译、不改文件。

配对启发式（docx）：正文按「中文占比 >0.3 的段落后紧跟占比 <0.3 的段落」
配成源/译对，与本项目双语输出「译文段插在源文段之后」的排版一致；表格
单元格与 xlsx 单元格用共享的 split_existing_bilingual_text 拆分同格双语。
该启发式在 9 份历史交付文档（1686 配对段）上实测 0 误配。
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from core.residual_classifier import (
    CATEGORY_NUMBERING_PREFIX,
    CATEGORY_QUANTITY_UNIT,
    check_heading_consistency,
    classify_residual_spans,
    detect_sibling_convention,
    deterministic_numbering_fix,
    is_section_heading_source,
)
from core.unit_ledger import (
    STATE_CLEAN,
    STATE_NEEDS_REVIEW,
    STATE_RELEASED_WITH_NOTE,
    UnitLedger,
)

_CJK_CHAR_RE = re.compile(r"[一-龥]")


@dataclass(frozen=True)
class ReplayFinding:
    """一处残留观测（source_anchor 为主锚点，与台账同基准）。"""

    unit_id: str
    source_anchor: str
    output_anchor: str
    source_text: str
    target_text: str
    category: str
    span_text: str
    deterministic_fix: str | None = None


@dataclass(frozen=True)
class HeadingOutlierFinding:
    unit_id: str
    target_text: str
    majority_form: str
    fix: str | None = None


@dataclass
class ReplayResult:
    path: str
    kind: str  # docx / xlsx
    target_lang: str
    convention: str = ""
    pair_count: int = 0
    findings: list[ReplayFinding] = field(default_factory=list)
    heading_outliers: list[HeadingOutlierFinding] = field(default_factory=list)
    ledger: UnitLedger = field(default_factory=UnitLedger)

    @property
    def clean(self) -> bool:
        return not self.findings and not self.heading_outliers


def _cjk_ratio(text: str) -> float:
    condensed = re.sub(r"\s+", "", str(text or ""))
    if not condensed:
        return 0.0
    hits = sum(1 for char in condensed if _CJK_CHAR_RE.match(char))
    return hits / len(condensed)


@dataclass(frozen=True)
class _Pair:
    source_anchor: str
    output_anchor: str
    source_text: str
    target_text: str


def _pair_adjacent_texts(entries, anchor_prefix: str) -> list[_Pair]:
    """entries: [(索引, 文本)]，索引为原文档中的位置（1 起）。"""
    pairs: list[_Pair] = []
    k = 0
    while k < len(entries) - 1:
        i1, t1 = entries[k]
        i2, t2 = entries[k + 1]
        if _cjk_ratio(t1) > 0.3 and _cjk_ratio(t2) < 0.3:
            pairs.append(
                _Pair(
                    source_anchor=f"{anchor_prefix}[{i1}]",
                    output_anchor=f"{anchor_prefix}[{i2}]",
                    source_text=t1,
                    target_text=t2,
                )
            )
            k += 2
        else:
            k += 1
    return pairs


def _pair_docx(path: Path, *, source_lang: str, target_lang: str) -> list[_Pair]:
    from docx import Document  # 延迟导入：xlsx 回放不应依赖 python-docx

    from core.translation_coverage import split_existing_bilingual_text

    doc = Document(str(path))
    body_entries = [
        (index + 1, paragraph.text.strip())
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip()
    ]
    pairs = _pair_adjacent_texts(body_entries, "body.paragraph")

    cell_index = 0
    for table_index, table in enumerate(doc.tables):
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen_cells:  # 合并单元格去重
                    continue
                seen_cells.add(id(cell._tc))
                anchor = f"table[{table_index}].cell[{cell_index}]"
                cell_index += 1
                text = cell.text.strip()
                if not text:
                    continue
                split = split_existing_bilingual_text(
                    text, source_lang=source_lang, target_lang=target_lang
                )
                if split is None:
                    continue
                source_text, target_text = split
                pairs.append(
                    _Pair(
                        source_anchor=anchor,
                        output_anchor=anchor,
                        source_text=source_text,
                        target_text=target_text,
                    )
                )
    return pairs


def _pair_xlsx(path: Path, *, source_lang: str, target_lang: str) -> list[_Pair]:
    import openpyxl

    from core.translation_coverage import split_existing_bilingual_text

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    pairs: list[_Pair] = []
    try:
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    value = cell.value
                    if not isinstance(value, str) or not value.strip():
                        continue
                    split = split_existing_bilingual_text(
                        value, source_lang=source_lang, target_lang=target_lang
                    )
                    if split is None:
                        continue
                    source_text, target_text = split
                    anchor = f"{sheet.title}!{cell.coordinate}"
                    pairs.append(
                        _Pair(
                            source_anchor=anchor,
                            output_anchor=anchor,
                            source_text=source_text,
                            target_text=target_text,
                        )
                    )
    finally:
        workbook.close()
    return pairs


def replay_pairs(
    pairs: list[_Pair],
    *,
    target_lang: str,
    result: ReplayResult,
) -> ReplayResult:
    """对配对结果跑分类器 + 标题一致性，填充 result 与台账。"""
    convention = detect_sibling_convention(
        [(pair.source_text, pair.target_text) for pair in pairs]
    )
    result.convention = convention
    result.pair_count = len(pairs)

    for pair in pairs:
        record = result.ledger.register(
            pair.source_anchor,
            source_text=pair.source_text,
            source_anchor=pair.source_anchor,
            output_anchor=pair.output_anchor,
        )
        spans = classify_residual_spans(pair.target_text, target_lang=target_lang)
        if not spans:
            result.ledger.set_state(record.unit_id, STATE_CLEAN)
            continue
        categories = sorted({span.category for span in spans})
        for span in spans:
            fix = None
            if span.category == CATEGORY_NUMBERING_PREFIX:
                fix = deterministic_numbering_fix(
                    pair.target_text,
                    convention=convention,
                    target_lang=target_lang,
                    source_text=pair.source_text,
                )
            result.findings.append(
                ReplayFinding(
                    unit_id=record.unit_id,
                    source_anchor=pair.source_anchor,
                    output_anchor=pair.output_anchor,
                    source_text=pair.source_text,
                    target_text=pair.target_text,
                    category=span.category,
                    span_text=span.text,
                    deterministic_fix=fix,
                )
            )
        if categories == [CATEGORY_QUANTITY_UNIT]:
            state = STATE_RELEASED_WITH_NOTE
        else:
            # 回放模式只诊断不修复：凡有非放行级残留一律标 needs_review，
            # 「本可自动修复」体现在 finding 的 deterministic_fix 字段里。
            state = STATE_NEEDS_REVIEW
        result.ledger.set_state(
            record.unit_id,
            state,
            categories=categories,
            evidence="residual: " + "、".join(f"{s.category}«{s.text}»" for s in spans),
        )

    heading_observations = [
        (pair.source_text, pair.target_text, pair.source_anchor)
        for pair in pairs
        if is_section_heading_source(pair.source_text)
    ]
    heading = check_heading_consistency(heading_observations, target_lang=target_lang)
    for outlier in heading.outliers:
        result.heading_outliers.append(
            HeadingOutlierFinding(
                unit_id=outlier.unit_key,
                target_text=outlier.target_text,
                majority_form=heading.majority_form or "",
                fix=heading.fixes.get(outlier.unit_key),
            )
        )
        record = result.ledger.get(outlier.unit_key)
        if record is not None:
            result.ledger.set_state(
                outlier.unit_key,
                STATE_NEEDS_REVIEW,
                evidence=f"heading form differs from majority ({heading.majority_form})",
            )
    return result


def replay_file(
    path: str | Path,
    *,
    target_lang: str,
    source_lang: str = "zh",
) -> ReplayResult:
    """对一份双语文档做回放体检。支持 .docx / .xlsx，按后缀分派。"""
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        kind = "docx"
        pairs = _pair_docx(file_path, source_lang=source_lang, target_lang=target_lang)
    elif suffix == ".xlsx":
        kind = "xlsx"
        pairs = _pair_xlsx(file_path, source_lang=source_lang, target_lang=target_lang)
    else:
        raise ValueError(f"unsupported file type for replay: {file_path.name}")
    result = ReplayResult(path=str(file_path), kind=kind, target_lang=target_lang)
    return replay_pairs(pairs, target_lang=target_lang, result=result)
